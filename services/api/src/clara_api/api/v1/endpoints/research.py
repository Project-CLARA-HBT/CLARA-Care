import asyncio
import hashlib
import json
import math
import os
import re
import secrets
import time
from collections.abc import Callable
from concurrent.futures import Future, ThreadPoolExecutor
from concurrent.futures import TimeoutError as FutureTimeoutError
from datetime import UTC, datetime, timedelta
from html import unescape
from html.parser import HTMLParser
from io import BytesIO
from threading import Lock, Thread
from typing import Any
from urllib.parse import quote, urljoin, urlparse
from uuid import uuid4

import httpx
from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile, status
from fastapi import Query as FastAPIQuery
from fastapi.responses import Response, StreamingResponse
from sqlalchemy import func, or_, select, update
from sqlalchemy.orm import Session

from clara_api.api.v1.endpoints.ml_proxy import proxy_ml_post
from clara_api.compliance.consent import PURPOSE_PERSONALIZATION, PURPOSE_RESEARCH
from clara_api.compliance.redaction import hash_user_ref
from clara_api.compliance.service import ComplianceService
from clara_api.core.attribution import (
    attach_attribution,
    build_attribution,
    normalize_source_errors,
    normalize_source_used,
)
from clara_api.core.config import get_settings
from clara_api.core.consent import PhrConsentService
from clara_api.core.control_tower import get_control_tower_config_service
from clara_api.core.control_tower.defaults import get_default_control_tower_config
from clara_api.core.flow_event_store import get_flow_event_store
from clara_api.core.rbac import require_roles
from clara_api.core.research_telemetry import sanitize_telemetry
from clara_api.core.research_upload_store import (
    ResearchUploadAuthorizationError,
    ResearchUploadNotFoundError,
    ResearchUploadStore,
    ResearchUploadStoreUnavailable,
)
from clara_api.core.security import TokenPayload
from clara_api.core.timeouts import resolve_sync_research_timeout
from clara_api.core.upload_safety import (
    UploadMalwareScannerUnavailable,
    UploadSafetyError,
    VerifiedUpload,
    read_upload_bytes_with_limit,
    verify_upload,
)
from clara_api.db.models import (
    FederatedSourceRecord,
    KnowledgeDocument,
    KnowledgeSource,
    MedicineCabinet,
    MedicineItem,
    PhrProfile,
    ResearchJob,
    SessionModel,
    SystemSetting,
    User,
    WorkspaceConversationShare,
)
from clara_api.db.models import (
    Query as QueryModel,
)
from clara_api.db.session import SessionLocal, get_db
from clara_api.observability.admin_audit import (
    ACTION_KB_DOCUMENT_STATUS,
    ACTION_KB_SOURCE_CREATE,
    ACTION_KB_SOURCE_UPLOAD,
    OUTCOME_FAILURE,
    OUTCOME_SUCCESS,
    record_admin_action,
)
from clara_api.phr.features import phr_features
from clara_api.schemas import (
    KnowledgeDocumentResponse,
    KnowledgeDocumentUpdateRequest,
    KnowledgeSourceCreateRequest,
    KnowledgeSourceResponse,
    KnowledgeSourceUpdateRequest,
    RagFlowConfig,
    ResearchClarifyQuestion,
    ResearchClarifyRequest,
    ResearchClarifyResponse,
    ResearchConversationCreateRequest,
    ResearchConversationListResponse,
    ResearchConversationMessageResponse,
    ResearchConversationMessagesResponse,
    ResearchConversationResponse,
    ResearchTier2JobCreateRequest,
    ResearchTier2JobResponse,
    ResearchTier2ShareResponse,
    SourceHubCatalogEntry,
    SourceHubRecord,
    SourceHubRecordsResponse,
    SourceHubSourceKey,
    SourceHubSyncRequest,
    SourceHubSyncResponse,
    WorkspaceConversationShareCreateRequest,
)

router = APIRouter()

_MAX_RESEARCH_UPLOADS = 200
_MAX_RESEARCH_UPLOAD_BYTES = 20 * 1024 * 1024
_PREVIEW_CHAR_LIMIT = 500
_MAX_EXTRACTED_TEXT_CHARS = 20_000
_DEFAULT_SOURCE_NAME = "General Uploads"
_TEXT_FILE_EXTENSIONS = {
    ".csv",
    ".json",
    ".log",
    ".markdown",
    ".md",
    ".txt",
    ".xml",
    ".yaml",
    ".yml",
}
_IMAGE_FILE_EXTENSIONS = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff", ".webp"}
_SOURCE_HUB_SETTING_KEY = "source_hub_records_v1"
_SOURCE_HUB_CATALOG_SETTING_KEY = "source_hub_catalog_v1"
_SOURCE_HUB_MAX_RECORDS = 500
_SOURCE_HUB_TIMEOUT_SECONDS = 12.0
_TRUE_VALUES = {"1", "true", "yes", "on"}
_SOURCE_HUB_SNIPPET_CHAR_LIMIT = 300
_DEFAULT_MARKDOWN_RENDER_HINTS: dict[str, Any] = {
    "markdown": True,
    "tables": True,
    "mermaid": False,
    "inline_references": False,
    "chart_spec_fences": [
        "chart-spec",
        "vega-lite",
        "echarts-option",
        "json",
        "yaml",
    ],
}
_VN_HTML_SOURCE_DEFINITIONS: dict[str, dict[str, Any]] = {
    "vn_moh": {
        "label": "Bộ Y tế Việt Nam",
        "docs_url": "https://moh.gov.vn/",
        "default_query": "huong dan chan doan dieu tri",
        "search_urls": (
            "https://moh.gov.vn/tim-kiem?q={query_q}",
            "https://moh.gov.vn/",
        ),
    },
    "vn_kcb": {
        "label": "Cục Quản lý Khám chữa bệnh",
        "docs_url": "https://kcb.vn/",
        "default_query": "huong dan kham chua benh",
        "search_urls": (
            "https://kcb.vn/?s={query_q}",
            "https://kcb.vn/",
        ),
    },
    "vn_canhgiacduoc": {
        "label": "Trung tâm Quốc gia về Thông tin thuốc và Theo dõi phản ứng có hại của thuốc",
        "docs_url": "https://canhgiacduoc.org.vn/",
        "default_query": "canh giac duoc ADR",
        "search_urls": (
            "https://canhgiacduoc.org.vn/?s={query_q}",
            "https://canhgiacduoc.org.vn/",
        ),
    },
    "vn_vbpl_byt": {
        "label": "VBPL Bộ Y tế",
        "docs_url": "https://vbpl.vn/boyte/Pages/home.aspx",
        "default_query": "thong tu bo y te",
        "search_urls": (
            "https://vbpl.vn/boyte/pages/default.aspx?keyword={query_q}",
            "https://vbpl.vn/boyte/Pages/home.aspx",
        ),
    },
    "vn_dav": {
        "label": "Cục Quản lý Dược Việt Nam",
        "docs_url": "https://dav.gov.vn/",
        "default_query": "thu hoi thuoc",
        "search_urls": (
            "https://dav.gov.vn/?s={query_q}",
            "https://dav.gov.vn/",
        ),
    },
}

_uploaded_research_files: dict[str, dict[str, Any]] = {}
_uploaded_research_lock = Lock()
_research_settings = get_settings()
_RESEARCH_JOB_MAX_WORKERS = max(
    1,
    min(32, int(_research_settings.research_job_max_workers)),
)
_RESEARCH_JOB_MAX_PENDING = max(
    1,
    min(2_000, int(_research_settings.research_job_max_pending)),
)
_RESEARCH_JOB_MAX_ACTIVE_PER_USER = max(
    1,
    min(100, int(_research_settings.research_job_max_active_per_user)),
)
_research_job_executor = ThreadPoolExecutor(
    max_workers=_RESEARCH_JOB_MAX_WORKERS,
    thread_name_prefix="research-tier2",
)
_research_job_futures: dict[str, Future[Any]] = {}
_research_job_lock = Lock()
_research_recovery_lock = Lock()
_research_recovery_started = False
_RESEARCH_JOB_LEASE_SECONDS = 120
_RESEARCH_JOB_RECOVERY_POLL_SECONDS = 30
_RESEARCH_MODE_ALLOWED = {"fast", "deep", "deep_beta"}
_RETRIEVAL_STACK_MODE_ALLOWED = {"auto", "full"}
_ANSWER_LANGUAGE_ALLOWED = {"vi", "en"}
_PROVIDER_SECRET_KEYS = frozenset(
    {
        "api_key",
        "apikey",
        "access_token",
        "client_secret",
        "password",
        "secret",
        "token",
        "authorization",
        "private_key",
    }
)


def _validate_upload_safety(
    *, file_name: str, content_type: str, file_bytes: bytes
) -> VerifiedUpload:
    """Apply the shared deterministic content and optional ClamAV boundary."""

    settings = get_settings()
    try:
        return verify_upload(
            filename=file_name,
            content_type=content_type,
            data=file_bytes,
            fallback_filename="uploaded-file",
            malware_scan_required=settings.upload_malware_scan_required,
            clamav_host=settings.upload_malware_clamav_host,
            clamav_port=settings.upload_malware_clamav_port,
        )
    except UploadMalwareScannerUnavailable as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Không thể kiểm tra an toàn tệp lúc này. Vui lòng thử lại sau.",
        ) from exc
    except UploadSafetyError as exc:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Tệp tải lên không khớp định dạng được phép.",
        ) from exc


def _load_research_rag_runtime(db: Session) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    try:
        control_tower = get_control_tower_config_service().load(db)
        return control_tower.rag_flow.model_dump(), [
            item.model_dump() for item in control_tower.rag_sources
        ]
    except Exception:  # pragma: no cover - defensive path for runtime resilience
        fallback = get_default_control_tower_config()
        return fallback.rag_flow.model_dump(), [item.model_dump() for item in fallback.rag_sources]


def _guess_extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return f".{filename.rsplit('.', 1)[-1].lower()}"


def _is_text_file(filename: str, content_type: str) -> bool:
    extension = _guess_extension(filename)
    if content_type.startswith("text/"):
        return True
    if extension in _TEXT_FILE_EXTENSIONS:
        return True
    return content_type in {"application/json", "application/xml"}


def _is_image_file(filename: str, content_type: str) -> bool:
    if content_type.startswith("image/"):
        return True
    return _guess_extension(filename) in _IMAGE_FILE_EXTENSIONS


def _is_pdf_file(filename: str, content_type: str) -> bool:
    return content_type == "application/pdf" or _guess_extension(filename) == ".pdf"


def _decode_text_payload(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("utf-8", errors="replace")


def _normalize_extracted_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+\n", "\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()[:_MAX_EXTRACTED_TEXT_CHARS]


def _extract_text_like_file(file_bytes: bytes, filename: str, content_type: str) -> str:
    decoded = _decode_text_payload(file_bytes)
    extension = _guess_extension(filename)
    if content_type == "application/json" or extension == ".json":
        try:
            parsed = json.loads(decoded)
        except json.JSONDecodeError:
            return decoded
        try:
            return json.dumps(parsed, ensure_ascii=False, indent=2)
        except (TypeError, ValueError):
            return _normalize_extracted_text(decoded)
    return _normalize_extracted_text(decoded)


def _extract_pdf_text(file_bytes: bytes) -> tuple[str, str]:
    try:
        from pypdf import PdfReader
    except Exception:
        return "", "PDF đã tải lên nhưng chưa thể trích xuất text vì thiếu parser `pypdf`."

    try:
        reader = PdfReader(BytesIO(file_bytes))
    except Exception as exc:
        return "", f"PDF đã được tải lên nhưng không đọc được nội dung ({exc.__class__.__name__})."

    pages_text: list[str] = []
    for page in reader.pages[:20]:
        try:
            extracted = page.extract_text() or ""
        except Exception:
            extracted = ""
        normalized = _normalize_extracted_text(extracted)
        if normalized:
            pages_text.append(normalized)

    if not pages_text:
        return "", "PDF đã tải lên nhưng không trích xuất được text hữu ích."
    return _normalize_extracted_text("\n\n".join(pages_text)), ""


def _extract_image_metadata_text(file_bytes: bytes, filename: str) -> tuple[str, str]:
    try:
        from PIL import Image, UnidentifiedImageError
    except Exception:
        return "", "Ảnh đã tải lên nhưng môi trường chưa có parser metadata."

    try:
        with Image.open(BytesIO(file_bytes)) as image_obj:
            width, height = image_obj.size
            image_format = (image_obj.format or _guess_extension(filename).lstrip(".")).upper()
            mode = image_obj.mode or "unknown"
            frame_count = int(getattr(image_obj, "n_frames", 1) or 1)
            parts = [
                f"image_format={image_format or 'UNKNOWN'}",
                f"size={width}x{height}",
                f"mode={mode}",
            ]
            if frame_count > 1:
                parts.append(f"frames={frame_count}")
            return _normalize_extracted_text("Image metadata: " + ", ".join(parts)), ""
    except (UnidentifiedImageError, OSError):
        return "", "Ảnh đã tải lên nhưng không đọc được metadata."
    except Exception as exc:
        return "", f"Ảnh đã tải lên nhưng parse metadata thất bại ({exc.__class__.__name__})."


def _is_truthy_env(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() in _TRUE_VALUES


def _extract_image_text_with_ocr(
    file_bytes: bytes,
    filename: str,
    content_type: str,
) -> tuple[str, str]:
    if not _is_truthy_env("RESEARCH_UPLOAD_IMAGE_OCR", default=True):
        return "", ""

    try:
        from clara_api.api.v1.endpoints.careguard import _scan_with_tgc_ocr
    except Exception:
        return "", ""

    try:
        extracted_text, _used_endpoint, _provider = _scan_with_tgc_ocr(
            file_bytes=file_bytes,
            file_name=filename,
            content_type=content_type,
        )
    except HTTPException as exc:
        detail = str(exc.detail).strip() if isinstance(exc.detail, str) else ""
        if detail:
            return "", f"Ảnh đã tải lên nhưng OCR không khả dụng: {detail}"
        return "", "Ảnh đã tải lên nhưng OCR không khả dụng."
    except Exception:
        return "", "Ảnh đã tải lên nhưng OCR không khả dụng."

    normalized = _normalize_extracted_text(extracted_text)
    if not normalized:
        return "", "Ảnh đã tải lên nhưng OCR không trích xuất được text hữu ích."
    return normalized, ""


def _extract_basic_text(file_bytes: bytes, filename: str, content_type: str) -> tuple[str, str]:
    if _is_text_file(filename, content_type):
        text = _extract_text_like_file(file_bytes, filename, content_type)
        return text, "text"

    if _is_pdf_file(filename, content_type):
        extracted_pdf_text, pdf_message = _extract_pdf_text(file_bytes)
        if extracted_pdf_text:
            return extracted_pdf_text, "text"
        return pdf_message, "pdf"

    if _is_image_file(filename, content_type):
        metadata_text, metadata_message = _extract_image_metadata_text(file_bytes, filename)
        extracted_image_text, image_message = _extract_image_text_with_ocr(
            file_bytes,
            filename,
            content_type,
        )
        if extracted_image_text and metadata_text:
            merged = _normalize_extracted_text(
                f"{metadata_text}\n\nOCR text:\n{extracted_image_text}"
            )
            return merged, "text"
        if extracted_image_text:
            return extracted_image_text, "text"
        if image_message:
            return image_message, "image"
        if metadata_text:
            return metadata_text, "text"
        return metadata_message, "image"

    return "File đã tải lên. Định dạng này chưa hỗ trợ trích xuất text tự động.", "other"


def _approx_token_count(text: str) -> int:
    stripped = text.strip()
    if not stripped:
        return 0
    return max(1, math.ceil(len(stripped) / 4))


def _store_uploaded_file(entry: dict[str, Any]) -> None:
    with _uploaded_research_lock:
        _uploaded_research_files[entry["file_id"]] = entry
        if len(_uploaded_research_files) <= _MAX_RESEARCH_UPLOADS:
            return

        oldest_file_id = min(
            _uploaded_research_files,
            key=lambda item_file_id: str(_uploaded_research_files[item_file_id]["created_at"]),
        )
        _uploaded_research_files.pop(oldest_file_id, None)


def _durable_uploads_enabled() -> bool:
    """Return True when the DB-backed durable upload store should be used (R2)."""

    return bool(get_settings().research_durable_uploads_enabled)


def _build_research_upload_store(db: Session) -> ResearchUploadStore:
    """Construct a durable upload store bound to the request DB session."""

    object_store_url = (get_settings().research_upload_object_store_url or "").strip()
    return ResearchUploadStore(db, object_store_url=object_store_url or None)


def _build_uploaded_documents_durable(
    uploaded_file_ids: list[Any],
    *,
    owner_user_id: int,
    db: Session,
) -> list[dict[str, Any]]:
    """Resolve uploaded documents from the durable, owner-isolated store (R2).

    A referenced ``file_id`` not owned by the requester raises a 403 and is
    excluded from the job (R2.4). An unknown ``file_id`` is skipped (mirrors the
    legacy in-memory behavior). If the configured backend is unavailable, a 503
    is raised so uploads are never silently dropped (R2.5).
    """

    store = _build_research_upload_store(db)
    documents: list[dict[str, Any]] = []
    for raw_file_id in uploaded_file_ids:
        if not isinstance(raw_file_id, str):
            continue
        try:
            stored = store.get(raw_file_id, owner_user_id)
        except ResearchUploadNotFoundError:
            continue
        except ResearchUploadAuthorizationError as exc:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Bạn không sở hữu file đã tải lên được tham chiếu.",
            ) from exc
        except ResearchUploadStoreUnavailable as exc:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Kho lưu trữ file research tạm thời không khả dụng.",
            ) from exc
        documents.append(stored.as_document())
    return documents


def _build_uploaded_documents(
    uploaded_file_ids: Any,
    *,
    owner_user_id: int,
    db: Session | None = None,
) -> list[dict[str, Any]]:
    if not isinstance(uploaded_file_ids, list):
        return []

    # Durable, owner-isolated backend behind the flag (R2). Falls back to the
    # in-memory dict when the flag is off so legacy behavior is preserved.
    if db is not None and _durable_uploads_enabled():
        return _build_uploaded_documents_durable(
            uploaded_file_ids,
            owner_user_id=owner_user_id,
            db=db,
        )

    documents: list[dict[str, Any]] = []
    with _uploaded_research_lock:
        for raw_file_id in uploaded_file_ids:
            if not isinstance(raw_file_id, str):
                continue
            cached = _uploaded_research_files.get(raw_file_id)
            if not cached:
                continue
            if int(cached.get("owner_user_id") or 0) != int(owner_user_id):
                continue

            documents.append(
                {
                    "file_id": raw_file_id,
                    "filename": cached["filename"],
                    "content_type": cached["content_type"],
                    "size": cached["size"],
                    "created_at": cached["created_at"],
                    "text": cached["text"],
                    "preview": cached["preview"],
                    "token_count": cached["token_count"],
                }
            )
    return documents


def _serialize_knowledge_document(document: KnowledgeDocument) -> KnowledgeDocumentResponse:
    return KnowledgeDocumentResponse(
        id=document.id,
        source_id=document.source_id,
        filename=document.filename,
        content_type=document.content_type,
        size=document.size,
        preview=document.preview,
        token_count=document.token_count,
        is_active=document.is_active,
        created_at=document.created_at,
        updated_at=document.updated_at,
    )


def _serialize_knowledge_source(
    source: KnowledgeSource,
    *,
    documents_count: int,
) -> KnowledgeSourceResponse:
    return KnowledgeSourceResponse(
        id=source.id,
        name=source.name,
        description=source.description,
        is_active=source.is_active,
        created_at=source.created_at,
        updated_at=source.updated_at,
        documents_count=documents_count,
    )


def _get_user_by_token(db: Session, token: TokenPayload) -> User:
    user = db.execute(select(User).where(User.email == token.sub)).scalar_one_or_none()
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED, detail="Người dùng không tồn tại"
        )
    return user


def _normalize_research_mode_value(raw_mode: Any, *, default: str = "fast") -> str:
    normalized = str(raw_mode or "").strip().lower().replace("-", "_")
    if normalized in {"deep", "deep_research", "long"}:
        return "deep"
    if normalized in {"deep_beta", "deepbeta"}:
        return "deep_beta"
    if normalized in _RESEARCH_MODE_ALLOWED:
        return normalized
    return default


def _normalize_retrieval_stack_mode_value(raw_mode: Any, *, default: str = "auto") -> str:
    normalized = str(raw_mode or "").strip().lower().replace("-", "_")
    if normalized in _RETRIEVAL_STACK_MODE_ALLOWED:
        return normalized
    return default


def _normalize_answer_language_value(raw_language: Any, *, default: str = "vi") -> str:
    normalized = str(raw_language or "").strip().lower().replace("-", "_")
    if normalized in _ANSWER_LANGUAGE_ALLOWED:
        return normalized
    return default


def _canonicalize_research_payload_contract(payload: dict[str, Any]) -> dict[str, Any]:
    normalized = dict(payload)
    metadata_obj = normalized.get("metadata")
    if not isinstance(metadata_obj, dict):
        metadata_obj = None

    has_mode_signal = (
        ("research_mode" in normalized)
        or ("mode" in normalized)
        or (metadata_obj is not None and metadata_obj.get("research_mode") is not None)
    )
    if has_mode_signal:
        mode = _normalize_research_mode_value(
            normalized.get("research_mode")
            or normalized.get("mode")
            or (metadata_obj.get("research_mode") if metadata_obj is not None else None),
            default="fast",
        )
        normalized["research_mode"] = mode
        if metadata_obj is not None:
            metadata_obj["research_mode"] = mode

    has_stack_mode_signal = (
        ("retrieval_stack_mode" in normalized)
        or ("stack_mode" in normalized)
        or (
            metadata_obj is not None
            and (
                metadata_obj.get("retrieval_stack_mode") is not None
                or metadata_obj.get("stack_mode") is not None
            )
        )
    )
    if has_stack_mode_signal:
        stack_mode = _normalize_retrieval_stack_mode_value(
            normalized.get("retrieval_stack_mode")
            or normalized.get("stack_mode")
            or (metadata_obj.get("retrieval_stack_mode") if metadata_obj is not None else None)
            or (metadata_obj.get("stack_mode") if metadata_obj is not None else None),
            default="auto",
        )
        normalized["retrieval_stack_mode"] = stack_mode
        if metadata_obj is not None:
            metadata_obj["retrieval_stack_mode"] = stack_mode

    fallback_reason = normalized.get("fallback_reason")
    if (not isinstance(fallback_reason, str) or not fallback_reason.strip()) and (
        metadata_obj is not None
    ):
        metadata_reason = metadata_obj.get("fallback_reason")
        if isinstance(metadata_reason, str) and metadata_reason.strip():
            fallback_reason = metadata_reason.strip()
            normalized["fallback_reason"] = fallback_reason
        else:
            fallback_reason = ""
    elif isinstance(fallback_reason, str):
        fallback_reason = fallback_reason.strip()
        if fallback_reason:
            normalized["fallback_reason"] = fallback_reason
    else:
        fallback_reason = ""

    has_fallback_signal = any(
        key in normalized for key in ("fallback", "fallback_used", "fallback_reason")
    ) or (
        metadata_obj is not None
        and any(key in metadata_obj for key in ("fallback", "fallback_used", "fallback_reason"))
    )
    if has_fallback_signal:
        fallback_used = bool(
            normalized.get("fallback_used")
            or (metadata_obj.get("fallback_used") if metadata_obj is not None else False)
            or normalized.get("fallback")
            or fallback_reason
        )
        normalized["fallback_used"] = fallback_used
        if "fallback" in normalized or fallback_used:
            normalized["fallback"] = fallback_used
        if metadata_obj is not None:
            metadata_obj["fallback_used"] = fallback_used
            if fallback_reason:
                metadata_obj["fallback_reason"] = fallback_reason

    return normalized


def _coerce_stored_result(raw_text: str) -> dict[str, Any]:
    stripped = raw_text.strip()
    if not stripped:
        return {"tier": "tier1", "answer": ""}

    try:
        parsed = json.loads(stripped)
    except json.JSONDecodeError:
        return {"tier": "tier1", "answer": stripped}

    if not isinstance(parsed, dict):
        return {"tier": "tier1", "answer": stripped}

    payload = parsed.get("result")
    result = payload if isinstance(payload, dict) else parsed
    if "tier" not in result:
        if any(
            key in result
            for key in (
                "citations",
                "flowEvents",
                "flow_events",
                "telemetry",
                "research_mode",
                "deep_pass_count",
                "source_attempts",
                "source_errors",
                "fallback_reason",
                "query_plan",
            )
        ):
            result = {"tier": "tier2", **result}
        else:
            result = {"tier": "tier1", **result}
    tier = str(result.get("tier") or "").strip().lower()
    if tier == "tier2":
        return _canonicalize_research_payload_contract(result)
    return result


def _serialize_research_conversation(
    *,
    session_obj: SessionModel,
    query_obj: QueryModel,
) -> ResearchConversationResponse:
    result_payload = _coerce_stored_result(query_obj.response_text)
    tier = str(result_payload.get("tier") or "tier1").strip().lower()
    if tier not in {"tier1", "tier2"}:
        tier = "tier1"
        result_payload["tier"] = tier

    created_at = query_obj.created_at or session_obj.created_at
    if created_at is None:
        created_at = datetime.now(tz=UTC)

    return ResearchConversationResponse(
        id=session_obj.id,
        query_id=query_obj.id,
        query=query_obj.user_input,
        result=result_payload,
        tier=tier,
        created_at=created_at,
    )


def _serialize_research_message(query_obj: QueryModel) -> ResearchConversationMessageResponse:
    result_payload = _coerce_stored_result(query_obj.response_text)
    tier = str(result_payload.get("tier") or "tier1").strip().lower()
    if tier not in {"tier1", "tier2"}:
        tier = "tier1"
        result_payload["tier"] = tier

    created_at = query_obj.created_at or datetime.now(tz=UTC)
    return ResearchConversationMessageResponse(
        query_id=query_obj.id,
        query=query_obj.user_input,
        tier=tier,  # type: ignore[arg-type]
        result=result_payload,
        created_at=created_at,
    )


def _validate_result_payload(result: dict[str, Any]) -> dict[str, Any]:
    payload = dict(result)
    tier = str(payload.get("tier") or "").strip().lower()
    if tier not in {"tier1", "tier2"}:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="result.tier phải là 'tier1' hoặc 'tier2'.",
        )
    payload["tier"] = tier
    if tier == "tier2":
        payload = _canonicalize_research_payload_contract(payload)
    return payload


def _get_owned_source(db: Session, *, source_id: int, owner_user_id: int) -> KnowledgeSource:
    source = db.execute(
        select(KnowledgeSource).where(
            KnowledgeSource.id == source_id,
            KnowledgeSource.owner_user_id == owner_user_id,
        )
    ).scalar_one_or_none()
    if not source:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Knowledge source không tồn tại"
        )
    return source


def _get_owned_document(db: Session, *, document_id: int, owner_user_id: int) -> KnowledgeDocument:
    document = db.execute(
        select(KnowledgeDocument).where(
            KnowledgeDocument.id == document_id,
            KnowledgeDocument.owner_user_id == owner_user_id,
        )
    ).scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document không tồn tại")
    return document


def _get_or_create_default_source(db: Session, owner_user_id: int) -> KnowledgeSource:
    source = db.execute(
        select(KnowledgeSource).where(
            KnowledgeSource.owner_user_id == owner_user_id,
            KnowledgeSource.name == _DEFAULT_SOURCE_NAME,
        )
    ).scalar_one_or_none()
    if source:
        return source

    source = KnowledgeSource(
        owner_user_id=owner_user_id,
        name=_DEFAULT_SOURCE_NAME,
        description="Nguồn mặc định cho upload nhanh từ màn hình chat/research",
        is_active=True,
    )
    db.add(source)
    db.commit()
    db.refresh(source)
    return source


def _build_source_documents(
    db: Session,
    *,
    owner_user_id: int,
    source_ids: list[int],
) -> list[dict[str, Any]]:
    if not source_ids:
        return []

    source_ids_set = sorted(set(source_ids))
    source_rows = (
        db.execute(
            select(KnowledgeSource.id).where(
                KnowledgeSource.owner_user_id == owner_user_id,
                KnowledgeSource.id.in_(source_ids_set),
                KnowledgeSource.is_active.is_(True),
            )
        )
        .scalars()
        .all()
    )
    valid_source_ids = set(source_rows)
    if not valid_source_ids:
        return []

    documents = (
        db.execute(
            select(KnowledgeDocument).where(
                KnowledgeDocument.owner_user_id == owner_user_id,
                KnowledgeDocument.source_id.in_(sorted(valid_source_ids)),
                KnowledgeDocument.is_active.is_(True),
            )
        )
        .scalars()
        .all()
    )

    return [
        {
            "file_id": f"knowledge-doc-{document.id}",
            "filename": document.filename,
            "content_type": document.content_type,
            "size": document.size,
            "created_at": document.created_at.isoformat(),
            "text": document.extracted_text,
            "preview": document.preview,
            "token_count": document.token_count,
            "source": f"knowledge-source-{document.source_id}",
        }
        for document in documents
    ]


def _extract_source_hub_sources(payload: dict[str, Any]) -> set[str]:
    raw = payload.get("source_hub_sources")
    values: list[str] = []
    if isinstance(raw, list):
        values = [str(item).strip().lower() for item in raw if str(item).strip()]
    elif isinstance(raw, str) and raw.strip():
        values = [raw.strip().lower()]
    allowed = {
        "pubmed",
        "europepmc",
        "semantic_scholar",
        "clinicaltrials",
        "rxnorm",
        "openfda",
        "dailymed",
        "vn_moh",
        "vn_kcb",
        "vn_canhgiacduoc",
        "vn_vbpl_byt",
        "vn_dav",
        "davidrug",
    }
    return {item for item in values if item in allowed}


def _build_source_hub_documents(
    db: Session,
    *,
    owner_user_id: int,
    query: str,
    source_filters: set[str],
    limit: int = 40,
) -> list[dict[str, Any]]:
    records = _load_source_hub_records(db, owner_user_id)
    query_terms = {term.strip().lower() for term in query.split() if len(term.strip()) >= 3}

    matched: list[SourceHubRecord] = []
    for record in records:
        if source_filters and record.source not in source_filters:
            continue
        haystack = " ".join(
            part for part in [record.title or "", record.snippet or "", record.query or ""] if part
        ).lower()
        if query_terms and not any(term in haystack for term in query_terms):
            continue
        matched.append(record)
        if len(matched) >= max(1, int(limit)):
            break

    docs: list[dict[str, Any]] = []
    for index, record in enumerate(matched, start=1):
        text_parts = [record.title]
        if record.snippet:
            text_parts.append(record.snippet)
        if record.metadata:
            compact_meta = ", ".join(
                f"{key}={value}"
                for key, value in list(record.metadata.items())[:4]
                if value not in (None, "", [])
            )
            if compact_meta:
                text_parts.append(compact_meta)
        docs.append(
            {
                "file_id": f"source-hub-{record.id}",
                "filename": f"{record.source}-{index}",
                "content_type": "text/plain",
                "size": 0,
                "created_at": record.synced_at or datetime.now(tz=UTC).isoformat(),
                "text": " | ".join(part for part in text_parts if part),
                "preview": (record.snippet or record.title or "")[:_PREVIEW_CHAR_LIMIT],
                "token_count": _approx_token_count(" ".join(text_parts)),
                "source": f"source_hub_{record.source}",
                "url": record.url,
            }
        )
    return docs


def _extract_source_ids(payload: dict[str, Any]) -> list[int]:
    raw_sources: list[Any] = []
    for key in ("source_ids", "knowledge_source_ids"):
        value = payload.get(key)
        if isinstance(value, list):
            raw_sources.extend(value)

    parsed: list[int] = []
    for item in raw_sources:
        if isinstance(item, int):
            parsed.append(item)
            continue
        if isinstance(item, str) and item.strip().isdigit():
            parsed.append(int(item.strip()))
    return parsed


def _coerce_research_mode(payload: dict[str, Any]) -> str:
    return _normalize_research_mode_value(
        payload.get("research_mode") or payload.get("mode"),
        default="fast",
    )


def _resolve_research_output_mode(
    value: Any,
    *,
    role: str | None,
) -> str:
    """Return a role-safe closed presentation mode.

    This selector is deliberately not a research, retrieval, verifier, or
    model-routing input. Consumers always receive plain language. Professional
    presentation is limited to the existing professional roles and remains
    subject to the same evidence-release boundary as the base answer.
    """

    requested = str(value or "plain_language").strip().lower()
    if requested != "professional":
        return "plain_language"
    if str(role or "").strip().lower() not in {"researcher", "doctor", "admin"}:
        return "plain_language"
    return "professional"


def _coerce_retrieval_stack_mode(payload: dict[str, Any]) -> str:
    return _normalize_retrieval_stack_mode_value(
        payload.get("retrieval_stack_mode") or payload.get("stack_mode"),
        default="auto",
    )


def _resolve_tier2_execution_modes(payload: dict[str, Any]) -> tuple[str, str]:
    research_mode = _coerce_research_mode(payload)
    retrieval_stack_mode = _coerce_retrieval_stack_mode(payload)
    if research_mode == "fast" and retrieval_stack_mode == "full":
        retrieval_stack_mode = "auto"
    return research_mode, retrieval_stack_mode


def _extract_tier2_query_text(payload: dict[str, Any]) -> str:
    for key in ("query", "message", "question"):
        value = payload.get(key)
        if isinstance(value, str):
            text = value.strip()
            if text:
                return text
    return ""


# Clarifying-questions support (clara-research R12). Ambiguity detection is a deterministic,
# API-side heuristic: deep research benefits from a well-scoped question, so an underspecified
# query (too few content words) is treated as ambiguous and triggers clarifying questions.
_CLARIFY_MIN_CONTENT_WORDS = 4
_CLARIFY_DEEP_MODES = {"deep", "deep_beta"}


def _detect_query_ambiguity(query: str) -> bool:
    """Return True when the query is underspecified enough to warrant clarification.

    A query is ambiguous when it carries fewer than ``_CLARIFY_MIN_CONTENT_WORDS`` content
    words (tokens of length >= 2). This keeps the gate deterministic and conservative: long,
    specific queries start immediately while short/vague prompts ask for scope first.
    """
    content_words = [token for token in re.split(r"\W+", query.lower()) if len(token) >= 2]
    return len(content_words) < _CLARIFY_MIN_CONTENT_WORDS


def _build_clarifying_questions(*, ui_language: str) -> list[ResearchClarifyQuestion]:
    """Build the curated clarifying-question set, localized to the requested UI language."""
    if ui_language == "en":
        specs = [
            (
                "population",
                "Who is this for (patient profile, age group, or condition)?",
                "Scoping the population focuses retrieval on the most relevant evidence.",
            ),
            (
                "scope",
                "What specifically would you like to compare or learn about?",
                "Narrowing the scope avoids an overly broad, unfocused report.",
            ),
            (
                "outcome",
                "What outcome or decision are you trying to support?",
                "The target outcome shapes which evidence and recommendations matter.",
            ),
        ]
    else:
        specs = [
            (
                "population",
                "Câu hỏi này dành cho đối tượng nào (hồ sơ bệnh nhân, nhóm tuổi, hoặc bệnh lý)?",
                "Xác định đối tượng giúp truy xuất đúng bằng chứng phù hợp nhất.",
            ),
            (
                "scope",
                "Bạn muốn so sánh hoặc tìm hiểu cụ thể về điều gì?",
                "Thu hẹp phạm vi tránh báo cáo quá rộng và thiếu trọng tâm.",
            ),
            (
                "outcome",
                "Bạn đang cần hỗ trợ cho kết quả hoặc quyết định nào?",
                "Kết quả mục tiêu quyết định bằng chứng và khuyến nghị nào là quan trọng.",
            ),
        ]
    return [
        ResearchClarifyQuestion(id=question_id, question=question, rationale=rationale)
        for question_id, question, rationale in specs
    ]


def _first_dict(*values: Any) -> dict[str, Any] | None:
    for value in values:
        if isinstance(value, dict):
            return value
    return None


def _first_value(
    sources: list[dict[str, Any] | None],
    *,
    keys: tuple[str, ...],
) -> Any:
    for source in sources:
        if source is None:
            continue
        for key in keys:
            if key in source and source[key] is not None:
                return source[key]
    return None


_VERIFICATION_MATRIX_KEYS: tuple[str, ...] = (
    "verification_matrix",
    "claim_verification_matrix",
    "claim_matrix",
    "claims_matrix",
)
_CONTRADICTION_SUMMARY_KEYS: tuple[str, ...] = (
    "contradiction_summary",
    "contradictions_summary",
    "contradiction_report",
    "contradiction_overview",
)
_TRACE_CONTAINER_KEYS: tuple[str, ...] = (
    "trace_metadata",
    "trace_context",
    "otel_trace_metadata",
    "otel_trace_context",
    "otel_trace",
    "trace",
    "otel",
)
_TRACE_SCALAR_KEYS: tuple[str, ...] = (
    "trace_id",
    "span_id",
    "parent_span_id",
    "trace_flags",
    "trace_state",
    "tracestate",
    "traceparent",
    "sampled",
    "service_name",
    "service",
    "component",
)
_TRACE_HINT_KEYS: set[str] = {key.lower() for key in (*_TRACE_CONTAINER_KEYS, *_TRACE_SCALAR_KEYS)}


def _is_trace_key(key: str) -> bool:
    normalized = key.strip().lower()
    if not normalized:
        return False
    return (
        normalized in _TRACE_HINT_KEYS
        or normalized.startswith("trace")
        or normalized.startswith("otel")
    )


def _normalize_trace_value(value: Any) -> Any:
    if isinstance(value, str | int | float | bool):
        return value
    if isinstance(value, list):
        compact = [item for item in value if isinstance(item, str | int | float | bool)]
        return compact[:20] if compact else None
    if isinstance(value, dict):
        compact: dict[str, Any] = {}
        for raw_key, nested_value in value.items():
            normalized_nested = _normalize_trace_value(nested_value)
            if normalized_nested is None:
                continue
            compact[str(raw_key)] = normalized_nested
        return compact or None
    return None


def _extract_trace_metadata(
    sources: list[dict[str, Any] | None],
) -> dict[str, Any] | None:
    trace_metadata: dict[str, Any] = {}

    for source in sources:
        if source is None:
            continue

        for container_key in _TRACE_CONTAINER_KEYS:
            candidate = source.get(container_key)
            if not isinstance(candidate, dict):
                continue
            for raw_key, raw_value in candidate.items():
                key_text = str(raw_key)
                if not _is_trace_key(key_text):
                    continue
                normalized_value = _normalize_trace_value(raw_value)
                if normalized_value is None:
                    continue
                trace_metadata.setdefault(key_text, normalized_value)

        for scalar_key in _TRACE_SCALAR_KEYS:
            raw_value = source.get(scalar_key)
            normalized_value = _normalize_trace_value(raw_value)
            if normalized_value is None:
                continue
            trace_metadata.setdefault(scalar_key, normalized_value)

    return trace_metadata or None


def _build_tier2_telemetry(
    *,
    normalized: dict[str, Any],
    metadata_obj: dict[str, Any] | None,
    context_debug_obj: dict[str, Any] | None,
) -> dict[str, Any] | None:
    telemetry_root = _first_dict(
        normalized.get("telemetry"),
        metadata_obj.get("telemetry") if metadata_obj else None,
        normalized.get("debug_telemetry"),
        metadata_obj.get("debug_telemetry") if metadata_obj else None,
        context_debug_obj.get("telemetry") if context_debug_obj else None,
    )
    telemetry = dict(telemetry_root) if telemetry_root else {}
    retrieval_trace = (
        context_debug_obj.get("retrieval_trace")
        if context_debug_obj and isinstance(context_debug_obj.get("retrieval_trace"), dict)
        else {}
    )
    retriever_debug = (
        retrieval_trace.get("retriever_debug")
        if isinstance(retrieval_trace.get("retriever_debug"), dict)
        else {}
    )

    debug_obj = _first_dict(
        normalized.get("debug"),
        metadata_obj.get("debug") if metadata_obj else None,
    )
    trace_metadata = _extract_trace_metadata(
        [
            telemetry,
            normalized,
            metadata_obj,
            context_debug_obj,
            retrieval_trace if isinstance(retrieval_trace, dict) else None,
            debug_obj,
        ]
    )
    sources: list[dict[str, Any] | None] = [
        telemetry,
        normalized,
        metadata_obj,
        context_debug_obj,
        debug_obj,
    ]

    if "keywords" not in telemetry:
        keywords = _first_value(
            sources,
            keys=(
                "keywords",
                "query_keywords",
                "keyword_list",
                "matched_keywords",
                "intent_keywords",
            ),
        )
        if keywords is not None:
            telemetry["keywords"] = keywords

    if "search_plan" not in telemetry:
        search_plan = _first_value(
            sources,
            keys=("search_plan", "search_trace", "query_plan"),
        )
        if search_plan is None and isinstance(retrieval_trace, dict):
            search_plan = retrieval_trace.get("search_plan")
        if search_plan is not None:
            telemetry["search_plan"] = search_plan

    if "query_plan" not in telemetry:
        query_plan = _first_value(
            sources,
            keys=("query_plan", "search_plan", "search_trace"),
        )
        if query_plan is None and isinstance(retrieval_trace, dict):
            query_plan = retrieval_trace.get("query_plan")
            if query_plan is None:
                query_plan = retrieval_trace.get("search_plan")
        if query_plan is not None:
            telemetry["query_plan"] = query_plan

    if "query_plan" not in telemetry and "search_plan" in telemetry:
        telemetry["query_plan"] = telemetry.get("search_plan")
    if "search_plan" not in telemetry and "query_plan" in telemetry:
        telemetry["search_plan"] = telemetry.get("query_plan")

    if "source_attempts" not in telemetry:
        source_attempts = _first_value(
            sources,
            keys=(
                "source_attempts",
                "connector_attempts",
                "provider_events",
                "retrieval_attempts",
            ),
        )
        if source_attempts is None and isinstance(retriever_debug, dict):
            source_attempts = retriever_debug.get("source_attempts")
            if source_attempts is None:
                source_attempts = retriever_debug.get("provider_events")
        if source_attempts is not None:
            telemetry["source_attempts"] = source_attempts

    if "index_summary" not in telemetry:
        index_summary = _first_value(
            sources,
            keys=("index_summary", "rerank_summary", "ranking_summary"),
        )
        if index_summary is None and isinstance(retrieval_trace, dict):
            index_summary = retrieval_trace.get("index_summary")
        if index_summary is None and isinstance(retriever_debug, dict):
            index_summary = retriever_debug.get("index_summary")
        if index_summary is not None:
            telemetry["index_summary"] = index_summary

    if "crawl_summary" not in telemetry:
        crawl_summary = _first_value(
            sources,
            keys=("crawl_summary", "web_crawl_summary", "crawl_trace"),
        )
        if crawl_summary is None and isinstance(retrieval_trace, dict):
            crawl_summary = retrieval_trace.get("crawl_summary")
        if crawl_summary is None and isinstance(retriever_debug, dict):
            crawl_summary = retriever_debug.get("crawl_summary")
        if crawl_summary is not None:
            telemetry["crawl_summary"] = crawl_summary

    if "docs" not in telemetry:
        docs = _first_value(
            sources,
            keys=(
                "docs",
                "documents",
                "retrieved_docs",
                "retrieved_context",
                "context_docs",
                "context_documents",
                "evidence_docs",
                "top_docs",
                "candidates",
            ),
        )
        if docs is None:
            docs = _first_value(
                [retriever_debug],
                keys=("top_documents", "documents", "context_docs"),
            )
        if docs is not None:
            telemetry["docs"] = docs

    if "scores" not in telemetry:
        scores = _first_value(
            sources,
            keys=(
                "scores",
                "score_breakdown",
                "score_map",
                "metrics",
                "context_scores",
                "ranking_scores",
                "source_scores",
            ),
        )
        if scores is not None:
            telemetry["scores"] = scores
        else:
            score_map: dict[str, Any] = {}
            relevance = _first_value(
                sources,
                keys=("relevance", "context_relevance", "retrieval_score"),
            )
            threshold = _first_value(sources, keys=("low_context_threshold", "threshold"))
            if relevance is not None:
                score_map["relevance"] = relevance
            if threshold is not None:
                score_map["low_context_threshold"] = threshold
            if score_map:
                telemetry["scores"] = score_map

    if "source_reasoning" not in telemetry:
        source_reasoning = _first_value(
            sources,
            keys=(
                "source_reasoning",
                "source_reasonings",
                "reasoning_by_source",
                "per_source_reasoning",
                "source_notes",
            ),
        )
        if source_reasoning is None:
            source_reasoning = _first_value(
                [retriever_debug],
                keys=("score_trace", "final_score_trace"),
            )
        if source_reasoning is not None:
            telemetry["source_reasoning"] = source_reasoning

    if "errors" not in telemetry:
        errors = _first_value(
            sources,
            keys=(
                "errors",
                "error",
                "error_list",
                "source_errors",
                "retrieval_errors",
                "failed_sources",
            ),
        )
        if errors is None:
            errors = _first_value([retriever_debug], keys=("source_errors",))
        if errors is not None:
            telemetry["errors"] = errors
        elif isinstance(normalized.get("fallback_reason"), str):
            telemetry["errors"] = [normalized["fallback_reason"]]

    if "verification_matrix" not in telemetry:
        verification_matrix = _first_value(
            sources,
            keys=_VERIFICATION_MATRIX_KEYS,
        )
        if verification_matrix is not None:
            telemetry["verification_matrix"] = verification_matrix

    if "contradiction_summary" not in telemetry:
        contradiction_summary = _first_value(
            sources,
            keys=_CONTRADICTION_SUMMARY_KEYS,
        )
        if contradiction_summary is not None:
            telemetry["contradiction_summary"] = contradiction_summary

    if "trace_metadata" not in telemetry and trace_metadata is not None:
        telemetry["trace_metadata"] = trace_metadata

    return telemetry or None


def _apply_role_gated_telemetry(
    result: dict[str, Any] | None,
    *,
    role: str | None,
) -> dict[str, Any] | None:
    """Replace the result's telemetry with a role-gated, PII-safe view (R3).

    Gated behind ``RESEARCH_ROLE_GATED_TELEMETRY_ENABLED``; when the flag is off
    the result is returned unchanged so legacy behavior is preserved (R20.2).
    Fail-closed: an unknown/unavailable role yields no telemetry at all.
    """

    if not isinstance(result, dict):
        return result
    if not get_settings().research_role_gated_telemetry_enabled:
        return result

    raw_telemetry = result.get("telemetry")
    telemetry_source: dict[str, Any] = raw_telemetry if isinstance(raw_telemetry, dict) else {}
    # Fold in the progress/flow signals so the sanitized summary can be built even
    # when the raw telemetry block carries no stage list of its own.
    for stage_key in ("flow_stages", "flow_events", "active_stage"):
        if stage_key not in telemetry_source and stage_key in result:
            telemetry_source = {**telemetry_source, stage_key: result[stage_key]}

    gated = result.copy()
    gated["telemetry"] = sanitize_telemetry(telemetry_source, role=role)
    return gated


def _normalize_tier2_response(payload: dict[str, Any]) -> dict[str, Any]:
    normalized = dict(payload)

    metadata = normalized.get("metadata")
    metadata_obj = metadata if isinstance(metadata, dict) else None
    if metadata is not None and metadata_obj is None:
        normalized["metadata"] = {}

    if "context_debug" not in normalized and metadata_obj is not None:
        nested_context_debug = metadata_obj.get("context_debug")
        if isinstance(nested_context_debug, dict):
            normalized["context_debug"] = nested_context_debug

    context_debug_obj = normalized.get("context_debug")
    if not isinstance(context_debug_obj, dict):
        context_debug_obj = None

    if "flow_events" not in normalized and metadata_obj is not None:
        nested_flow_events = metadata_obj.get("flow_events")
        if isinstance(nested_flow_events, list):
            normalized["flow_events"] = nested_flow_events

    metadata_telemetry_obj = (
        metadata_obj.get("telemetry")
        if metadata_obj is not None and isinstance(metadata_obj.get("telemetry"), dict)
        else None
    )
    retrieval_trace_obj = (
        context_debug_obj.get("retrieval_trace")
        if context_debug_obj is not None
        and isinstance(context_debug_obj.get("retrieval_trace"), dict)
        else None
    )

    if "source_attempts" not in normalized:
        source_attempts = _first_value(
            [metadata_obj, metadata_telemetry_obj, context_debug_obj, retrieval_trace_obj],
            keys=(
                "source_attempts",
                "connector_attempts",
                "provider_events",
                "retrieval_attempts",
            ),
        )
        if source_attempts is not None:
            normalized["source_attempts"] = source_attempts

    if "source_errors" not in normalized:
        source_errors = _first_value(
            [metadata_obj, metadata_telemetry_obj, context_debug_obj, retrieval_trace_obj],
            keys=("source_errors", "retrieval_errors"),
        )
        if source_errors is not None:
            normalized["source_errors"] = source_errors

    if "fallback_reason" not in normalized:
        fallback_reason = _first_value(
            [metadata_obj, metadata_telemetry_obj, context_debug_obj, retrieval_trace_obj],
            keys=("fallback_reason",),
        )
        if isinstance(fallback_reason, str):
            stripped_reason = fallback_reason.strip()
            if stripped_reason:
                normalized["fallback_reason"] = stripped_reason

    if "verification_matrix" not in normalized:
        verification_matrix = _first_value(
            [metadata_obj, metadata_telemetry_obj, context_debug_obj, retrieval_trace_obj],
            keys=_VERIFICATION_MATRIX_KEYS,
        )
        if verification_matrix is not None:
            normalized["verification_matrix"] = verification_matrix

    if "contradiction_summary" not in normalized:
        contradiction_summary = _first_value(
            [metadata_obj, metadata_telemetry_obj, context_debug_obj, retrieval_trace_obj],
            keys=_CONTRADICTION_SUMMARY_KEYS,
        )
        if contradiction_summary is not None:
            normalized["contradiction_summary"] = contradiction_summary

    if "trace_metadata" not in normalized:
        trace_metadata = _extract_trace_metadata(
            [
                normalized,
                metadata_obj,
                metadata_telemetry_obj,
                context_debug_obj,
                retrieval_trace_obj,
            ]
        )
        if trace_metadata is not None:
            normalized["trace_metadata"] = trace_metadata

    if "query_plan" not in normalized:
        query_plan = _first_value(
            [metadata_obj, metadata_telemetry_obj, context_debug_obj, retrieval_trace_obj],
            keys=("query_plan", "search_plan", "search_trace"),
        )
        if query_plan is not None:
            normalized["query_plan"] = query_plan

    if "query_plan" not in normalized and normalized.get("search_plan") is not None:
        normalized["query_plan"] = normalized.get("search_plan")
    if "search_plan" not in normalized and normalized.get("query_plan") is not None:
        normalized["search_plan"] = normalized.get("query_plan")

    telemetry = _build_tier2_telemetry(
        normalized=normalized,
        metadata_obj=metadata_obj,
        context_debug_obj=context_debug_obj,
    )
    if telemetry is not None:
        normalized["telemetry"] = telemetry
        if "source_attempts" not in normalized and telemetry.get("source_attempts") is not None:
            normalized["source_attempts"] = telemetry.get("source_attempts")
        if "query_plan" not in normalized and telemetry.get("query_plan") is not None:
            normalized["query_plan"] = telemetry.get("query_plan")
        if "search_plan" not in normalized and telemetry.get("search_plan") is not None:
            normalized["search_plan"] = telemetry.get("search_plan")
        if (
            "verification_matrix" not in normalized
            and telemetry.get("verification_matrix") is not None
        ):
            normalized["verification_matrix"] = telemetry.get("verification_matrix")
        if (
            "contradiction_summary" not in normalized
            and telemetry.get("contradiction_summary") is not None
        ):
            normalized["contradiction_summary"] = telemetry.get("contradiction_summary")
        if "trace_metadata" not in normalized:
            telemetry_trace_metadata = _extract_trace_metadata([telemetry])
            if telemetry_trace_metadata is not None:
                normalized["trace_metadata"] = telemetry_trace_metadata
        if "source_errors" not in normalized:
            telemetry_source_errors = telemetry.get("source_errors")
            if telemetry_source_errors is None and isinstance(telemetry.get("errors"), dict):
                telemetry_source_errors = telemetry.get("errors")
            if telemetry_source_errors is not None:
                normalized["source_errors"] = telemetry_source_errors

    if "source_errors" in normalized:
        normalized["source_errors"] = normalize_source_errors(normalized.get("source_errors"))

    answer_markdown = normalized.get("answer_markdown")
    if not isinstance(answer_markdown, str) or not answer_markdown.strip():
        for key in ("answer", "summary", "message"):
            candidate = normalized.get(key)
            if isinstance(candidate, str) and candidate.strip():
                normalized["answer_markdown"] = candidate
                break

    if not isinstance(normalized.get("answer_format"), str):
        normalized["answer_format"] = "markdown"
    if not isinstance(normalized.get("render_hints"), dict):
        normalized["render_hints"] = dict(_DEFAULT_MARKDOWN_RENDER_HINTS)

    return _canonicalize_research_payload_contract(normalized)


def _extract_research_source_used(normalized: dict[str, Any]) -> list[str]:
    metadata_obj = (
        normalized.get("metadata") if isinstance(normalized.get("metadata"), dict) else {}
    )
    telemetry_obj = (
        normalized.get("telemetry") if isinstance(normalized.get("telemetry"), dict) else {}
    )

    source_used = normalize_source_used(
        normalized.get("source_used") or metadata_obj.get("source_used") or []
    )
    source_attempts = normalized.get("source_attempts")
    if not isinstance(source_attempts, list):
        source_attempts = telemetry_obj.get("source_attempts")
    if not isinstance(source_attempts, list):
        source_attempts = metadata_obj.get("source_attempts")
    if isinstance(source_attempts, list):
        for attempt in source_attempts:
            if not isinstance(attempt, dict):
                continue
            for key in ("source", "provider", "connector", "name"):
                raw_value = attempt.get(key)
                if raw_value is None:
                    continue
                normalized_value = str(raw_value).strip().lower()
                if normalized_value and normalized_value not in source_used:
                    source_used.append(normalized_value)
                break

    citations = normalized.get("citations")
    if isinstance(citations, list):
        for citation in citations:
            if isinstance(citation, str):
                normalized_value = citation.strip().lower()
            elif isinstance(citation, dict):
                normalized_value = (
                    str(citation.get("source") or citation.get("id") or citation.get("title") or "")
                    .strip()
                    .lower()
                )
            else:
                normalized_value = ""
            if normalized_value and normalized_value not in source_used:
                source_used.append(normalized_value)

    return source_used


def _attach_research_attribution(normalized: dict[str, Any]) -> dict[str, Any]:
    normalized = _canonicalize_research_payload_contract(normalized)
    rich_citations = normalized.get("citations")
    metadata_obj = (
        normalized.get("metadata") if isinstance(normalized.get("metadata"), dict) else {}
    )
    telemetry_obj = (
        normalized.get("telemetry") if isinstance(normalized.get("telemetry"), dict) else {}
    )
    source_used = _extract_research_source_used(normalized)
    source_errors = normalize_source_errors(
        normalized.get("source_errors")
        or metadata_obj.get("source_errors")
        or telemetry_obj.get("source_errors")
        or (telemetry_obj.get("errors") if isinstance(telemetry_obj.get("errors"), dict) else {})
        or {}
    )
    if "source_errors" not in normalized and source_errors:
        normalized["source_errors"] = source_errors
    mode = _normalize_research_mode_value(
        normalized.get("research_mode") or metadata_obj.get("research_mode"),
        default="fast",
    )
    normalized["research_mode"] = mode
    if isinstance(metadata_obj, dict):
        metadata_obj["research_mode"] = mode

    sources = [
        {
            "id": source_id,
            "name": source_id.replace("_", " ").title(),
            "type": "retrieval",
            "category": "research",
        }
        for source_id in source_used
    ]
    fallback_reason = normalized.get("fallback_reason")
    if not isinstance(fallback_reason, str) or not fallback_reason.strip():
        metadata_fallback_reason = metadata_obj.get("fallback_reason")
        if isinstance(metadata_fallback_reason, str) and metadata_fallback_reason.strip():
            fallback_reason = metadata_fallback_reason.strip()
            normalized["fallback_reason"] = fallback_reason
        else:
            fallback_reason = ""
    fallback_used = bool(
        normalized.get("fallback_used")
        or metadata_obj.get("fallback_used")
        or normalized.get("fallback")
        or fallback_reason
    )
    if "fallback" in normalized or fallback_used:
        normalized["fallback"] = fallback_used
    normalized["fallback_used"] = fallback_used
    metadata_obj["fallback_used"] = fallback_used
    if fallback_reason:
        metadata_obj["fallback_reason"] = fallback_reason

    attribution = build_attribution(
        channel="research",
        mode=mode,
        sources=sources,
        citations_payload=normalized.get("citations"),
        source_used=source_used,
        source_errors=source_errors,
        fallback_used=fallback_used,
    )
    response = attach_attribution(normalized, attribution=attribution)
    if isinstance(rich_citations, list):
        response["citations"] = rich_citations
    return response


def _coerce_personal_mode(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, int | float):
        return value != 0
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "on"}
    return False


def _enforce_never_fast_and_personal(payload: dict[str, Any]) -> None:
    """Reject any request that sets ``personal_mode`` while the mode is fast.

    clara-research R15.2 / Property 30: preserve the invariant "never (fast && personal)".
    Personalization (PHR + medicine cabinet) is valid only in tier2 deep/deep_beta runs, so a
    fast-mode request carrying ``personal_mode`` is rejected here rather than silently coerced.
    Applied unconditionally (independent of personalization feature flags) so the invariant holds
    even when personalization is disabled.
    """
    if _coerce_personal_mode(payload.get("personal_mode")) and (
        _coerce_research_mode(payload) == "fast"
    ):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "personal_mode is not allowed when research_mode is 'fast' "
                "(invariant: never (fast && personal))."
            ),
        )


def _as_dict_list(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    return [item for item in value if isinstance(item, dict)]


def _compact_text(value: Any, *, limit: int = 160) -> str:
    text = str(value or "").strip()
    if len(text) <= limit:
        return text
    return f"{text[:limit].rstrip()}..."


def _age_band(date_of_birth: Any) -> str:
    """Return a coarse age band for personalised research prompts.

    Deep/DeepBeta only need clinically relevant context.  Names and exact birth
    dates are neither necessary nor appropriate to send to an LLM provider, so
    retain the useful life-stage signal without forwarding a direct identifier.
    """

    if date_of_birth is None:
        return ""
    try:
        today = datetime.now(tz=UTC).date()
        age = today.year - date_of_birth.year - (
            (today.month, today.day) < (date_of_birth.month, date_of_birth.day)
        )
    except (AttributeError, TypeError, ValueError):
        return ""
    if age < 0:
        return ""
    if age < 18:
        return "under_18"
    if age < 40:
        return "18_39"
    if age < 65:
        return "40_64"
    return "65_plus"


def _build_personal_context_payload(
    db: Session,
    *,
    user_id: int,
    answer_language: str,
) -> dict[str, Any]:
    profile = db.execute(
        select(PhrProfile).where(PhrProfile.user_id == user_id)
    ).scalar_one_or_none()
    allergies: list[dict[str, Any]] = []
    conditions: list[dict[str, Any]] = []
    profile_medications: list[dict[str, Any]] = []
    profile_payload: dict[str, Any] = {}

    if profile is not None:
        allergies = _as_dict_list(profile.allergies_json)
        conditions = _as_dict_list(profile.conditions_json)
        profile_medications = _as_dict_list(profile.medications_json)
        profile_payload = {
            # Data minimisation: personal-mode synthesis receives only the
            # life-stage and clinical fields it can legitimately use. Never
            # relay direct identifiers, free-text notes, or an exact DOB.
            "age_band": _age_band(profile.date_of_birth),
            "gender": profile.gender or "",
        }

    cabinet = db.execute(
        select(MedicineCabinet).where(MedicineCabinet.user_id == user_id)
    ).scalar_one_or_none()
    cabinet_items: list[dict[str, Any]] = []
    if cabinet is not None:
        rows = (
            db.execute(
                select(MedicineItem)
                .where(MedicineItem.cabinet_id == cabinet.id)
                .order_by(MedicineItem.updated_at.desc(), MedicineItem.id.desc())
                .limit(120)
            )
            .scalars()
            .all()
        )
        cabinet_items = [
            {
                "name": item.drug_name or item.normalized_name or "",
                "normalized_name": item.normalized_name or "",
                "dose": item.dosage or "",
                "dosage_form": item.dosage_form or "",
                "quantity": item.quantity,
                "source": item.source or "",
                "expires_on": item.expires_on.isoformat() if item.expires_on else None,
                "note": item.note or "",
            }
            for item in rows
        ]

    merged_medications: list[dict[str, Any]] = []
    seen_medication_keys: set[str] = set()
    for item in profile_medications:
        name = _compact_text(item.get("name"), limit=120)
        if not name:
            continue
        key = name.lower()
        if key in seen_medication_keys:
            continue
        seen_medication_keys.add(key)
        merged_medications.append(
            {
                "name": name,
                "dose": _compact_text(item.get("dose"), limit=80),
                "frequency": _compact_text(item.get("frequency"), limit=80),
                "note": _compact_text(item.get("note"), limit=140),
                "source": "phr",
            }
        )
    for item in cabinet_items:
        name = _compact_text(item.get("name"), limit=120)
        if not name:
            continue
        key = name.lower()
        if key in seen_medication_keys:
            continue
        seen_medication_keys.add(key)
        merged_medications.append(
            {
                "name": name,
                "dose": _compact_text(item.get("dose"), limit=80),
                "frequency": "",
                "note": _compact_text(item.get("note"), limit=140),
                "source": "cabinet",
            }
        )

    if answer_language == "en":
        header = "Personal health context (PHR + medicine cabinet):"
        allergies_label = "Allergies"
        conditions_label = "Conditions"
        meds_label = "Current meds"
    else:
        header = "Bối cảnh sức khỏe cá nhân (PHR + tủ thuốc):"
        allergies_label = "Dị ứng"
        conditions_label = "Bệnh nền"
        meds_label = "Thuốc hiện có"

    allergy_names = [str(item.get("name") or "").strip() for item in allergies][:12]
    condition_names = [str(item.get("name") or "").strip() for item in conditions][:12]
    med_names = [str(item.get("name") or "").strip() for item in merged_medications][:16]
    summary_lines = [f"- {header}"]
    if allergy_names:
        allergy_str = ", ".join(name for name in allergy_names if name)
        summary_lines.append(f"- {allergies_label}: {allergy_str}.")
    if condition_names:
        condition_str = ", ".join(name for name in condition_names if name)
        summary_lines.append(f"- {conditions_label}: {condition_str}.")
    if med_names:
        summary_lines.append(f"- {meds_label}: {', '.join([name for name in med_names if name])}.")
    summary_markdown = "\n".join(summary_lines)

    return {
        "profile": profile_payload,
        "allergies": allergies,
        "conditions": conditions,
        "medications": merged_medications,
        "medicine_cabinet": {
            "exists": cabinet is not None,
            "label": cabinet.label if cabinet is not None else "",
            "items": cabinet_items,
        },
        "summary_markdown": summary_markdown,
    }


# Gap-fill pass-count keys an upstream caller may supply to request additional bounded
# gap-fill retrieval passes. The Research_API independently clamps any such request to
# ``RESEARCH_API_GAP_FILL_HARD_MAX`` before forwarding to ML (clara-research R5.5,
# defense in depth). Runtime LLM overrides are deliberately not supported; only
# bounded top-level request keys are inspected.
_GAP_FILL_PASS_REQUEST_KEYS: tuple[str, ...] = (
    "gap_fill_max_passes",
    "research_gap_fill_max_passes",
)


def _clamp_gap_fill_request_keys(container: dict[str, Any], *, ceiling: int) -> None:
    """Clamp in place any gap-fill pass-count keys present in ``container`` to ``ceiling``."""

    for key in _GAP_FILL_PASS_REQUEST_KEYS:
        if key not in container:
            continue
        try:
            requested = int(container[key])
        except (TypeError, ValueError):
            # Non-numeric request values are left untouched for the schema layer to reject.
            continue
        container[key] = max(0, min(requested, ceiling))


def _enforce_api_gap_fill_ceiling(
    payload: dict[str, Any], *, hard_max: int
) -> dict[str, Any]:
    """Clamp any requested gap-fill pass count to the API hard ceiling (clara-research R5.5).

    The Research_API enforces ``RESEARCH_API_GAP_FILL_HARD_MAX`` externally so a
    misbehaving or over-eager caller cannot drive unbounded gap-fill retrieval. When no
    gap-fill pass count is requested, ``payload`` is returned unchanged so legacy
    (flag-off) requests keep their exact shape (clara-research R20.2).
    """

    ceiling = max(0, int(hard_max))
    _clamp_gap_fill_request_keys(payload, ceiling=ceiling)
    return payload


def _build_tier2_upstream_payload(
    payload: dict[str, Any],
    *,
    db: Session,
    user: User,
    token: TokenPayload,
) -> dict[str, Any]:
    settings = get_settings()
    upstream_payload = dict(payload)
    # Provider, endpoint, model and credentials are deployment-only settings.
    # Never relay a request-supplied runtime block to ML: even an admin-facing
    # web control can be compromised, and this boundary must not become an
    # SSRF, credential-forwarding, or model-governance bypass.  Registered
    # DeepSeek V4 task contracts resolve the actual model server-side.
    upstream_payload.pop("llm_runtime", None)
    requested_language = (
        str(upstream_payload.get("ui_language") or upstream_payload.get("answer_language") or "vi")
        .strip()
        .lower()
    )
    answer_language = "en" if requested_language == "en" else "vi"
    upstream_payload["ui_language"] = answer_language
    upstream_payload["answer_language"] = answer_language
    explicit_research_mode = "research_mode" in upstream_payload or "mode" in upstream_payload
    research_mode, retrieval_stack_mode = _resolve_tier2_execution_modes(upstream_payload)
    if explicit_research_mode:
        upstream_payload["research_mode"] = research_mode
    upstream_payload["retrieval_stack_mode"] = retrieval_stack_mode
    upstream_payload.pop("stack_mode", None)
    answer_language = _normalize_answer_language_value(
        upstream_payload.get("ui_language") or upstream_payload.get("answer_language"),
        default="vi",
    )
    personal_mode = _coerce_personal_mode(upstream_payload.get("personal_mode"))
    upstream_payload["personal_mode"] = personal_mode
    upstream_payload["ui_language"] = answer_language
    upstream_payload["answer_language"] = answer_language
    upstream_payload["answer_format"] = str(upstream_payload.get("answer_format") or "markdown")
    upstream_payload["response_format"] = str(upstream_payload.get("response_format") or "markdown")
    # Output modes are a separately dark presentation feature. The selected
    # closed value is persisted with the durable job only when the API gate is
    # enabled, and is echoed through ML solely for a second gate. It is never
    # passed to the LLM prompt or allowed to alter retrieval/model policy.
    if settings.research_output_modes_enabled:
        upstream_payload["output_mode"] = _resolve_research_output_mode(
            upstream_payload.get("output_mode"), role=token.role
        )
    else:
        upstream_payload.pop("output_mode", None)
    incoming_render_hints = upstream_payload.get("render_hints")
    if isinstance(incoming_render_hints, dict):
        merged_render_hints = {
            **_DEFAULT_MARKDOWN_RENDER_HINTS,
            **incoming_render_hints,
        }
    else:
        merged_render_hints = dict(_DEFAULT_MARKDOWN_RENDER_HINTS)
    upstream_payload["render_hints"] = merged_render_hints
    transient_documents = _build_uploaded_documents(
        payload.get("uploaded_file_ids"),
        owner_user_id=user.id,
        db=db,
    )
    source_ids = _extract_source_ids(payload)
    source_documents = _build_source_documents(db, owner_user_id=user.id, source_ids=source_ids)
    source_hub_filters = _extract_source_hub_sources(payload)
    source_hub_documents = _build_source_hub_documents(
        db,
        owner_user_id=user.id,
        query=str(payload.get("query") or payload.get("message") or ""),
        source_filters=source_hub_filters,
    )
    uploaded_documents = [*transient_documents, *source_documents, *source_hub_documents]

    if uploaded_documents or payload.get("source_mode") in {"uploaded_files", "knowledge_sources"}:
        upstream_payload["uploaded_documents"] = uploaded_documents

    # clara-research R15.1 / Property 29: personalization (PHR + medicine cabinet) is
    # incorporated into the synthesis only when personal_mode is set, the research mode is
    # deep or deep_beta, AND the user has granted consent. Fast mode never personalizes
    # (R15.2 invariant — rejected upstream and re-checked here as defense in depth so personal
    # context can never leak into a fast run). No consent ⇒ run without personalization
    # (R15.3), which is not an error.
    if personal_mode and research_mode in {"deep", "deep_beta"}:
        # PHR consent gate (Req 2.2, 2.3, 2.4): when consent enforcement is on, only feed
        # personal PHR context if personalization + research consent are present; skip the
        # personal-mode context if either consent is absent. The clara-research personalization
        # surface (RESEARCH_PERSONALIZATION_ENABLED) likewise requires explicit personalization
        # + research consent regardless of the legacy PHR enforcement flag (R15.1/R15.3). The
        # gate reads current consent each request, so revocation takes effect on the next
        # request. Both flags off ⇒ legacy behavior (always include).
        phr_flags = phr_features(settings)
        include_personal = True
        if phr_flags.consent_enforcement or settings.research_personalization_enabled:
            include_personal = PhrConsentService.is_granted(
                db, user_id=user.id, purpose="personalization"
            ) and PhrConsentService.is_granted(db, user_id=user.id, purpose="research")
        # Compliance granular-consent gate (Req 2.1, 2.3): when
        # COMPLIANCE_GRANULAR_CONSENT_ENABLED is on, also require the
        # compliance-ledger personalization + research grants. Flag off ⇒
        # has_consent returns True, so legacy behavior is preserved exactly.
        compliance = ComplianceService(db, settings=settings)
        include_personal = (
            include_personal
            and compliance.has_consent(user_id=user.id, purpose=PURPOSE_PERSONALIZATION)
            and compliance.has_consent(user_id=user.id, purpose=PURPOSE_RESEARCH)
        )
        if include_personal:
            personal_context = _build_personal_context_payload(
                db,
                user_id=user.id,
                answer_language=answer_language,
            )
            upstream_payload["personal_context"] = personal_context
            metadata = upstream_payload.get("metadata")
            if not isinstance(metadata, dict):
                metadata = {}
            metadata["personal_mode"] = True
            metadata["personal_context_available"] = bool(personal_context.get("summary_markdown"))
            metadata["personal_context_medication_count"] = len(
                personal_context.get("medications", [])
            )
            upstream_payload["metadata"] = metadata

    upstream_payload["role"] = token.role
    upstream_payload["strict_deepseek_required"] = bool(settings.deepseek_strict_mode)
    runtime_rag_flow, runtime_rag_sources = _load_research_rag_runtime(db)

    try:
        upstream_payload["rag_flow"] = RagFlowConfig.model_validate(runtime_rag_flow).model_dump()
    except Exception:
        upstream_payload["rag_flow"] = RagFlowConfig().model_dump()
    upstream_payload["rag_sources"] = runtime_rag_sources

    # R5.5: clamp any requested gap-fill pass count to the API hard ceiling before the
    # payload is forwarded to ML. No-op (legacy shape) when no count is requested.
    upstream_payload = _enforce_api_gap_fill_ceiling(
        upstream_payload, hard_max=settings.research_api_gap_fill_hard_max
    )

    return upstream_payload


def _enforce_request_execution_contract(
    normalized: dict[str, Any],
    *,
    request_payload: dict[str, Any],
) -> dict[str, Any]:
    response = dict(normalized)
    metadata_obj = response.get("metadata")
    if not isinstance(metadata_obj, dict):
        metadata_obj = {}
        response["metadata"] = metadata_obj

    research_mode, retrieval_stack_mode = _resolve_tier2_execution_modes(request_payload)
    answer_language = _normalize_answer_language_value(
        request_payload.get("ui_language") or request_payload.get("answer_language"),
        default="vi",
    )
    response["research_mode"] = research_mode
    metadata_obj["research_mode"] = research_mode
    response["retrieval_stack_mode"] = retrieval_stack_mode
    metadata_obj["retrieval_stack_mode"] = retrieval_stack_mode
    response["ui_language"] = answer_language
    metadata_obj["answer_language"] = answer_language

    # API and ML must both acknowledge a closed output mode before any
    # presentation payload can be composed. A missing/mismatched ML
    # acknowledgement fails closed to the established answer-only response.
    settings = get_settings()
    if settings.research_output_modes_enabled:
        expected_output_mode = _resolve_research_output_mode(
            request_payload.get("output_mode"),
            role=str(request_payload.get("role") or ""),
        )
        candidate_output_mode = str(
            response.get("output_mode") or metadata_obj.get("output_mode") or ""
        ).strip().lower()
        if candidate_output_mode == expected_output_mode:
            response["output_mode"] = expected_output_mode
            metadata_obj["output_mode"] = expected_output_mode
        else:
            response.pop("output_mode", None)
            metadata_obj.pop("output_mode", None)

    # R5.5 (defense in depth): the API forcibly caps the reported gap-fill pass count at
    # the configured hard ceiling so a misbehaving orchestrator cannot surface or persist
    # an unbounded count. Only acts when the field is present (legacy shape preserved).
    hard_max = max(0, int(get_settings().research_api_gap_fill_hard_max))
    if "gap_fill_passes" in response:
        try:
            reported_passes = int(response["gap_fill_passes"])
        except (TypeError, ValueError):
            reported_passes = None
        if reported_passes is not None:
            response["gap_fill_passes"] = max(0, min(reported_passes, hard_max))

    return response


def _empty_job_progress() -> dict[str, Any]:
    return {
        "flow_events": [],
        "flow_stages": [],
        "active_stage": "",
        "active_status": "",
        "status_note": "",
        "reasoning_steps": [],
    }


def _safe_failure_stage(progress: Any) -> str:
    """Expose a bounded pipeline stage in failure telemetry, never upstream text."""

    if not isinstance(progress, dict):
        return "unknown"
    value = str(progress.get("active_stage") or "").strip().lower()
    if not value or not re.fullmatch(r"[a-z0-9_]{1,64}", value):
        return "unknown"
    return value


def _stage_from_events(events: list[dict[str, Any]]) -> list[dict[str, Any]]:
    stage_map: dict[str, dict[str, Any]] = {}
    for event in events:
        stage = str(event.get("stage") or "").strip()
        if not stage:
            continue
        item = stage_map.get(stage) or {
            "id": stage,
            "label": stage.replace("_", " ").title(),
            "status": "pending",
            "detail": "",
            "source": "flow_events",
        }
        event_status = str(event.get("status") or "").strip().lower() or "pending"
        item["status"] = event_status
        if isinstance(event.get("note"), str) and event.get("note"):
            item["detail"] = event["note"]
        stage_map[stage] = item
    return list(stage_map.values())


def _extract_stage_status_note(value: Any) -> tuple[str, str, str] | None:
    if not isinstance(value, dict):
        return None
    stage = str(value.get("stage") or value.get("phase") or value.get("active_stage") or "").strip()
    if not stage:
        return None
    status_text = str(
        value.get("status") or value.get("state") or value.get("active_status") or ""
    ).strip()
    note_raw = (
        value.get("note")
        or value.get("detail")
        or value.get("message")
        or value.get("status_note")
        or ""
    )
    note = str(note_raw).strip() if note_raw is not None else ""
    return stage, status_text.lower() or "in_progress", note


def _latest_stage_status_note(events: Any) -> tuple[str, str, str] | None:
    if not isinstance(events, list):
        return None
    for raw_event in reversed(events):
        signal = _extract_stage_status_note(raw_event)
        if signal is not None:
            return signal
    return None


def _extract_ml_progress_signal(payload: dict[str, Any]) -> tuple[str, str, str] | None:
    signal = _latest_stage_status_note(payload.get("flow_events"))
    if signal is not None:
        return signal

    metadata_obj = payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None
    if metadata_obj is not None:
        signal = _latest_stage_status_note(metadata_obj.get("flow_events"))
        if signal is not None:
            return signal

    telemetry_obj = (
        metadata_obj.get("telemetry")
        if metadata_obj is not None and isinstance(metadata_obj.get("telemetry"), dict)
        else None
    )
    context_debug_obj = (
        payload.get("context_debug") if isinstance(payload.get("context_debug"), dict) else None
    )

    for candidate in (payload, metadata_obj, telemetry_obj, context_debug_obj):
        signal = _extract_stage_status_note(candidate)
        if signal is not None:
            return signal
    return None


def _merge_progress_with_ml_flow_events(
    *,
    progress: dict[str, Any],
    ml_flow_events: list[dict[str, Any]],
) -> dict[str, Any]:
    history_events = progress.get("flow_events")
    if not isinstance(history_events, list):
        history_events = []
    else:
        history_events = [item for item in history_events if isinstance(item, dict)]

    clean_ml_events = [item for item in ml_flow_events if isinstance(item, dict)]
    merged_events = [*history_events, *clean_ml_events][-120:]
    progress["flow_events"] = merged_events
    progress["flow_stages"] = _stage_from_events(merged_events)

    signal = _latest_stage_status_note(clean_ml_events)
    if signal is None:
        signal = _latest_stage_status_note(merged_events)
    if signal is not None:
        stage, status_text, note = signal
        progress["active_stage"] = stage
        progress["active_status"] = status_text
        if note:
            progress["status_note"] = note
    return progress


def _append_job_event(
    db: Session,
    *,
    job: ResearchJob,
    stage: str,
    status_text: str,
    note: str,
    payload: dict[str, Any] | None = None,
) -> None:
    if isinstance(job.progress_json, dict):
        progress = dict(job.progress_json)
    else:
        progress = _empty_job_progress()
    events = progress.get("flow_events")
    if not isinstance(events, list):
        events = []
    else:
        events = [item for item in events if isinstance(item, dict)]
    event_item: dict[str, Any] = {
        "id": str(uuid4()),
        "stage": stage,
        "status": status_text,
        "note": note,
        "component": "research_job",
        "timestamp": datetime.now(tz=UTC).isoformat(),
    }
    if payload:
        event_item["payload"] = payload
    events.append(event_item)
    progress["flow_events"] = list(events[-80:])
    progress["flow_stages"] = _stage_from_events(progress["flow_events"])
    progress["active_stage"] = stage
    progress["active_status"] = status_text
    progress["status_note"] = note
    reasoning_steps = progress.get("reasoning_steps")
    if not isinstance(reasoning_steps, list):
        reasoning_steps = []
    else:
        reasoning_steps = [item for item in reasoning_steps if isinstance(item, dict)]
    reasoning_steps.append(
        {
            "stage": stage,
            "status": status_text,
            "note": note,
            "timestamp": event_item["timestamp"],
        }
    )
    progress["reasoning_steps"] = list(reasoning_steps[-40:])
    job.progress_json = json.loads(json.dumps(progress, ensure_ascii=False))
    job.updated_at = datetime.now(tz=UTC)
    if job.status == "running":
        job.lease_heartbeat_at = job.updated_at
    db.add(job)
    db.commit()

    try:
        # The global observability stream is explicitly no-PII.  Job query,
        # notes, raw upstream errors, retrieval traces and verification rows
        # can contain patient/research text, so they stay in the owner-scoped
        # result only and are never mirrored here.
        store_event: dict[str, Any] = {
            "job_id": str(job.job_id),
            "stage": stage,
            "status": status_text,
            "timestamp": event_item["timestamp"],
        }
        if isinstance(payload, dict):
            store_event["fallback_used"] = bool(payload.get("fallback_used"))
            unsupported_claims = payload.get("unsupported_claims")
            if isinstance(unsupported_claims, list):
                store_event["unsupported_claim_count"] = len(unsupported_claims)
            verification_matrix = payload.get("verification_matrix")
            if isinstance(verification_matrix, list):
                store_event["verification_claim_count"] = len(verification_matrix)
        get_flow_event_store().append(
            source="research",
            user_id=str(job.user_id),
            role=str(job.role or "normal"),
            intent="research_tier2",
            model_used=(
                str(payload.get("model_used"))
                if isinstance(payload, dict) and payload.get("model_used") is not None
                else None
            ),
            event=store_event,
            occurred_at=event_item["timestamp"],
        )
    except Exception:
        # Best effort only; never break research job flow due to telemetry stream persistence.
        pass


def _serialize_research_job(
    job: ResearchJob,
    *,
    role: str | None = None,
) -> ResearchTier2JobResponse:
    progress = job.progress_json if isinstance(job.progress_json, dict) else _empty_job_progress()
    result = job.result_json if isinstance(job.result_json, dict) else None
    result = _apply_role_gated_telemetry(result, role=role)
    return ResearchTier2JobResponse(
        job_id=job.job_id,
        status=str(job.status or "queued"),  # type: ignore[arg-type]
        query=job.query_text,
        started_at=job.started_at,
        completed_at=job.completed_at,
        created_at=job.created_at,
        updated_at=job.updated_at,
        progress=progress,
        result=result,
        error=job.error_text or None,
        run_manifest=(
            job.run_manifest_json if isinstance(job.run_manifest_json, dict) else None
        ),
        evidence_snapshot=(
            job.evidence_snapshot_json
            if isinstance(job.evidence_snapshot_json, dict)
            else None
        ),
        attempt_count=int(job.attempt_count or 0),
        recovery_count=int(job.recovery_count or 0),
    )


def _canonical_sha256(value: Any) -> str:
    serialized = json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
    return hashlib.sha256(serialized.encode("utf-8")).hexdigest()


def _sanitize_provider_secrets(value: Any) -> Any:
    """Redact provider credentials before a request is persisted or hashed."""

    if isinstance(value, dict):
        sanitized: dict[str, Any] = {}
        for raw_key, raw_value in value.items():
            key = str(raw_key)
            normalized_key = key.lower().replace("-", "_")
            if normalized_key in _PROVIDER_SECRET_KEYS or any(
                marker in normalized_key
                for marker in ("api_key", "secret", "password", "access_token")
            ):
                sanitized[key] = "[REDACTED]"
            else:
                sanitized[key] = _sanitize_provider_secrets(raw_value)
        return sanitized
    if isinstance(value, list):
        return [_sanitize_provider_secrets(item) for item in value]
    if isinstance(value, tuple):
        return [_sanitize_provider_secrets(item) for item in value]
    return value


def _apply_research_quality_gates(
    result: dict[str, Any], *, request_payload: dict[str, Any]
) -> dict[str, Any]:
    """Attach deterministic quality metadata and fail-closed degraded labels.

    This is intentionally an internal harness gate: it never fabricates a
    citation and never blocks a valid answer solely because a provider omitted
    optional evidence. Instead, an answer that cannot demonstrate evidence
    coverage is explicitly marked degraded for the UI and downstream policy.
    """

    gated = dict(result)
    citations = gated.get("citations")
    if not isinstance(citations, list):
        citations = gated.get("sources") if isinstance(gated.get("sources"), list) else []
    def is_resolvable_citation(item: Any) -> bool:
        if not isinstance(item, dict):
            return False
        source = str(item.get("source") or "").strip().lower()
        if source in {"", "system_fallback", "fallback", "unknown"}:
            return False
        # A label alone is not evidence. Require a stable retriever/study ID or
        # a resolvable URL so fabricated source names cannot pass the gate.
        return any(
            str(item.get(key) or "").strip()
            for key in ("source_id", "study_id", "pmid", "doi", "url")
        )

    def safe_nonnegative_int(value: Any) -> int:
        if isinstance(value, bool):
            return 0
        try:
            numeric = float(value)
        except (TypeError, ValueError):
            return 0
        if not math.isfinite(numeric) or numeric < 0:
            return 0
        return int(numeric)

    def is_nonnegative_integral_number(value: Any) -> bool:
        if isinstance(value, bool):
            return False
        try:
            numeric = float(value)
        except (TypeError, ValueError):
            return False
        return math.isfinite(numeric) and numeric >= 0 and numeric.is_integer()

    def is_unit_interval_number(value: Any) -> bool:
        if isinstance(value, bool):
            return False
        try:
            numeric = float(value)
        except (TypeError, ValueError):
            return False
        return math.isfinite(numeric) and 0.0 <= numeric <= 1.0

    real_citations = [item for item in citations if is_resolvable_citation(item)]
    unsupported = gated.get("unsupported_claims")
    unsupported_count = len(unsupported) if isinstance(unsupported, list) else 0
    citation_count = len(real_citations)
    answer_present = bool(
        str(gated.get("answer_markdown") or gated.get("answer") or "").strip()
    )
    mode = _coerce_research_mode(request_payload)
    raw_verification_matrix = gated.get("verification_matrix")
    verification_matrix = (
        raw_verification_matrix if isinstance(raw_verification_matrix, dict) else {}
    )
    verification_summary = (
        verification_matrix.get("summary")
        if isinstance(verification_matrix.get("summary"), dict)
        else {}
    )
    verification_rows = (
        verification_matrix.get("rows")
        if isinstance(verification_matrix.get("rows"), list)
        else []
    )
    verification_state = str(verification_matrix.get("state") or "").strip().lower()
    verification_version = str(verification_matrix.get("version") or "").strip()
    verifier_counts_well_formed = (
        is_nonnegative_integral_number(verification_summary.get("total_claims"))
        and is_nonnegative_integral_number(verification_summary.get("supported_claims"))
    )
    verifier_total_claims = safe_nonnegative_int(verification_summary.get("total_claims"))
    verifier_supported_claims = safe_nonnegative_int(
        verification_summary.get("supported_claims")
    )
    verifier_support_ratio = verification_summary.get("support_ratio")
    verifier_ratio_matches_counts = False
    if is_unit_interval_number(verifier_support_ratio) and verifier_counts_well_formed:
        if verifier_total_claims == 0:
            verifier_ratio_matches_counts = float(verifier_support_ratio) == 0.0
        elif verifier_supported_claims <= verifier_total_claims:
            verifier_ratio_matches_counts = (
                abs(
                    float(verifier_support_ratio)
                    - (verifier_supported_claims / verifier_total_claims)
                )
                <= 0.0001
            )
    verifier_summary_complete = (
        verifier_counts_well_formed
        and verifier_supported_claims <= verifier_total_claims
        and verifier_ratio_matches_counts
    )
    verifier_contract_reason: str | None = None
    if answer_present:
        if not isinstance(raw_verification_matrix, dict):
            verifier_contract_reason = "verification_unavailable"
        elif verification_state in {"skipped", "disabled"}:
            verifier_contract_reason = "verification_skipped"
        elif verification_state in {"unavailable", "error"}:
            verifier_contract_reason = "verification_unavailable"
        elif (
            verification_state not in {"verified", "warning"}
            or not verification_version
            or not verifier_summary_complete
            or not isinstance(verification_matrix.get("rows"), list)
        ):
            verifier_contract_reason = "verification_invalid"
    unsupported_from_verifier = sum(
        1
        for row in verification_rows
        if isinstance(row, dict)
        and str(
            row.get("support_status") or row.get("status") or row.get("verdict") or ""
        )
        .strip()
        .lower()
        in {"unsupported", "insufficient", "contradicted", "failed", "error"}
    )
    unsupported_count = max(unsupported_count, unsupported_from_verifier)
    metadata_input = gated.get("metadata") if isinstance(gated.get("metadata"), dict) else {}
    source_target = (
        gated.get("source_target_achieved")
        if isinstance(gated.get("source_target_achieved"), dict)
        else metadata_input.get("source_target_achieved")
        if isinstance(metadata_input.get("source_target_achieved"), dict)
        else {}
    )
    achieved_documents = safe_nonnegative_int(source_target.get("achieved_document_count"))
    retrieved_ids = gated.get("retrieved_ids")
    has_retrieved_evidence = bool(
        achieved_documents > 0
        or (isinstance(retrieved_ids, list) and any(str(item).strip() for item in retrieved_ids))
        or real_citations
    )
    support_ratio_raw = verification_summary.get("support_ratio")
    support_ratio = (
        float(support_ratio_raw)
        if isinstance(support_ratio_raw, int | float)
        else None
    )
    total_claims = safe_nonnegative_int(verification_summary.get("total_claims"))
    gate_reasons: list[str] = []
    if answer_present and citation_count == 0:
        gate_reasons.append("no_citations")
    if answer_present and not has_retrieved_evidence:
        gate_reasons.append("no_retrieved_evidence")
    if verifier_contract_reason:
        gate_reasons.append(verifier_contract_reason)
    if unsupported_count:
        gate_reasons.append("unsupported_claims")
    if (
        answer_present
        and support_ratio is not None
        and support_ratio <= 0.0
        and total_claims > 0
    ):
        gate_reasons.append("zero_claim_support")
    quality_gate = {
        "schema_version": "1.0",
        "passed": not gate_reasons,
        "citation_count": citation_count,
        "unsupported_claim_count": unsupported_count,
        "verifier": {
            "state": verification_state or "unavailable",
            "version": verification_version or None,
            "row_count": len(verification_rows),
            "total_claims": total_claims,
        },
        "answer_present": answer_present,
        "reasons": gate_reasons,
        "mode": mode,
    }
    metadata = dict(metadata_input)
    metadata["quality_gate"] = quality_gate
    gated["metadata"] = metadata
    gated["quality_gate"] = quality_gate
    if gate_reasons:
        gated["degraded"] = True
        gated["degraded_reason"] = ";".join(gate_reasons)
        # This is a policy abstention, not a model or research fallback.  Do
        # not let UI telemetry imply that unverified clinical prose was
        # replaced by another generated answer.
        gated["fallback_used"] = False
        metadata["fallback_used"] = False
        metadata["degraded_path"] = True
        gated["metadata"] = metadata
        # A failed medical-evidence gate is a release block, not merely a badge.
        # Even when some citations or claims are supported, never publish the
        # remaining confident factual prose while unsupported claims survive.
        # Keep the evidence and verifier diagnostics so the UI can explain the
        # abstention and a researcher can refine the request.
        language = str(
            request_payload.get("ui_language")
            or request_payload.get("answer_language")
            or "vi"
        ).lower()
        safe_answer = (
            "CLARA could not retrieve enough verifiable evidence for this research "
            "request. No clinical conclusion is released. Please retry or narrow the "
            "question; for care decisions, consult a qualified clinician."
            if language == "en"
            else "CLARA chưa truy xuất được đủ bằng chứng có thể kiểm chứng cho yêu cầu "
            "nghiên cứu này. Hệ thống không phát hành kết luận y khoa. Vui lòng thử lại "
            "hoặc thu hẹp câu hỏi; với quyết định điều trị, hãy trao đổi với bác sĩ."
        )
        gated["answer"] = safe_answer
        gated["answer_markdown"] = safe_answer
        gated["citations"] = real_citations
        gated["sources"] = real_citations
        # A degraded evidence response is normally an explicit warning.  It
        # must never downgrade an independent safety verdict that already
        # blocked or escalated the request (for example, a contradicted dosage
        # claim from the ML verifier).  The policy action is consumed by the
        # clinical-answer renderer, so preserving the stricter action here is
        # part of the API safety boundary rather than presentation metadata.
        existing_policy_action = str(gated.get("policy_action") or "").strip().lower()
        gated["policy_action"] = (
            existing_policy_action
            if existing_policy_action in {"block", "escalate"}
            else "warn"
        )
    return gated


def _attach_verified_research_presentation(
    result: dict[str, Any], *, request_payload: dict[str, Any]
) -> dict[str, Any]:
    """Attach a deterministic read-only presentation after evidence release.

    The original ``answer_markdown`` and ``citations`` are deliberately left
    untouched. This function adds no medical prose, makes no model call, and
    only exposes citation *identifiers* already present in the released
    envelope. Any failed, malformed, skipped, or unavailable verification path
    keeps the existing safe abstention byte-for-byte from the quality gate.
    """

    settings = get_settings()
    if not settings.research_output_modes_enabled:
        return result

    gated = dict(result)
    quality_gate = gated.get("quality_gate")
    if not isinstance(quality_gate, dict) or quality_gate.get("passed") is not True:
        return gated

    metadata = gated.get("metadata") if isinstance(gated.get("metadata"), dict) else {}
    expected_mode = _resolve_research_output_mode(
        request_payload.get("output_mode"),
        role=str(request_payload.get("role") or ""),
    )
    ml_mode = str(gated.get("output_mode") or metadata.get("output_mode") or "").strip().lower()
    if ml_mode != expected_mode:
        # ML did not pass its independent closed-mode gate. Do not synthesize a
        # presentation from a one-sided configuration rollout.
        return gated

    answer = str(gated.get("answer_markdown") or gated.get("answer") or "").strip()
    citations = gated.get("citations")
    if not answer or not isinstance(citations, list):
        return gated
    citation_ids = [
        str(item.get("source_id") or "").strip()
        for item in citations
        if isinstance(item, dict) and str(item.get("source_id") or "").strip()
    ]
    if len(citation_ids) != len(citations):
        # The quality gate should already reject unresolvable evidence. Keep
        # the old released shape rather than presenting a partial citation map.
        return gated

    # The mode only controls reader chrome: the answer is a reference to the
    # same released markdown and professional mode may reveal the already
    # released citation list. No verifier rows, prompts, confidence values,
    # telemetry, PII, or provider data are copied into this projection.
    gated["presentation"] = {
        "schema_version": "research-presentation-v1",
        "mode": expected_mode,
        "answer_markdown": answer,
        "citation_ids": citation_ids,
        "citation_visibility": "expanded" if expected_mode == "professional" else "compact",
    }
    return gated


def _build_research_run_manifest(
    *, job_id: str, request_payload: dict[str, Any], created_at: datetime
) -> dict[str, Any]:
    """Create a privacy-conscious, immutable description of an execution request."""

    return {
        "schema_version": "1.0",
        "run_id": job_id,
        "created_at": created_at.isoformat(),
        "input_sha256": _canonical_sha256(request_payload),
        "research_mode": _coerce_research_mode(request_payload),
        "source_mode": str(request_payload.get("source_mode") or "hybrid"),
        "retrieval_stack_mode": str(
            request_payload.get("retrieval_stack_mode") or "auto"
        ),
        "answer_language": str(
            request_payload.get("ui_language")
            or request_payload.get("answer_language")
            or "vi"
        ),
        "selected_sources": {
            "knowledge_source_ids": list(request_payload.get("source_ids") or []),
            "source_hubs": list(request_payload.get("source_hub_sources") or []),
            "uploaded_file_ids": list(request_payload.get("uploaded_file_ids") or []),
        },
    }


def _build_evidence_snapshot(
    *, job_id: str, result: dict[str, Any], captured_at: datetime
) -> dict[str, Any]:
    """Freeze the exact evidence references used by a completed synthesis."""

    citations = result.get("citations")
    sources = result.get("sources")
    registry = result.get("citation_registry")
    snapshot_payload = {
        "citations": citations if isinstance(citations, list) else [],
        "sources": sources if isinstance(sources, list) else [],
        "citation_registry": registry if isinstance(registry, list) else [],
        "retrieval_trace": result.get("retrieval_trace")
        if isinstance(result.get("retrieval_trace"), dict | list)
        else None,
    }
    return {
        "schema_version": "1.0",
        "run_id": job_id,
        "captured_at": captured_at.isoformat(),
        "evidence_sha256": _canonical_sha256(snapshot_payload),
        **snapshot_payload,
    }


def _build_fail_soft_response_local(
    fail_soft_payload: dict[str, Any],
    reason: str,
) -> dict[str, Any]:
    response = dict(fail_soft_payload)
    response.setdefault("metadata", {})
    response.setdefault("citations", [])
    response.setdefault("fallback", True)
    response.setdefault("fallback_reason", reason)
    return response


def _estimate_reasoning_phase(elapsed_seconds: float) -> tuple[str, str, int]:
    if elapsed_seconds < 15:
        return ("scope_question", "Đang phân tích câu hỏi và xác định phạm vi.", 15)
    if elapsed_seconds < 35:
        return ("collect_evidence", "Đang truy xuất evidence từ nguồn nội bộ và nguồn live.", 40)
    if elapsed_seconds < 60:
        return ("synthesize_findings", "Đang tổng hợp điểm đồng thuận và phát hiện mâu thuẫn.", 70)
    if elapsed_seconds < 90:
        return ("verification", "Đang kiểm chứng claim theo bằng chứng truy xuất.", 88)
    return ("final_response", "Đang hoàn thiện câu trả lời và chuẩn hóa citation.", 95)


def _invoke_ml_tier2_with_progress(
    *,
    ml_payload: dict[str, Any],
    fail_soft_payload: dict[str, Any] | None,
    heartbeat: Callable[[float, dict[str, Any] | None], None],
) -> dict[str, Any]:
    settings = get_settings()
    url = f"{settings.ml_service_url.rstrip('/')}/v1/research/tier2"
    headers: dict[str, str] = {}
    if settings.ml_internal_api_key.strip():
        headers["X-ML-Internal-Key"] = settings.ml_internal_api_key.strip()
    # The async job boundary must be longer than one ML synthesis attempt, but
    # it must also terminate a stalled dependency instead of pinning a worker
    # for eight minutes.  Failures are surfaced as failed jobs; no response is
    # fabricated or retried through another model.
    timeout_seconds = max(settings.ml_service_timeout_seconds * 1.5, 30.0)
    started = datetime.now(tz=UTC)
    request_kwargs: dict[str, Any] = {"json": ml_payload, "timeout": timeout_seconds}
    if headers:
        request_kwargs["headers"] = headers

    with ThreadPoolExecutor(max_workers=1, thread_name_prefix="research-ml-call") as executor:
        future = executor.submit(httpx.post, url, **request_kwargs)
        while True:
            try:
                response = future.result(timeout=2.0)
                break
            except FutureTimeoutError:
                elapsed = (datetime.now(tz=UTC) - started).total_seconds()
                heartbeat(elapsed, None)
                continue
    if response.status_code >= 500:
        if fail_soft_payload is not None:
            return _build_fail_soft_response_local(
                fail_soft_payload, f"status_{response.status_code}"
            )
        raise RuntimeError(f"ml_upstream_status_{response.status_code}")
    if response.status_code >= 400:
        if fail_soft_payload is not None:
            return _build_fail_soft_response_local(
                fail_soft_payload, f"status_{response.status_code}"
            )
        raise RuntimeError(f"ml_rejected_status_{response.status_code}")

    try:
        data = response.json()
    except ValueError as exc:
        if fail_soft_payload is not None:
            return _build_fail_soft_response_local(fail_soft_payload, "InvalidJSON")
        raise RuntimeError("ml_invalid_json") from exc

    if not isinstance(data, dict):
        if fail_soft_payload is not None:
            return _build_fail_soft_response_local(fail_soft_payload, "UnexpectedPayloadFormat")
        raise RuntimeError("ml_unexpected_payload_format")

    elapsed_seconds = (datetime.now(tz=UTC) - started).total_seconds()
    ml_signal = _extract_ml_progress_signal(data)
    if ml_signal is not None:
        stage, status_text, note = ml_signal
        heartbeat(
            elapsed_seconds,
            {
                "stage": stage,
                "status": status_text,
                "note": note,
                "source": "ml_event",
            },
        )
    return data


def _claim_research_job(db: Session, *, job_id: str, worker_id: str) -> bool:
    """Atomically claim a queued job so API workers cannot execute it twice."""

    now = datetime.now(tz=UTC)
    claimed = db.execute(
        update(ResearchJob)
        .where(ResearchJob.job_id == job_id, ResearchJob.status == "queued")
        .values(
            status="running",
            worker_id=worker_id,
            lease_heartbeat_at=now,
            started_at=now,
            updated_at=now,
            attempt_count=ResearchJob.attempt_count + 1,
            error_text="",
        )
    )
    db.commit()
    return bool(claimed.rowcount)


def _run_research_job(job_id: str) -> None:
    db = SessionLocal()
    worker_id = f"research-{os.getpid()}-{uuid4().hex[:12]}"
    try:
        if not _claim_research_job(db, job_id=job_id, worker_id=worker_id):
            return
        job = db.execute(
            select(ResearchJob).where(ResearchJob.job_id == job_id)
        ).scalar_one_or_none()
        if job is None:
            return
        _append_job_event(
            db,
            job=job,
            stage="dispatch_ml",
            status_text="in_progress",
            note="Đã gửi yêu cầu lên ML service, đang chạy reasoning nhiều bước.",
        )
        request_payload = job.request_payload if isinstance(job.request_payload, dict) else {}
        last_heartbeat_bucket = -1

        def _heartbeat(elapsed_seconds: float, ml_event: dict[str, Any] | None) -> None:
            nonlocal last_heartbeat_bucket
            ml_signal = _extract_stage_status_note(ml_event)
            if ml_signal is not None:
                stage, status_text, note = ml_signal
                _append_job_event(
                    db,
                    job=job,
                    stage=stage,
                    status_text=status_text,
                    note=note or "ML đang cập nhật tiến trình reasoning.",
                    payload={
                        "elapsed_seconds": round(elapsed_seconds, 1),
                        "source": "ml_event",
                        "research_mode": _coerce_research_mode(request_payload),
                        "source_mode": str(request_payload.get("source_mode") or "hybrid"),
                    },
                )
                return

            # Emit heartbeat every ~10s and avoid duplicate messages in same 10s bucket.
            bucket = int(elapsed_seconds // 10)
            if bucket <= 0 or bucket == last_heartbeat_bucket:
                return
            last_heartbeat_bucket = bucket
            phase, note, progress_percent = _estimate_reasoning_phase(elapsed_seconds)
            _append_job_event(
                db,
                job=job,
                stage="reasoning",
                status_text="in_progress",
                note=note,
                payload={
                    "elapsed_seconds": round(elapsed_seconds, 1),
                    "heartbeat_seq": bucket,
                    "phase": phase,
                    "progress_percent": progress_percent,
                    "research_mode": _coerce_research_mode(request_payload),
                    "source_mode": str(request_payload.get("source_mode") or "hybrid"),
                },
            )

        ml_response = _invoke_ml_tier2_with_progress(
            ml_payload=request_payload,
            # Research failures must surface as failures.  Never substitute a
            # generic model/research response for an unavailable upstream run.
            fail_soft_payload=None,
            heartbeat=_heartbeat,
        )
        normalized = _normalize_tier2_response(ml_response)
        normalized = _enforce_request_execution_contract(
            normalized,
            request_payload=request_payload,
        )
        enriched = _attach_research_attribution(normalized)
        enriched = _apply_research_quality_gates(
            enriched,
            request_payload=request_payload,
        )
        enriched = _attach_verified_research_presentation(
            enriched,
            request_payload=request_payload,
        )
        quality_gate = enriched.get("quality_gate")
        _append_job_event(
            db,
            job=job,
            stage="quality_gate",
            status_text="completed"
            if isinstance(quality_gate, dict) and quality_gate.get("passed")
            else "degraded",
            note=(
                "Đã kiểm tra độ bao phủ citation và claim; kết quả đủ điều kiện."
                if isinstance(quality_gate, dict) and quality_gate.get("passed")
                else "Kết quả được gắn nhãn degraded vì chưa đủ bằng chứng kiểm chứng."
            ),
            payload={"quality_gate": quality_gate},
        )
        completed_at = datetime.now(tz=UTC)
        job.result_json = enriched
        job.evidence_snapshot_json = _build_evidence_snapshot(
            job_id=job_id,
            result=enriched,
            captured_at=completed_at,
        )
        job.status = "completed"
        job.completed_at = completed_at
        job.worker_id = None
        job.lease_heartbeat_at = None
        db.add(job)
        db.commit()
        _append_job_event(
            db,
            job=job,
            stage="final_response",
            status_text="completed",
            note="Đã hoàn tất trả lời, có thể render Markdown đầy đủ.",
            payload={
                "fallback_used": bool(enriched.get("fallback") or enriched.get("fallback_reason")),
                "fallback_reason": enriched.get("fallback_reason"),
                "source_errors": enriched.get("source_errors"),
                "verification_matrix": enriched.get("verification_matrix"),
                "unsupported_claims": enriched.get("unsupported_claims"),
                "model_used": enriched.get("model_used"),
                "source_count": len(enriched.get("sources", []))
                if isinstance(enriched.get("sources"), list)
                else 0,
            },
        )
        flow_events = enriched.get("flow_events")
        if isinstance(flow_events, list) and flow_events:
            if isinstance(job.progress_json, dict):
                progress = dict(job.progress_json)
            else:
                progress = _empty_job_progress()
            progress = _merge_progress_with_ml_flow_events(
                progress=progress,
                ml_flow_events=[item for item in flow_events if isinstance(item, dict)],
            )
            if not str(progress.get("status_note") or "").strip():
                progress["status_note"] = "Đã nhận flow events đầy đủ từ ML."
            job.progress_json = json.loads(json.dumps(progress, ensure_ascii=False))
            db.add(job)
            db.commit()
    except Exception as exc:  # pragma: no cover - defensive runtime protection
        try:
            job = db.execute(
                select(ResearchJob).where(ResearchJob.job_id == job_id)
            ).scalar_one_or_none()
            if job is not None:
                job.status = "failed"
                # Do not persist a provider exception verbatim: SDKs and
                # upstream gateways may echo prompt content or request details.
                # Keep the error PII-safe while making production failures
                # actionable.  A bounded execution stage distinguishes a
                # timeout during retrieval from one during report synthesis;
                # provider text, query text and prompt content are never kept.
                job.error_text = (
                    f"research_job_failed:{exc.__class__.__name__}:"
                    f"{_safe_failure_stage(job.progress_json)}"
                )
                job.completed_at = datetime.now(tz=UTC)
                job.worker_id = None
                job.lease_heartbeat_at = None
                db.add(job)
                db.commit()
                _append_job_event(
                    db,
                    job=job,
                    stage="final_response",
                    status_text="failed",
                    note="Research job không thể hoàn tất do lỗi upstream.",
                )
        except Exception:
            pass
    finally:
        db.close()
        with _research_job_lock:
            _research_job_futures.pop(job_id, None)


def _queue_research_job(job_id: str) -> None:
    with _research_job_lock:
        stale_job_ids = [
            future_job_id
            for future_job_id, future in _research_job_futures.items()
            if future.done()
        ]
        for stale_job_id in stale_job_ids:
            _research_job_futures.pop(stale_job_id, None)
        if len(_research_job_futures) >= _RESEARCH_JOB_MAX_PENDING:
            raise RuntimeError("research_job_queue_full")
        future = _research_job_executor.submit(_run_research_job, job_id)
        _research_job_futures[job_id] = future


def _count_pending_research_jobs() -> int:
    with _research_job_lock:
        stale_job_ids = [
            job_id for job_id, future in _research_job_futures.items() if future.done()
        ]
        for job_id in stale_job_ids:
            _research_job_futures.pop(job_id, None)
        return len(_research_job_futures)


def recover_research_jobs_once() -> int:
    """Requeue abandoned leases and dispatch every durable queued job.

    The status transition is database-backed, while ``_claim_research_job`` is
    atomic. Consequently this remains safe when several API processes run the
    recovery pass concurrently.
    """

    cutoff = datetime.now(tz=UTC) - timedelta(seconds=_RESEARCH_JOB_LEASE_SECONDS)
    now = datetime.now(tz=UTC)
    with SessionLocal() as db:
        db.execute(
            update(ResearchJob)
            .where(
                ResearchJob.status == "running",
                or_(
                    ResearchJob.lease_heartbeat_at.is_(None),
                    ResearchJob.lease_heartbeat_at < cutoff,
                ),
            )
            .values(
                status="queued",
                worker_id=None,
                lease_heartbeat_at=None,
                updated_at=now,
                recovery_count=ResearchJob.recovery_count + 1,
            )
        )
        db.commit()
        job_ids = list(
            db.scalars(
                select(ResearchJob.job_id)
                .where(ResearchJob.status == "queued")
                .order_by(ResearchJob.created_at.asc())
                .limit(_RESEARCH_JOB_MAX_PENDING)
            ).all()
        )

    dispatched = 0
    for durable_job_id in job_ids:
        try:
            _queue_research_job(str(durable_job_id))
            dispatched += 1
        except RuntimeError:
            break
    return dispatched


def start_research_job_recovery() -> None:
    """Start the daemon that makes queued/restarted research runs recoverable."""

    global _research_recovery_started
    with _research_recovery_lock:
        if _research_recovery_started:
            return
        _research_recovery_started = True

    def _sweep() -> None:
        while True:
            try:
                recover_research_jobs_once()
            except Exception:
                # Recovery is best-effort per sweep; the next interval retries.
                pass
            time.sleep(_RESEARCH_JOB_RECOVERY_POLL_SECONDS)

    Thread(target=_sweep, name="research-job-recovery", daemon=True).start()


_SOURCE_HUB_CATALOG: tuple[SourceHubCatalogEntry, ...] = (
    SourceHubCatalogEntry(
        key="pubmed",
        label="PubMed",
        description="NCBI PubMed biomedical literature via E-utilities",
        docs_url="https://www.ncbi.nlm.nih.gov/books/NBK25501/",
        default_query="diabetes type 2 guideline",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="rxnorm",
        label="RxNorm",
        description="NLM RxNorm normalized clinical drug names via RxNav",
        docs_url="https://lhncbc.nlm.nih.gov/RxNav/APIs/index.html",
        default_query="metformin",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="openfda",
        label="openFDA",
        description="US FDA drug label and safety data",
        docs_url="https://open.fda.gov/apis/",
        default_query="statin",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="dailymed",
        label="DailyMed",
        description="NLM/FDA SPL drug label feed",
        docs_url="https://dailymed.nlm.nih.gov/dailymed/webservices-help.cfm",
        default_query="warfarin",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="europepmc",
        label="Europe PMC",
        description="Biomedical literature and abstracts",
        docs_url="https://europepmc.org/RestfulWebService",
        default_query="warfarin ibuprofen interaction",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="semantic_scholar",
        label="Semantic Scholar",
        description="Academic graph search for biomedical papers",
        docs_url="https://api.semanticscholar.org/api-docs/graph",
        default_query="warfarin nsaid bleeding risk",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="clinicaltrials",
        label="ClinicalTrials.gov",
        description="Clinical study registry and metadata",
        docs_url="https://clinicaltrials.gov/data-api/about-api",
        default_query="warfarin interaction",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="vn_moh",
        label="Bộ Y tế Việt Nam",
        description="Tin tức và văn bản điều hành chính thức từ Cổng thông tin Bộ Y tế",
        docs_url="https://moh.gov.vn/",
        default_query="huong dan chan doan dieu tri",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="vn_kcb",
        label="Cục Quản lý Khám chữa bệnh",
        description="Thông báo, công văn và hướng dẫn khám chữa bệnh từ kcb.vn",
        docs_url="https://kcb.vn/",
        default_query="huong dan kham chua benh",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="vn_canhgiacduoc",
        label="Cảnh giác Dược Quốc gia",
        description="Bản tin cảnh giác dược và theo dõi phản ứng có hại của thuốc",
        docs_url="https://canhgiacduoc.org.vn/",
        default_query="canh giac duoc ADR",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="vn_vbpl_byt",
        label="VBPL Bộ Y tế",
        description="Văn bản pháp quy lĩnh vực y tế trên hệ thống VBPL",
        docs_url="https://vbpl.vn/boyte/Pages/home.aspx",
        default_query="thong tu bo y te",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="vn_dav",
        label="Cục Quản lý Dược Việt Nam",
        description="Thông tin quản lý dược, công bố và thông báo chuyên ngành từ dav.gov.vn",
        docs_url="https://dav.gov.vn/",
        default_query="thu hoi thuoc",
        supports_live_sync=True,
    ),
    SourceHubCatalogEntry(
        key="davidrug",
        label="DAVIDrug",
        description="Cục Quản lý Dược Việt Nam (public web data fallback)",
        docs_url="https://dichvucong.dav.gov.vn/congbothuoc/index",
        default_query="paracetamol",
        supports_live_sync=True,
    ),
)
_SOURCE_HUB_CATALOG_BY_KEY: dict[str, SourceHubCatalogEntry] = {
    entry.key: entry for entry in _SOURCE_HUB_CATALOG
}
_SUPPORTED_SOURCE_HUB_SOURCE_KEYS: set[str] = set(_SOURCE_HUB_CATALOG_BY_KEY.keys())


def _source_hub_setting_key(owner_user_id: int) -> str:
    return f"{_SOURCE_HUB_SETTING_KEY}:{owner_user_id}"


def _to_text(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, int | float) and not isinstance(value, bool):
        return str(value)
    return ""


def _to_bool(value: Any, *, default: bool) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, int | float) and not isinstance(value, bool):
        return value != 0
    if isinstance(value, str):
        return value.strip().lower() in _TRUE_VALUES
    return default


def _serialize_source_hub_catalog(entries: list[SourceHubCatalogEntry]) -> dict[str, Any]:
    return {
        "entries": [entry.model_dump() for entry in entries],
        "updated_at": datetime.now(tz=UTC).isoformat(),
    }


def _normalize_source_hub_catalog_entries(raw_entries: Any) -> list[SourceHubCatalogEntry]:
    if not isinstance(raw_entries, list):
        return []
    normalized: list[SourceHubCatalogEntry] = []
    seen_keys: set[str] = set()
    for row in raw_entries:
        if not isinstance(row, dict):
            continue
        key = _to_text(row.get("key")).lower()
        if key not in _SUPPORTED_SOURCE_HUB_SOURCE_KEYS or key in seen_keys:
            continue
        fallback = _SOURCE_HUB_CATALOG_BY_KEY.get(key)
        try:
            item = SourceHubCatalogEntry(
                key=key,  # type: ignore[arg-type]
                label=_to_text(row.get("label")) or (fallback.label if fallback else key.upper()),
                description=_to_text(row.get("description"))
                or (fallback.description if fallback else ""),
                docs_url=_to_text(row.get("docs_url")) or (fallback.docs_url if fallback else None),
                default_query=_to_text(row.get("default_query"))
                or (fallback.default_query if fallback else None),
                supports_live_sync=_to_bool(
                    row.get("supports_live_sync"),
                    default=(fallback.supports_live_sync if fallback else True),
                ),
            )
        except Exception:
            continue
        normalized.append(item)
        seen_keys.add(key)
    return normalized


def _save_source_hub_catalog(
    db: Session, entries: list[SourceHubCatalogEntry]
) -> list[SourceHubCatalogEntry]:
    dedup: dict[str, SourceHubCatalogEntry] = {}
    for entry in entries:
        dedup[entry.key] = entry
    ordered_entries = list(dedup.values())

    setting = db.execute(
        select(SystemSetting).where(SystemSetting.key == _SOURCE_HUB_CATALOG_SETTING_KEY)
    ).scalar_one_or_none()
    if setting is None:
        setting = SystemSetting(
            key=_SOURCE_HUB_CATALOG_SETTING_KEY,
            value_json=_serialize_source_hub_catalog(ordered_entries),
        )
        db.add(setting)
    else:
        setting.value_json = _serialize_source_hub_catalog(ordered_entries)
        db.add(setting)
    db.commit()
    return ordered_entries


def _load_source_hub_catalog(db: Session) -> list[SourceHubCatalogEntry]:
    setting = db.execute(
        select(SystemSetting).where(SystemSetting.key == _SOURCE_HUB_CATALOG_SETTING_KEY)
    ).scalar_one_or_none()
    if setting and isinstance(setting.value_json, dict):
        normalized = _normalize_source_hub_catalog_entries(setting.value_json.get("entries"))
        if normalized:
            return normalized
    return _save_source_hub_catalog(db, list(_SOURCE_HUB_CATALOG))


def _normalize_source_hub_record(record: dict[str, Any]) -> SourceHubRecord | None:
    source = _to_text(record.get("source")).lower()
    if source not in _SUPPORTED_SOURCE_HUB_SOURCE_KEYS:
        return None
    title = _to_text(record.get("title"))
    if not title:
        return None

    record_id = _to_text(record.get("id")) or str(uuid4())
    metadata_raw = record.get("metadata")
    metadata = metadata_raw if isinstance(metadata_raw, dict) else {}

    return SourceHubRecord(
        id=record_id,
        source=source,  # type: ignore[arg-type]
        title=title,
        url=_to_text(record.get("url")) or None,
        snippet=_to_text(record.get("snippet")) or None,
        external_id=_to_text(record.get("external_id")) or None,
        query=_to_text(record.get("query")) or None,
        published_at=_to_text(record.get("published_at")) or None,
        synced_at=_to_text(record.get("synced_at")) or None,
        metadata=metadata,
    )


def _parse_iso_datetime(value: str | None) -> datetime | None:
    if not value:
        return None
    normalized = value.strip()
    if not normalized:
        return None
    if normalized.endswith("Z"):
        normalized = f"{normalized[:-1]}+00:00"
    try:
        parsed = datetime.fromisoformat(normalized)
    except ValueError:
        return None
    if parsed.tzinfo is None:
        return parsed.replace(tzinfo=UTC)
    return parsed


def _serialize_source_hub_db_row(row: FederatedSourceRecord) -> SourceHubRecord | None:
    payload = {
        "id": row.record_id,
        "source": row.source,
        "title": row.title,
        "url": row.url or None,
        "snippet": row.snippet or None,
        "external_id": row.external_id or None,
        "query": row.query or None,
        "published_at": row.published_at or None,
        "synced_at": row.synced_at.isoformat() if row.synced_at else None,
        "metadata": row.metadata_json if isinstance(row.metadata_json, dict) else {},
    }
    return _normalize_source_hub_record(payload)


def _load_source_hub_records_from_legacy_settings(
    db: Session, owner_user_id: int
) -> list[SourceHubRecord]:
    setting = db.execute(
        select(SystemSetting).where(SystemSetting.key == _source_hub_setting_key(owner_user_id))
    ).scalar_one_or_none()
    if not setting or not isinstance(setting.value_json, dict):
        return []

    raw_records = setting.value_json.get("records")
    if not isinstance(raw_records, list):
        return []

    parsed: list[SourceHubRecord] = []
    for item in raw_records:
        if not isinstance(item, dict):
            continue
        normalized = _normalize_source_hub_record(item)
        if normalized is not None:
            parsed.append(normalized)
    return parsed


def _load_source_hub_records(db: Session, owner_user_id: int) -> list[SourceHubRecord]:
    rows = (
        db.execute(
            select(FederatedSourceRecord)
            .where(FederatedSourceRecord.owner_user_id == owner_user_id)
            .order_by(FederatedSourceRecord.synced_at.desc(), FederatedSourceRecord.id.desc())
            .limit(_SOURCE_HUB_MAX_RECORDS)
        )
        .scalars()
        .all()
    )
    if rows:
        parsed = []
        for row in rows:
            normalized = _serialize_source_hub_db_row(row)
            if normalized is not None:
                parsed.append(normalized)
        return parsed

    # Legacy fallback: migrate old JSON-in-settings payload into dedicated table.
    legacy_records = _load_source_hub_records_from_legacy_settings(db, owner_user_id)
    if legacy_records:
        _save_source_hub_records(db, owner_user_id, legacy_records)
    return legacy_records


def _save_source_hub_records(
    db: Session, owner_user_id: int, records: list[SourceHubRecord]
) -> None:
    pruned = records[:_SOURCE_HUB_MAX_RECORDS]
    db.query(FederatedSourceRecord).filter(
        FederatedSourceRecord.owner_user_id == owner_user_id
    ).delete(synchronize_session=False)

    now = datetime.now(tz=UTC)
    for record in pruned:
        parsed_synced_at = _parse_iso_datetime(record.synced_at) or now
        db.add(
            FederatedSourceRecord(
                owner_user_id=owner_user_id,
                record_id=record.id,
                source=record.source,
                title=record.title,
                url=record.url or "",
                snippet=record.snippet or "",
                external_id=record.external_id or "",
                query=record.query or "",
                published_at=record.published_at or "",
                synced_at=parsed_synced_at,
                metadata_json=record.metadata if isinstance(record.metadata, dict) else {},
            )
        )
    db.commit()


def _merge_source_hub_records(
    existing: list[SourceHubRecord], incoming: list[SourceHubRecord]
) -> list[SourceHubRecord]:
    dedup: dict[str, SourceHubRecord] = {}
    for item in [*incoming, *existing]:
        dedup[item.id] = item

    merged = sorted(
        dedup.values(),
        key=lambda record: record.synced_at or "",
        reverse=True,
    )
    return merged[:_SOURCE_HUB_MAX_RECORDS]


def _http_get_json(url: str, *, params: dict[str, Any] | None = None) -> dict[str, Any]:
    with httpx.Client(timeout=_SOURCE_HUB_TIMEOUT_SECONDS) as client:
        response = client.get(url, params=params)
    response.raise_for_status()
    payload = response.json()
    return payload if isinstance(payload, dict) else {}


def _http_get_text(url: str, *, params: dict[str, Any] | None = None) -> str:
    with httpx.Client(timeout=_SOURCE_HUB_TIMEOUT_SECONDS) as client:
        response = client.get(url, params=params)
    response.raise_for_status()
    return response.text


def _http_post_json(url: str, *, payload: dict[str, Any]) -> dict[str, Any]:
    with httpx.Client(timeout=_SOURCE_HUB_TIMEOUT_SECONDS) as client:
        response = client.post(url, json=payload)
    response.raise_for_status()
    body = response.json()
    return body if isinstance(body, dict) else {}


class _LightAnchorExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.items: list[dict[str, str]] = []
        self._skip_depth = 0
        self._in_anchor = False
        self._anchor_href = ""
        self._anchor_title = ""
        self._anchor_text_chunks: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        normalized_tag = tag.lower()
        if normalized_tag in {"script", "style", "noscript"}:
            self._skip_depth += 1
            return
        if self._skip_depth > 0:
            return
        if normalized_tag != "a":
            return
        attributes = {str(key).lower(): _to_text(value) for key, value in attrs if key}
        self._in_anchor = True
        self._anchor_href = _to_text(attributes.get("href"))
        self._anchor_title = _to_text(attributes.get("title")) or _to_text(
            attributes.get("aria-label")
        )
        self._anchor_text_chunks = []

    def handle_endtag(self, tag: str) -> None:
        normalized_tag = tag.lower()
        if normalized_tag in {"script", "style", "noscript"}:
            if self._skip_depth > 0:
                self._skip_depth -= 1
            return
        if self._skip_depth > 0:
            return
        if normalized_tag != "a" or not self._in_anchor:
            return

        text = _normalize_html_text(" ".join(self._anchor_text_chunks))
        title = _normalize_html_text(self._anchor_title) or text
        if self._anchor_href and title:
            self.items.append(
                {
                    "href": self._anchor_href,
                    "title": title,
                    "snippet": text[:_SOURCE_HUB_SNIPPET_CHAR_LIMIT],
                }
            )

        self._in_anchor = False
        self._anchor_href = ""
        self._anchor_title = ""
        self._anchor_text_chunks = []

    def handle_data(self, data: str) -> None:
        if self._skip_depth > 0 or not self._in_anchor:
            return
        text = _normalize_html_text(data)
        if text:
            self._anchor_text_chunks.append(text)


def _normalize_html_text(value: str) -> str:
    cleaned = unescape(value or "")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def _build_query_terms(query: str) -> list[str]:
    terms = [part for part in re.split(r"[^\w]+", query.lower()) if len(part.strip()) >= 2]
    if not terms and query.strip():
        return [query.strip().lower()]
    return list(dict.fromkeys(terms))


def _resolve_source_hub_url(page_url: str, href: str) -> str:
    raw_href = href.strip()
    if not raw_href or raw_href.startswith("#"):
        return ""
    if raw_href.lower().startswith(("javascript:", "mailto:", "tel:")):
        return ""
    resolved = urljoin(page_url, raw_href)
    parsed = urlparse(resolved)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return ""
    return resolved


def _extract_anchor_candidates(html_text: str, page_url: str) -> list[dict[str, str]]:
    parser = _LightAnchorExtractor()
    parser.feed(html_text)
    parser.close()

    candidates: list[dict[str, str]] = []
    for item in parser.items:
        title = _to_text(item.get("title"))
        href = _to_text(item.get("href"))
        if not title or not href:
            continue
        resolved_url = _resolve_source_hub_url(page_url, href)
        if not resolved_url:
            continue
        snippet = _to_text(item.get("snippet"))
        if snippet == title:
            snippet = ""
        candidates.append(
            {
                "title": title,
                "url": resolved_url,
                "snippet": snippet[:_SOURCE_HUB_SNIPPET_CHAR_LIMIT],
            }
        )
    return candidates


def _extract_published_date(value: str) -> str | None:
    for pattern, year_idx, month_idx, day_idx in (
        (r"(?<!\d)(20\d{2})[-/.](\d{1,2})[-/.](\d{1,2})(?!\d)", 1, 2, 3),
        (r"(?<!\d)(\d{1,2})[-/.](\d{1,2})[-/.](20\d{2})(?!\d)", 3, 2, 1),
    ):
        matched = re.search(pattern, value)
        if not matched:
            continue
        try:
            year = int(matched.group(year_idx))
            month = int(matched.group(month_idx))
            day = int(matched.group(day_idx))
            if 1 <= month <= 12 and 1 <= day <= 31:
                return f"{year:04d}-{month:02d}-{day:02d}"
        except ValueError:
            continue
    return None


def _external_id_from_url(url: str) -> str:
    parsed = urlparse(url)
    path_value = f"{parsed.netloc}{parsed.path}?{parsed.query}".strip("?").lower()
    normalized = re.sub(r"[^a-z0-9]+", "-", path_value).strip("-")
    if normalized:
        return normalized[:120]
    return hashlib.sha1(url.encode("utf-8")).hexdigest()[:32]


def _build_vn_source_urls(source: str, query: str) -> list[str]:
    source_def = _VN_HTML_SOURCE_DEFINITIONS.get(source)
    if not isinstance(source_def, dict):
        return []
    search_urls = source_def.get("search_urls")
    if not isinstance(search_urls, list | tuple):
        return []

    query_text = query.strip()
    query_q = quote(query_text)
    urls: list[str] = []
    for template in search_urls:
        if not isinstance(template, str):
            continue
        text = template.strip()
        if not text:
            continue
        try:
            candidate = text.format(query=query_text, query_q=query_q)
        except Exception:
            candidate = text
        urls.append(candidate)
    return list(dict.fromkeys(urls))


def _fetch_vn_html_source_records(
    source: SourceHubSourceKey,
    query: str,
    limit: int,
    synced_at: str,
) -> tuple[list[SourceHubRecord], list[str]]:
    source_def = _VN_HTML_SOURCE_DEFINITIONS.get(source)
    if not isinstance(source_def, dict):
        return [], [f"Nguồn không được hỗ trợ: {source}"]

    source_label = _to_text(source_def.get("label")) or source
    safe_limit = max(1, min(500, int(limit)))
    query_terms = _build_query_terms(query)
    page_urls = _build_vn_source_urls(source, query)

    records: list[SourceHubRecord] = []
    warnings: list[str] = []
    seen_urls: set[str] = set()
    seen_ids: set[str] = set()
    crawl_errors = 0

    for page_url in page_urls:
        try:
            html = _http_get_text(page_url)
        except Exception:
            crawl_errors += 1
            continue

        for item in _extract_anchor_candidates(html, page_url):
            title = _to_text(item.get("title"))
            record_url = _to_text(item.get("url"))
            snippet = _to_text(item.get("snippet"))
            if not title or not record_url:
                continue
            if record_url in seen_urls:
                continue

            haystack = f"{title} {snippet} {record_url}".lower()
            if query_terms and not any(term in haystack for term in query_terms):
                continue

            external_id = _external_id_from_url(record_url)
            record_id = f"{source}:{external_id}"
            if record_id in seen_ids:
                continue

            seen_ids.add(record_id)
            seen_urls.add(record_url)
            records.append(
                SourceHubRecord(
                    id=record_id,
                    source=source,
                    title=title,
                    url=record_url,
                    snippet=snippet[:_SOURCE_HUB_SNIPPET_CHAR_LIMIT] or None,
                    external_id=external_id,
                    query=query,
                    published_at=_extract_published_date(f"{title} {snippet} {record_url}"),
                    synced_at=synced_at,
                    metadata={
                        "crawl_url": page_url,
                        "source_label": source_label,
                        "query_terms": query_terms[:8],
                    },
                )
            )
            if len(records) >= safe_limit:
                break
        if len(records) >= safe_limit:
            break

    if not records:
        if crawl_errors >= len(page_urls) and page_urls:
            warnings.append(f"{source_label} hiện không truy cập được để crawl HTML.")
        else:
            warnings.append(f"{source_label} không có kết quả phù hợp cho query này.")
    elif crawl_errors:
        warnings.append(
            f"{source_label} có một số URL crawl lỗi ({crawl_errors}/{len(page_urls)})."
        )

    return records[:safe_limit], warnings


def _ncbi_eutils_params(params: dict[str, Any]) -> dict[str, Any]:
    """Augment NCBI E-utilities params with the API key when configured.

    A configured ``NCBI_API_KEY`` raises the per-IP rate limit from 3 to 10
    requests/second. When unset, requests are sent unauthenticated (the source
    hub still works, just at the lower anonymous rate limit).
    """
    api_key = _research_settings.ncbi_api_key.strip()
    if not api_key:
        return params
    return {**params, "api_key": api_key}


def _fetch_pubmed_records(
    query: str, limit: int, synced_at: str
) -> tuple[list[SourceHubRecord], list[str]]:
    warnings: list[str] = []
    search = _http_get_json(
        "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi",
        params=_ncbi_eutils_params(
            {"db": "pubmed", "term": query, "retmax": limit, "retmode": "json"}
        ),
    )
    search_result = search.get("esearchresult")
    if not isinstance(search_result, dict):
        return [], ["PubMed trả dữ liệu không đúng định dạng esearchresult."]

    id_list_raw = search_result.get("idlist")
    id_list = (
        [str(item).strip() for item in id_list_raw if str(item).strip()]
        if isinstance(id_list_raw, list)
        else []
    )
    if not id_list:
        return [], ["PubMed không có kết quả phù hợp cho query hiện tại."]

    summary = _http_get_json(
        "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi",
        params=_ncbi_eutils_params(
            {"db": "pubmed", "id": ",".join(id_list[:limit]), "retmode": "json"}
        ),
    )
    result = summary.get("result")
    if not isinstance(result, dict):
        return [], ["PubMed không trả được summary chi tiết."]

    uids_raw = result.get("uids")
    uids = (
        [str(item).strip() for item in uids_raw if str(item).strip()]
        if isinstance(uids_raw, list)
        else []
    )

    records: list[SourceHubRecord] = []
    for uid in uids[:limit]:
        item = result.get(uid)
        if not isinstance(item, dict):
            continue
        title = _to_text(item.get("title"))
        if not title:
            continue
        journal = _to_text(item.get("fulljournalname")) or _to_text(item.get("source"))
        pubdate = _to_text(item.get("pubdate"))
        snippet = " | ".join(part for part in [journal, pubdate] if part).strip()
        records.append(
            SourceHubRecord(
                id=f"pubmed:{uid}",
                source="pubmed",
                title=title,
                url=f"https://pubmed.ncbi.nlm.nih.gov/{uid}/",
                snippet=snippet or None,
                external_id=uid,
                query=query,
                published_at=pubdate or None,
                synced_at=synced_at,
                metadata={
                    "authors": item.get("authors"),
                    "pubtype": item.get("pubtype"),
                    "doi": item.get("elocationid"),
                },
            )
        )

    if not records:
        warnings.append("PubMed trả về bản ghi rỗng sau bước summary.")
    return records, warnings


def _fetch_rxnorm_records(
    query: str, limit: int, synced_at: str
) -> tuple[list[SourceHubRecord], list[str]]:
    warnings: list[str] = []
    payload = _http_get_json(
        "https://rxnav.nlm.nih.gov/REST/approximateTerm.json",
        params={"term": query, "maxEntries": limit},
    )
    group = payload.get("approximateGroup")
    if not isinstance(group, dict):
        return [], ["RxNorm không trả về approximateGroup."]
    candidates = group.get("candidate")
    if not isinstance(candidates, list) or not candidates:
        return [], ["RxNorm không có candidate cho query này."]

    records: list[SourceHubRecord] = []
    for index, item in enumerate(candidates[:limit]):
        if not isinstance(item, dict):
            continue
        rxcui = _to_text(item.get("rxcui"))
        rank = _to_text(item.get("rank"))
        score = _to_text(item.get("score"))
        name = _to_text(item.get("name")) or f"RxNorm candidate {index + 1}"
        records.append(
            SourceHubRecord(
                id=f"rxnorm:{rxcui or index}",
                source="rxnorm",
                title=name,
                url=f"https://mor.nlm.nih.gov/RxNav/search?searchBy=RXCUI&searchTerm={rxcui}"
                if rxcui
                else None,
                snippet=f"rxcui={rxcui or '-'} | score={score or '-'} | rank={rank or '-'}",
                external_id=rxcui or None,
                query=query,
                published_at=None,
                synced_at=synced_at,
                metadata=item,
            )
        )

    return records, warnings


def _fetch_openfda_records(
    query: str, limit: int, synced_at: str
) -> tuple[list[SourceHubRecord], list[str]]:
    warnings: list[str] = []
    escaped = query.replace('"', '\\"')
    payload = _http_get_json(
        "https://api.fda.gov/drug/label.json",
        params={"search": f'openfda.brand_name:"{escaped}"', "limit": limit},
    )
    results = payload.get("results")
    if not isinstance(results, list) or not results:
        return [], ["openFDA không có kết quả cho query này."]

    records: list[SourceHubRecord] = []
    for index, item in enumerate(results[:limit]):
        if not isinstance(item, dict):
            continue
        openfda = item.get("openfda")
        openfda_obj = openfda if isinstance(openfda, dict) else {}
        brand_names = openfda_obj.get("brand_name")
        title = (
            _to_text(brand_names[0]) if isinstance(brand_names, list) and brand_names else ""
        ) or f"openFDA label {index + 1}"
        set_id_list = openfda_obj.get("set_id")
        set_id = _to_text(set_id_list[0]) if isinstance(set_id_list, list) and set_id_list else ""
        purpose = item.get("purpose")
        warning_text = item.get("warnings")
        snippet = _to_text(purpose[0]) if isinstance(purpose, list) and purpose else ""
        if not snippet:
            snippet = (
                _to_text(warning_text[0]) if isinstance(warning_text, list) and warning_text else ""
            )
        records.append(
            SourceHubRecord(
                id=f"openfda:{set_id or index}",
                source="openfda",
                title=title,
                url=None,
                snippet=snippet[:280] or None,
                external_id=set_id or None,
                query=query,
                published_at=None,
                synced_at=synced_at,
                metadata={"openfda": openfda_obj},
            )
        )
    return records, warnings


def _fetch_dailymed_records(
    query: str, limit: int, synced_at: str
) -> tuple[list[SourceHubRecord], list[str]]:
    warnings: list[str] = []
    escaped_query = quote(query.strip())
    payload = _http_get_json(
        f"https://dailymed.nlm.nih.gov/dailymed/services/v1/drugname/{escaped_query}/spls.json"
    )
    rows = payload.get("data")
    if not isinstance(rows, list) or not rows:
        return [], ["DailyMed không có kết quả cho query này."]

    records: list[SourceHubRecord] = []
    for index, item in enumerate(rows[:limit]):
        if not isinstance(item, list):
            continue
        set_id = _to_text(item[0] if len(item) > 0 else "")
        title = _to_text(item[1] if len(item) > 1 else "") or f"DailyMed label {index + 1}"
        version = _to_text(item[2] if len(item) > 2 else "")
        published = _to_text(item[3] if len(item) > 3 else "")
        records.append(
            SourceHubRecord(
                id=f"dailymed:{set_id or index}",
                source="dailymed",
                title=title,
                url=(
                    f"https://dailymed.nlm.nih.gov/dailymed/drugInfo.cfm?setid={set_id}"
                    if set_id
                    else "https://dailymed.nlm.nih.gov/"
                ),
                snippet=" | ".join(part for part in [version, published] if part) or None,
                external_id=set_id or None,
                query=query,
                published_at=published or None,
                synced_at=synced_at,
                metadata={},
            )
        )

    if not records:
        warnings.append("DailyMed trả dữ liệu không hợp lệ sau khi parse.")
    return records, warnings


def _fetch_europepmc_records(
    query: str, limit: int, synced_at: str
) -> tuple[list[SourceHubRecord], list[str]]:
    warnings: list[str] = []
    payload = _http_get_json(
        "https://www.ebi.ac.uk/europepmc/webservices/rest/search",
        params={
            "query": query,
            "format": "json",
            "resultType": "core",
            "pageSize": max(1, min(50, int(limit))),
        },
    )
    result_list = payload.get("resultList")
    results = result_list.get("result") if isinstance(result_list, dict) else []
    if not isinstance(results, list) or not results:
        return [], ["Europe PMC không có kết quả cho query này."]

    records: list[SourceHubRecord] = []
    for index, item in enumerate(results[:limit]):
        if not isinstance(item, dict):
            continue
        source = _to_text(item.get("source")).lower() or "europepmc"
        source_id = _to_text(item.get("id"))
        title = _to_text(item.get("title")) or f"Europe PMC record {index + 1}"
        journal = _to_text(item.get("journalTitle"))
        pub_year = _to_text(item.get("pubYear"))
        if source == "med" and source_id:
            url = f"https://pubmed.ncbi.nlm.nih.gov/{source_id}/"
        elif source_id:
            url = f"https://europepmc.org/article/{source.upper()}/{source_id}"
        else:
            url = "https://europepmc.org/"
        records.append(
            SourceHubRecord(
                id=f"europepmc:{source}:{source_id or index}",
                source="europepmc",
                title=title,
                url=url,
                snippet=" | ".join(part for part in [journal, pub_year] if part) or None,
                external_id=source_id or None,
                query=query,
                published_at=pub_year or None,
                synced_at=synced_at,
                metadata={"source_provider": source},
            )
        )

    if not records:
        warnings.append("Europe PMC trả dữ liệu không hợp lệ sau khi parse.")
    return records, warnings


def _fetch_semantic_scholar_records(
    query: str, limit: int, synced_at: str
) -> tuple[list[SourceHubRecord], list[str]]:
    warnings: list[str] = []
    payload = _http_get_json(
        "https://api.semanticscholar.org/graph/v1/paper/search",
        params={
            "query": query,
            "limit": max(1, min(50, int(limit))),
            "fields": "paperId,title,year,url,venue,journal",
        },
    )
    rows = payload.get("data")
    if not isinstance(rows, list) or not rows:
        return [], ["Semantic Scholar không có kết quả cho query này."]

    records: list[SourceHubRecord] = []
    for index, item in enumerate(rows[:limit]):
        if not isinstance(item, dict):
            continue
        paper_id = _to_text(item.get("paperId"))
        title = _to_text(item.get("title")) or f"Semantic Scholar record {index + 1}"
        year = _to_text(item.get("year"))
        url = _to_text(item.get("url"))
        venue = _to_text(item.get("venue"))
        journal_obj = item.get("journal")
        journal = _to_text(journal_obj.get("name")) if isinstance(journal_obj, dict) else ""
        records.append(
            SourceHubRecord(
                id=f"semantic_scholar:{paper_id or index}",
                source="semantic_scholar",
                title=title,
                url=(
                    url
                    or (f"https://www.semanticscholar.org/paper/{paper_id}" if paper_id else None)
                ),
                snippet=" | ".join(part for part in [venue, journal, year] if part) or None,
                external_id=paper_id or None,
                query=query,
                published_at=year or None,
                synced_at=synced_at,
                metadata={},
            )
        )

    if not records:
        warnings.append("Semantic Scholar trả dữ liệu không hợp lệ sau khi parse.")
    return records, warnings


def _fetch_clinicaltrials_records(
    query: str, limit: int, synced_at: str
) -> tuple[list[SourceHubRecord], list[str]]:
    warnings: list[str] = []
    payload = _http_get_json(
        "https://clinicaltrials.gov/api/v2/studies",
        params={
            "query.term": query,
            "pageSize": max(1, min(50, int(limit))),
            "format": "json",
        },
    )
    studies = payload.get("studies")
    if not isinstance(studies, list) or not studies:
        return [], ["ClinicalTrials.gov không có kết quả cho query này."]

    records: list[SourceHubRecord] = []
    for index, item in enumerate(studies[:limit]):
        if not isinstance(item, dict):
            continue
        protocol = item.get("protocolSection")
        if not isinstance(protocol, dict):
            continue
        identification = protocol.get("identificationModule")
        status_module = protocol.get("statusModule")
        identification = identification if isinstance(identification, dict) else {}
        status_module = status_module if isinstance(status_module, dict) else {}

        nct_id = _to_text(identification.get("nctId"))
        title = _to_text(identification.get("briefTitle")) or f"Clinical trial {index + 1}"
        overall_status = _to_text(status_module.get("overallStatus"))
        start_date_obj = status_module.get("startDateStruct")
        start_date = (
            _to_text(start_date_obj.get("date")) if isinstance(start_date_obj, dict) else ""
        )

        records.append(
            SourceHubRecord(
                id=f"clinicaltrials:{nct_id or index}",
                source="clinicaltrials",
                title=title,
                url=(f"https://clinicaltrials.gov/study/{nct_id}" if nct_id else None),
                snippet=" | ".join(part for part in [overall_status, start_date] if part) or None,
                external_id=nct_id or None,
                query=query,
                published_at=start_date or None,
                synced_at=synced_at,
                metadata={},
            )
        )

    if not records:
        warnings.append("ClinicalTrials.gov trả dữ liệu không hợp lệ sau khi parse.")
    return records, warnings


def _fetch_davidrug_records(
    query: str, limit: int, synced_at: str
) -> tuple[list[SourceHubRecord], list[str]]:
    warnings: list[str] = []
    payload = _http_post_json(
        "https://dichvucong.dav.gov.vn/api/services/app/soDangKy/GetAllPublicServerPaging",
        payload={
            "filterText": query,
            "SoDangKyThuoc": {},
            "KichHoat": True,
            "skipCount": 0,
            "maxResultCount": max(1, min(100, int(limit))),
            "sorting": None,
        },
    )
    result_obj = payload.get("result")
    result = result_obj if isinstance(result_obj, dict) else {}
    rows_obj = result.get("items")
    rows = rows_obj if isinstance(rows_obj, list) else []
    if not rows:
        warnings.append("DAVIDrug không có kết quả phù hợp cho query này.")
        return [], warnings

    records: list[SourceHubRecord] = []
    for index, item in enumerate(rows[:limit]):
        if not isinstance(item, dict):
            continue
        external_id = _to_text(item.get("id")) or _to_text(item.get("soDangKy"))
        title = _to_text(item.get("tenThuoc")) or f"DAVIDrug record {index + 1}"
        so_dang_ky = _to_text(item.get("soDangKy"))

        company_obj = item.get("congTyDangKy")
        company = company_obj if isinstance(company_obj, dict) else {}
        registrant = _to_text(company.get("tenCongTyDangKy"))

        info_obj = item.get("thongTinThuocCoBan")
        info = info_obj if isinstance(info_obj, dict) else {}
        active_ingredient = _to_text(info.get("hoatChatChinh"))

        snippet_parts = [
            part
            for part in (
                f"SĐK: {so_dang_ky}" if so_dang_ky else "",
                f"Hoạt chất: {active_ingredient}" if active_ingredient else "",
                f"Đơn vị đăng ký: {registrant}" if registrant else "",
            )
            if part
        ]
        snippet = " | ".join(snippet_parts)[:300] or None

        register_obj = item.get("thongTinDangKyThuoc")
        register = register_obj if isinstance(register_obj, dict) else {}
        published_at = _to_text(register.get("ngayCapSoDangKy")) or None

        record_id = f"davidrug:{external_id or index}"
        records.append(
            SourceHubRecord(
                id=record_id,
                source="davidrug",
                title=title,
                url="https://dichvucong.dav.gov.vn/congbothuoc/index",
                snippet=snippet,
                external_id=external_id or None,
                query=query,
                published_at=published_at,
                synced_at=synced_at,
                metadata={
                    "so_dang_ky": so_dang_ky,
                    "dang_bao_che": _to_text(info.get("dangBaoChe")) or None,
                    "ham_luong": _to_text(info.get("hamLuong")) or None,
                    "active_ingredient": active_ingredient or None,
                    "registrant": registrant or None,
                },
            )
        )

    if not records:
        warnings.append("DAVIDrug trả dữ liệu không hợp lệ sau khi parse.")
    return records, warnings


def _fetch_source_hub_records(
    source: SourceHubSourceKey, query: str, limit: int
) -> tuple[list[SourceHubRecord], list[str]]:
    synced_at = datetime.now(tz=UTC).isoformat()
    if source == "pubmed":
        return _fetch_pubmed_records(query, limit, synced_at)
    if source == "rxnorm":
        return _fetch_rxnorm_records(query, limit, synced_at)
    if source == "openfda":
        return _fetch_openfda_records(query, limit, synced_at)
    if source == "dailymed":
        return _fetch_dailymed_records(query, limit, synced_at)
    if source == "europepmc":
        return _fetch_europepmc_records(query, limit, synced_at)
    if source == "semantic_scholar":
        return _fetch_semantic_scholar_records(query, limit, synced_at)
    if source == "clinicaltrials":
        return _fetch_clinicaltrials_records(query, limit, synced_at)
    if source in {
        "vn_moh",
        "vn_kcb",
        "vn_canhgiacduoc",
        "vn_vbpl_byt",
        "vn_dav",
    }:
        return _fetch_vn_html_source_records(source, query, limit, synced_at)
    if source == "davidrug":
        return _fetch_davidrug_records(query, limit, synced_at)
    return [], [f"Nguồn không được hỗ trợ: {source}"]


@router.get("/conversations")
def list_research_conversations(
    limit: int = 50,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> ResearchConversationListResponse:
    user = _get_user_by_token(db, token)
    safe_limit = max(1, min(200, int(limit)))

    sessions = (
        db.execute(
            select(SessionModel)
            .where(SessionModel.user_id == user.id)
            .order_by(SessionModel.created_at.desc(), SessionModel.id.desc())
            .limit(max(safe_limit * 3, safe_limit))
        )
        .scalars()
        .all()
    )

    items: list[ResearchConversationResponse] = []
    for session_obj in sessions:
        query_obj = db.execute(
            select(QueryModel)
            .where(QueryModel.session_id == session_obj.id)
            .order_by(QueryModel.created_at.desc(), QueryModel.id.desc())
            .limit(1)
        ).scalar_one_or_none()
        if query_obj is None:
            continue
        items.append(_serialize_research_conversation(session_obj=session_obj, query_obj=query_obj))
        if len(items) >= safe_limit:
            break

    return ResearchConversationListResponse(items=items)


@router.post("/conversations")
def create_research_conversation(
    payload: ResearchConversationCreateRequest,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> ResearchConversationResponse:
    user = _get_user_by_token(db, token)
    query_text = payload.query.strip()
    if not query_text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="query không được rỗng.",
        )

    result_payload = _validate_result_payload(payload.result)
    try:
        stored_result = json.dumps({"result": result_payload}, ensure_ascii=False)
    except TypeError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"result chứa dữ liệu không thể lưu JSON: {exc}",
        ) from exc

    session_obj = SessionModel(
        user_id=user.id,
        title=query_text[:255],
    )
    db.add(session_obj)
    db.flush()

    query_obj = QueryModel(
        session_id=session_obj.id,
        role=token.role,
        user_input=query_text,
        response_text=stored_result,
    )
    db.add(query_obj)
    db.commit()
    db.refresh(session_obj)
    db.refresh(query_obj)

    return _serialize_research_conversation(session_obj=session_obj, query_obj=query_obj)


@router.get("/conversations/{conversation_id}/messages")
def list_research_conversation_messages(
    conversation_id: int,
    limit: int = 100,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> ResearchConversationMessagesResponse:
    user = _get_user_by_token(db, token)
    session_obj = db.execute(
        select(SessionModel).where(
            SessionModel.id == conversation_id,
            SessionModel.user_id == user.id,
        )
    ).scalar_one_or_none()
    if session_obj is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conversation không tồn tại.",
        )

    safe_limit = max(1, min(500, int(limit)))
    rows = (
        db.execute(
            select(QueryModel)
            .where(QueryModel.session_id == session_obj.id)
            .order_by(QueryModel.created_at.asc(), QueryModel.id.asc())
            .limit(safe_limit)
        )
        .scalars()
        .all()
    )
    return ResearchConversationMessagesResponse(
        conversation_id=session_obj.id,
        items=[_serialize_research_message(row) for row in rows],
    )


@router.post("/conversations/{conversation_id}/messages")
def append_research_conversation_message(
    conversation_id: int,
    payload: ResearchConversationCreateRequest,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> ResearchConversationResponse:
    user = _get_user_by_token(db, token)
    session_obj = db.execute(
        select(SessionModel).where(
            SessionModel.id == conversation_id,
            SessionModel.user_id == user.id,
        )
    ).scalar_one_or_none()
    if session_obj is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conversation không tồn tại.",
        )

    query_text = payload.query.strip()
    if not query_text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="query không được rỗng.",
        )
    result_payload = _validate_result_payload(payload.result)
    try:
        stored_result = json.dumps({"result": result_payload}, ensure_ascii=False)
    except TypeError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"result chứa dữ liệu không thể lưu JSON: {exc}",
        ) from exc

    query_obj = QueryModel(
        session_id=session_obj.id,
        role=token.role,
        user_input=query_text,
        response_text=stored_result,
    )
    db.add(query_obj)
    session_obj.title = query_text[:255]
    db.add(session_obj)
    db.commit()
    db.refresh(session_obj)
    db.refresh(query_obj)
    return _serialize_research_conversation(session_obj=session_obj, query_obj=query_obj)


@router.delete("/conversations/{conversation_id}")
def delete_research_conversation(
    conversation_id: int,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> dict[str, bool]:
    user = _get_user_by_token(db, token)
    session_obj = db.execute(
        select(SessionModel).where(
            SessionModel.id == conversation_id,
            SessionModel.user_id == user.id,
        )
    ).scalar_one_or_none()
    if session_obj is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conversation không tồn tại.",
        )

    db.delete(session_obj)
    db.commit()
    return {"deleted": True}


@router.get("/knowledge-sources")
def list_knowledge_sources(
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> dict[str, list[KnowledgeSourceResponse]]:
    user = _get_user_by_token(db, token)
    sources = (
        db.execute(
            select(KnowledgeSource)
            .where(KnowledgeSource.owner_user_id == user.id)
            .order_by(KnowledgeSource.updated_at.desc())
        )
        .scalars()
        .all()
    )

    counts_rows = db.execute(
        select(KnowledgeDocument.source_id, func.count(KnowledgeDocument.id))
        .where(KnowledgeDocument.owner_user_id == user.id)
        .group_by(KnowledgeDocument.source_id)
    ).all()
    count_by_source = {int(source_id): int(total) for source_id, total in counts_rows}

    return {
        "items": [
            _serialize_knowledge_source(
                source,
                documents_count=count_by_source.get(source.id, 0),
            )
            for source in sources
        ]
    }


def _audit_kb_mutation(
    db: Session,
    token: TokenPayload,
    action: str,
    *,
    target: str = "",
    outcome: str,
    meta: dict[str, Any] | None = None,
) -> None:
    """Append one admin-action audit row for a knowledge-source mutation.

    Records both success and failure outcomes (Requirement 9.5) with an opaque,
    salted actor reference (the hash of the caller id — never a raw user id /
    email, Requirement 9.3) and a PII-free ``meta`` (counts/flags/mime only).
    The write is a no-op when ``admin_audit_log_enabled`` is off, preserving the
    flags-off baseline (Requirement 12.2); the audit row is committed on its own
    so it is durable independently of the mutation's own transaction.
    """

    record_admin_action(
        db,
        hash_user_ref(token.sub),
        action,
        target=target,
        outcome=outcome,
        meta=meta,
    )
    db.commit()


@router.post("/knowledge-sources")
def create_knowledge_source(
    payload: KnowledgeSourceCreateRequest,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> KnowledgeSourceResponse:
    user = _get_user_by_token(db, token)
    try:
        source = KnowledgeSource(
            owner_user_id=user.id,
            name=payload.name.strip(),
            description=payload.description.strip(),
            is_active=True,
        )
        db.add(source)
        db.commit()
        db.refresh(source)
    except Exception:
        db.rollback()
        _audit_kb_mutation(
            db,
            token,
            ACTION_KB_SOURCE_CREATE,
            outcome=OUTCOME_FAILURE,
            meta={"name_length": len(payload.name.strip())},
        )
        raise
    _audit_kb_mutation(
        db,
        token,
        ACTION_KB_SOURCE_CREATE,
        target=str(source.id),
        outcome=OUTCOME_SUCCESS,
        meta={"name_length": len(payload.name.strip())},
    )
    return _serialize_knowledge_source(source, documents_count=0)


@router.patch("/knowledge-sources/{source_id}")
def update_knowledge_source(
    source_id: int,
    payload: KnowledgeSourceUpdateRequest,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> KnowledgeSourceResponse:
    user = _get_user_by_token(db, token)
    source = _get_owned_source(db, source_id=source_id, owner_user_id=user.id)

    if payload.name is not None:
        source.name = payload.name.strip()
    if payload.description is not None:
        source.description = payload.description.strip()
    if payload.is_active is not None:
        source.is_active = payload.is_active

    db.add(source)
    db.commit()
    db.refresh(source)

    documents_count = db.execute(
        select(func.count(KnowledgeDocument.id)).where(
            KnowledgeDocument.source_id == source.id,
            KnowledgeDocument.owner_user_id == user.id,
        )
    ).scalar_one()

    return _serialize_knowledge_source(source, documents_count=int(documents_count or 0))


@router.delete("/knowledge-sources/{source_id}")
def delete_knowledge_source(
    source_id: int,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> dict[str, bool]:
    user = _get_user_by_token(db, token)
    source = _get_owned_source(db, source_id=source_id, owner_user_id=user.id)
    db.delete(source)
    db.commit()
    return {"deleted": True}


@router.get("/knowledge-sources/{source_id}/documents")
def list_knowledge_documents(
    source_id: int,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> dict[str, list[KnowledgeDocumentResponse]]:
    user = _get_user_by_token(db, token)
    _get_owned_source(db, source_id=source_id, owner_user_id=user.id)
    documents = (
        db.execute(
            select(KnowledgeDocument)
            .where(
                KnowledgeDocument.source_id == source_id,
                KnowledgeDocument.owner_user_id == user.id,
            )
            .order_by(KnowledgeDocument.created_at.desc())
        )
        .scalars()
        .all()
    )

    return {"items": [_serialize_knowledge_document(item) for item in documents]}


@router.post("/knowledge-sources/{source_id}/upload-file")
async def upload_file_to_knowledge_source(
    source_id: int,
    file: UploadFile = File(...),
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    user = _get_user_by_token(db, token)
    try:
        source = _get_owned_source(db, source_id=source_id, owner_user_id=user.id)

        file_name = file.filename or "uploaded-file"
        content_type = file.content_type or "application/octet-stream"
        file_bytes = await read_upload_bytes_with_limit(
            file, max_bytes=_MAX_RESEARCH_UPLOAD_BYTES
        )
        if not file_bytes:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="File upload rỗng")
        verified = _validate_upload_safety(
            file_name=file_name, content_type=content_type, file_bytes=file_bytes
        )
        file_name = verified.filename
        content_type = verified.media_type

        extracted_text, file_kind = _extract_basic_text(file_bytes, file_name, content_type)
        preview = extracted_text[:_PREVIEW_CHAR_LIMIT]
        token_count = _approx_token_count(extracted_text if file_kind == "text" else "")

        document = KnowledgeDocument(
            source_id=source.id,
            owner_user_id=user.id,
            filename=file_name,
            content_type=content_type,
            size=len(file_bytes),
            extracted_text=extracted_text if file_kind == "text" else "",
            preview=preview,
            token_count=token_count,
            is_active=True,
        )
        db.add(document)
        db.commit()
        db.refresh(document)
    except Exception:
        db.rollback()
        _audit_kb_mutation(
            db,
            token,
            ACTION_KB_SOURCE_UPLOAD,
            target=str(source_id),
            outcome=OUTCOME_FAILURE,
        )
        raise

    _audit_kb_mutation(
        db,
        token,
        ACTION_KB_SOURCE_UPLOAD,
        target=str(source.id),
        outcome=OUTCOME_SUCCESS,
        meta={
            "size": len(file_bytes),
            "token_count": token_count,
            "content_type": content_type,
        },
    )

    return {
        "document": _serialize_knowledge_document(document),
        "source_id": source.id,
    }


@router.patch("/documents/{document_id}")
def update_knowledge_document(
    document_id: int,
    payload: KnowledgeDocumentUpdateRequest,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> KnowledgeDocumentResponse:
    user = _get_user_by_token(db, token)
    try:
        document = _get_owned_document(db, document_id=document_id, owner_user_id=user.id)
        document.is_active = payload.is_active
        db.add(document)
        db.commit()
        db.refresh(document)
    except Exception:
        db.rollback()
        _audit_kb_mutation(
            db,
            token,
            ACTION_KB_DOCUMENT_STATUS,
            target=str(document_id),
            outcome=OUTCOME_FAILURE,
            meta={"is_active": payload.is_active},
        )
        raise
    _audit_kb_mutation(
        db,
        token,
        ACTION_KB_DOCUMENT_STATUS,
        target=str(document_id),
        outcome=OUTCOME_SUCCESS,
        meta={"is_active": payload.is_active},
    )
    return _serialize_knowledge_document(document)


@router.post("/upload-file")
async def upload_research_file(
    file: UploadFile = File(...),
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    user = _get_user_by_token(db, token)

    file_name = file.filename or "uploaded-file"
    content_type = file.content_type or "application/octet-stream"
    file_bytes = await read_upload_bytes_with_limit(file, max_bytes=_MAX_RESEARCH_UPLOAD_BYTES)
    if not file_bytes:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="File upload rỗng")
    verified = _validate_upload_safety(
        file_name=file_name, content_type=content_type, file_bytes=file_bytes
    )
    file_name = verified.filename
    content_type = verified.media_type

    extracted_text, file_kind = _extract_basic_text(file_bytes, file_name, content_type)
    preview = extracted_text[:_PREVIEW_CHAR_LIMIT]
    token_count = _approx_token_count(extracted_text if file_kind == "text" else "")
    created_at = datetime.now(tz=UTC).isoformat()
    file_id = str(uuid4())
    stored_text = extracted_text if file_kind == "text" else ""

    if _durable_uploads_enabled():
        # Durable, owner-isolated persistence (R2.1-R2.3, R2.6). A broken but
        # enabled backend surfaces a 503 so the upload is never silently lost.
        store = _build_research_upload_store(db)
        try:
            stored = store.put(
                user.id,
                file_bytes,
                stored_text,
                meta={
                    "file_id": file_id,
                    "filename": file_name,
                    "content_type": content_type,
                    "size": len(file_bytes),
                    "preview": preview,
                    "token_count": token_count,
                    "ocr_bridge_kind": file_kind,
                },
            )
        except ResearchUploadStoreUnavailable as exc:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Kho lưu trữ file research tạm thời không khả dụng.",
            ) from exc
        file_id = stored.file_id
    else:
        _store_uploaded_file(
            {
                "file_id": file_id,
                "filename": file_name,
                "content_type": content_type,
                "size": len(file_bytes),
                "created_at": created_at,
                "owner_user_id": user.id,
                "text": stored_text,
                "preview": preview,
                "token_count": token_count,
            }
        )

    default_source = _get_or_create_default_source(db, user.id)
    document = KnowledgeDocument(
        source_id=default_source.id,
        owner_user_id=user.id,
        filename=file_name,
        content_type=content_type,
        size=len(file_bytes),
        extracted_text=extracted_text if file_kind == "text" else "",
        preview=preview,
        token_count=token_count,
        is_active=True,
    )
    db.add(document)
    db.commit()

    return {
        "file_id": file_id,
        "preview": preview,
        "token_count": token_count,
        "metadata": {
            "filename": file_name,
            "size": len(file_bytes),
            "created_at": created_at,
            "knowledge_source_id": default_source.id,
        },
    }


@router.post("/tier2")
def research_tier2(
    payload: dict[str, Any],
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    settings = get_settings()
    user = _get_user_by_token(db, token)
    query_text = _extract_tier2_query_text(payload)
    normalized_input = dict(payload)
    # clara-research R15.2: enforce "never (fast && personal)" on the raw-dict surface too.
    _enforce_never_fast_and_personal(normalized_input)
    if query_text:
        normalized_input["query"] = query_text
        normalized_input["message"] = query_text
    upstream_payload = _build_tier2_upstream_payload(
        normalized_input, db=db, user=user, token=token
    )

    response = proxy_ml_post(
        "/v1/research/tier2",
        upstream_payload,
        # Keep the synchronous surface fail-closed as well; callers receive a
        # structured upstream error, never a synthetic research answer.
        fail_soft_payload=None,
        timeout_seconds=resolve_sync_research_timeout(settings.ml_research_timeout_seconds),
    )
    normalized = _normalize_tier2_response(response)
    normalized = _enforce_request_execution_contract(
        normalized,
        request_payload=upstream_payload,
    )
    attributed = _attach_research_attribution(normalized)
    # Keep the synchronous endpoint on the same evidence-release boundary as
    # the durable job worker.  Without this step, a caller could receive
    # factual-looking prose from ``POST /tier2`` even when the ML verifier had
    # marked one or more claims unsupported or contradicted.  The gate retains
    # provenance and verifier diagnostics, but replaces the conclusion with an
    # explicit abstention; it never manufactures a citation or downgrades the
    # FIDES/safety override result.
    gated = _apply_research_quality_gates(
        attributed,
        request_payload=upstream_payload,
    )
    gated = _attach_verified_research_presentation(
        gated,
        request_payload=upstream_payload,
    )
    return _apply_role_gated_telemetry(gated, role=token.role) or gated


@router.post("/clarify")
def research_clarify(
    payload: ResearchClarifyRequest,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> ResearchClarifyResponse:
    """Return clarifying questions for an ambiguous deep-research query (clara-research R12.1).

    Gated on ``RESEARCH_CLARIFYING_QUESTIONS_ENABLED`` and research mode ∈ {deep, deep_beta}.
    When the flag is off, the mode is not deep/deep_beta, or the query is unambiguous, the
    endpoint reports ``ambiguous=false`` with no questions so the UI starts without prompting.
    """
    settings = get_settings()
    # Authenticate the caller against a real user, consistent with the other research endpoints.
    _get_user_by_token(db, token)

    research_mode = _normalize_research_mode_value(payload.research_mode, default="deep")
    query_text = _extract_tier2_query_text(payload.model_dump())

    gate_open = (
        bool(settings.research_clarifying_questions_enabled)
        and research_mode in _CLARIFY_DEEP_MODES
    )
    if not gate_open or not query_text:
        return ResearchClarifyResponse(
            ambiguous=False, research_mode=research_mode, questions=[]
        )

    ambiguous = _detect_query_ambiguity(query_text)
    questions = (
        _build_clarifying_questions(ui_language=payload.ui_language) if ambiguous else []
    )
    return ResearchClarifyResponse(
        ambiguous=ambiguous,
        research_mode=research_mode,
        questions=questions,
    )


@router.post("/tier2/jobs")
def create_research_tier2_job(
    payload: ResearchTier2JobCreateRequest,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> ResearchTier2JobResponse:
    user = _get_user_by_token(db, token)
    input_payload = payload.model_dump()
    query_text = _extract_tier2_query_text(input_payload)
    if not query_text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="query không được rỗng.",
        )
    user_active_jobs = db.scalar(
        select(func.count())
        .select_from(ResearchJob)
        .where(
            ResearchJob.user_id == user.id,
            ResearchJob.status.in_(("queued", "running")),
        )
    )
    if int(user_active_jobs or 0) >= _RESEARCH_JOB_MAX_ACTIVE_PER_USER:
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail=(
                "Bạn đang có quá nhiều research job đang xử lý. "
                "Vui lòng chờ job hiện tại hoàn tất rồi thử lại."
            ),
        )
    if _count_pending_research_jobs() >= _RESEARCH_JOB_MAX_PENDING:
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail="Hệ thống đang bận xử lý research jobs. Vui lòng thử lại sau ít phút.",
        )

    input_payload["query"] = query_text
    input_payload["message"] = query_text
    upstream_payload = _build_tier2_upstream_payload(input_payload, db=db, user=user, token=token)
    persisted_payload = _sanitize_provider_secrets(upstream_payload)

    now = datetime.now(tz=UTC)
    job_id = uuid4().hex
    job = ResearchJob(
        job_id=job_id,
        user_id=user.id,
        role=token.role,
        status="queued",
        query_text=query_text,
        request_payload=persisted_payload,
        progress_json=_empty_job_progress(),
        result_json=None,
        run_manifest_json=_build_research_run_manifest(
            job_id=job_id,
            request_payload=persisted_payload,
            created_at=now,
        ),
        evidence_snapshot_json=None,
        error_text="",
        created_at=now,
        updated_at=now,
    )
    db.add(job)
    db.commit()
    db.refresh(job)
    try:
        _queue_research_job(job_id)
    except RuntimeError as exc:
        db.delete(job)
        db.commit()
        if str(exc) == "research_job_queue_full":
            raise HTTPException(
                status_code=status.HTTP_429_TOO_MANY_REQUESTS,
                detail="Hệ thống đang bận xử lý research jobs. Vui lòng thử lại sau ít phút.",
            ) from exc
        raise
    _append_job_event(
        db,
        job=job,
        stage="queue",
        status_text="completed",
        note="Đã tạo research job. Chuẩn bị chạy truy xuất chuyên sâu.",
    )
    db.refresh(job)
    return _serialize_research_job(job, role=token.role)


@router.get("/tier2/jobs/{job_id}")
def get_research_tier2_job(
    job_id: str,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> ResearchTier2JobResponse:
    user = _get_user_by_token(db, token)
    job = db.execute(
        select(ResearchJob).where(
            ResearchJob.job_id == job_id,
            ResearchJob.user_id == user.id,
        )
    ).scalar_one_or_none()
    if job is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Research job không tồn tại.",
        )
    return _serialize_research_job(job, role=token.role)


_RESEARCH_EXPORT_FORMATS = ("md", "docx", "pdf")
_RESEARCH_EXPORT_MEDIA_TYPES = {
    "md": "text/markdown; charset=utf-8",
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ),
    "pdf": "application/pdf",
}


def _research_export_enabled() -> bool:
    """Return True when the export surface is enabled (R16, default-off)."""

    return bool(get_settings().research_export_enabled)


def _export_report_body(result: dict[str, Any]) -> str:
    """Resolve the report body markdown/text from a stored tier2 result."""

    for key in ("answer_markdown", "answer", "summary", "message"):
        candidate = result.get(key)
        if isinstance(candidate, str) and candidate.strip():
            return candidate.strip()
    return ""


def _export_citations(result: dict[str, Any]) -> list[dict[str, Any]]:
    """Normalize the citations list into dicts for rendering (R16.2)."""

    raw_citations = result.get("citations")
    if not isinstance(raw_citations, list):
        return []

    citations: list[dict[str, Any]] = []
    for index, citation in enumerate(raw_citations, start=1):
        if isinstance(citation, dict):
            citation_id = (
                citation.get("source_id")
                or citation.get("id")
                or citation.get("citation_id")
                or f"c{index}"
            )
            citations.append(
                {
                    "citation_id": str(citation_id),
                    "title": citation.get("title"),
                    "source": citation.get("source"),
                    "url": citation.get("url"),
                    "study_id": citation.get("study_id"),
                    "source_type": citation.get("source_type"),
                    "trust_tier": citation.get("trust_tier"),
                    "published_at": citation.get("published_at"),
                }
            )
        elif isinstance(citation, str) and citation.strip():
            citations.append({"citation_id": f"c{index}", "title": citation.strip()})
    return citations


def _export_citation_registry(
    result: dict[str, Any],
    citations: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Resolve the Citation Registry appendix, deriving it from citations when absent.

    The exported report ALWAYS includes a Citation Registry (R16.2); when the ML
    layer did not emit an explicit ``citation_registry`` it is reconstructed from
    the report citations so the appendix is never empty for a cited report.
    """

    registry = result.get("citation_registry")
    if isinstance(registry, list) and registry:
        normalized: list[dict[str, Any]] = []
        for index, entry in enumerate(registry, start=1):
            if isinstance(entry, dict):
                citation_id = entry.get("citation_id") or entry.get("id") or f"c{index}"
                merged = dict(entry)
                merged["citation_id"] = str(citation_id)
                normalized.append(merged)
            elif isinstance(entry, str) and entry.strip():
                normalized.append({"citation_id": f"c{index}", "study_id": entry.strip()})
        if normalized:
            return normalized

    # Derive from citations so the appendix is always present (R16.2).
    return [
        {
            "citation_id": citation["citation_id"],
            "study_id": citation.get("study_id"),
            "title": citation.get("title"),
            "url": citation.get("url"),
            "source": citation.get("source"),
            "source_type": citation.get("source_type"),
            "trust_tier": citation.get("trust_tier"),
            "published_at": citation.get("published_at"),
        }
        for citation in citations
    ]


def _format_citation_line(entry: dict[str, Any]) -> str:
    """Render a single citation/registry entry as a one-line summary."""

    label = entry.get("title") or entry.get("source") or entry.get("study_id") or "Nguồn"
    parts = [str(label).strip()]
    detail_bits: list[str] = []
    for key in ("source", "study_id", "source_type"):
        value = entry.get(key)
        if value not in (None, "") and str(value).strip():
            detail_bits.append(f"{key}={str(value).strip()}")
    trust_tier = entry.get("trust_tier")
    if trust_tier not in (None, ""):
        detail_bits.append(f"trust_tier={trust_tier}")
    published_at = entry.get("published_at")
    if published_at not in (None, "") and str(published_at).strip():
        detail_bits.append(f"date={str(published_at).strip()}")
    if detail_bits:
        parts.append(f"({', '.join(detail_bits)})")
    url = entry.get("url")
    if url not in (None, "") and str(url).strip():
        parts.append(str(url).strip())
    return " ".join(parts)


def _build_export_markdown(*, query_text: str, result: dict[str, Any]) -> str:
    """Build the canonical Markdown export, always including citations + registry."""

    citations = _export_citations(result)
    registry = _export_citation_registry(result, citations)

    title = (query_text or "Research Report").strip() or "Research Report"
    sections: list[str] = [f"# {title}"]

    body = _export_report_body(result)
    if body:
        sections.append(body)

    citation_lines = ["## Citations"]
    if citations:
        for index, citation in enumerate(citations, start=1):
            citation_lines.append(
                f"{index}. [{citation['citation_id']}] {_format_citation_line(citation)}"
            )
    else:
        citation_lines.append("_Không có trích dẫn cho báo cáo này._")
    sections.append("\n".join(citation_lines))

    registry_lines = ["## Citation Registry"]
    if registry:
        for entry in registry:
            registry_lines.append(f"- [{entry['citation_id']}] {_format_citation_line(entry)}")
    else:
        registry_lines.append("_Không có mục nào trong sổ trích dẫn._")
    sections.append("\n".join(registry_lines))

    return "\n\n".join(sections).strip() + "\n"


def _markdown_to_plain_lines(markdown_text: str) -> list[str]:
    """Reduce export markdown to plain text lines for PDF rendering."""

    plain_lines: list[str] = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.lstrip("#").strip() if line.startswith("#") else line
        plain_lines.append(stripped)
    return plain_lines


def _wrap_export_lines(lines: list[str], *, max_chars: int = 95) -> list[str]:
    wrapped: list[str] = []
    for line in lines:
        if not line:
            wrapped.append("")
            continue
        current = ""
        for word in line.split(" "):
            if not current:
                current = word
            elif len(current) + 1 + len(word) <= max_chars:
                current = f"{current} {word}"
            else:
                wrapped.append(current)
                current = word
        wrapped.append(current)
    return wrapped


def _render_export_docx(*, query_text: str, result: dict[str, Any]) -> bytes:
    """Render the report (with citations + registry) as a DOCX document."""

    from docx import Document

    citations = _export_citations(result)
    registry = _export_citation_registry(result, citations)

    document = Document()
    document.add_heading((query_text or "Research Report").strip() or "Research Report", level=0)

    body = _export_report_body(result)
    if body:
        for block in body.split("\n\n"):
            block = block.strip()
            if block:
                document.add_paragraph(block)

    document.add_heading("Citations", level=1)
    if citations:
        for citation in citations:
            document.add_paragraph(
                f"[{citation['citation_id']}] {_format_citation_line(citation)}",
                style="List Number",
            )
    else:
        document.add_paragraph("Không có trích dẫn cho báo cáo này.")

    document.add_heading("Citation Registry", level=1)
    if registry:
        for entry in registry:
            document.add_paragraph(
                f"[{entry['citation_id']}] {_format_citation_line(entry)}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("Không có mục nào trong sổ trích dẫn.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _render_export_pdf(*, query_text: str, result: dict[str, Any]) -> bytes:
    """Render the report as a minimal, dependency-free PDF (citations + registry)."""

    markdown_text = _build_export_markdown(query_text=query_text, result=result)
    lines = _wrap_export_lines(_markdown_to_plain_lines(markdown_text))

    lines_per_page = 52
    pages = [lines[i : i + lines_per_page] for i in range(0, len(lines), lines_per_page)] or [[]]

    # Object numbering: 1=Catalog, 2=Pages, 3=Font, then per page (page, contents).
    page_object_numbers: list[int] = []
    objects: dict[int, bytes] = {}
    next_object_number = 4
    for page_lines in pages:
        page_number = next_object_number
        contents_number = next_object_number + 1
        next_object_number += 2
        page_object_numbers.append(page_number)

        content = "BT\n/F1 11 Tf\n14 TL\n50 770 Td\n"
        for index, line in enumerate(page_lines):
            if index > 0:
                content += "T*\n"
            content += f"({_pdf_escape(line)}) Tj\n"
        content += "ET"
        content_bytes = content.encode("latin-1", errors="replace")

        objects[page_number] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {contents_number} 0 R >>"
        ).encode("latin-1")
        objects[contents_number] = (
            f"<< /Length {len(content_bytes)} >>\nstream\n".encode("latin-1")
            + content_bytes
            + b"\nendstream"
        )

    kids = " ".join(f"{number} 0 R" for number in page_object_numbers)
    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    objects[2] = (
        f"<< /Type /Pages /Kids [{kids}] /Count {len(page_object_numbers)} >>"
    ).encode("latin-1")
    objects[3] = (
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica "
        b"/Encoding /WinAnsiEncoding >>"
    )

    pdf = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for number in sorted(objects):
        offsets[number] = len(pdf)
        pdf += f"{number} 0 obj\n".encode("latin-1")
        pdf += objects[number]
        pdf += b"\nendobj\n"

    xref_offset = len(pdf)
    total_objects = len(objects) + 1
    pdf += f"xref\n0 {total_objects}\n".encode("latin-1")
    pdf += b"0000000000 65535 f \n"
    for number in sorted(objects):
        pdf += f"{offsets[number]:010d} 00000 n \n".encode("latin-1")
    pdf += (
        f"trailer\n<< /Size {total_objects} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF"
    ).encode("latin-1")
    return bytes(pdf)


@router.post("/tier2/jobs/{job_id}/export")
def export_research_tier2_job(
    job_id: str,
    export_format: str = FastAPIQuery("md", alias="format"),
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> Response:
    """Export a completed research report as md/docx/pdf (R16.1, R16.2, R16.4).

    Default-off behind ``RESEARCH_EXPORT_ENABLED``: when disabled the surface does
    not exist (404), preserving legacy behavior. Owner-isolated. Export is rejected
    until the report has completed (R16.4), and every artifact always includes the
    citations and the Citation Registry (R16.2).
    """

    if not _research_export_enabled():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Research export chưa được bật.",
        )

    normalized_format = str(export_format or "").strip().lower()
    if normalized_format not in _RESEARCH_EXPORT_FORMATS:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="format phải là 'md', 'docx' hoặc 'pdf'.",
        )

    user = _get_user_by_token(db, token)
    job = db.execute(
        select(ResearchJob).where(
            ResearchJob.job_id == job_id,
            ResearchJob.user_id == user.id,
        )
    ).scalar_one_or_none()
    if job is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Research job không tồn tại.",
        )

    if str(job.status or "").strip().lower() != "completed":
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Báo cáo chưa hoàn tất nên không thể export.",
        )

    result = job.result_json if isinstance(job.result_json, dict) else None
    if not result:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Báo cáo chưa hoàn tất nên không thể export.",
        )

    query_text = job.query_text or ""
    if normalized_format == "md":
        body = _build_export_markdown(query_text=query_text, result=result).encode("utf-8")
    elif normalized_format == "docx":
        body = _render_export_docx(query_text=query_text, result=result)
    else:  # pdf
        body = _render_export_pdf(query_text=query_text, result=result)

    filename = f"research_{job_id}.{normalized_format}"
    return Response(
        content=body,
        media_type=_RESEARCH_EXPORT_MEDIA_TYPES[normalized_format],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _research_share_enabled() -> bool:
    """Return True when the read-only share surface is enabled (R16.3, default-off)."""

    return bool(get_settings().research_share_enabled)


def _research_share_public_url(share_token: str) -> str:
    """Build the public ``/share/{token}`` URL (reuses the workspace mechanism)."""

    base = get_settings().auth_public_web_base_url.rstrip("/")
    return f"{base}/share/{share_token}"


def _research_share_token_hash(share_token: str) -> str:
    return hashlib.sha256(share_token.encode("utf-8")).hexdigest()


def _generate_research_share_token(db: Session) -> str:
    """Generate a unique share token, mirroring the workspace share mechanism."""

    for _ in range(8):
        candidate = secrets.token_urlsafe(24)
        exists = db.execute(
            select(WorkspaceConversationShare.id).where(
                WorkspaceConversationShare.token_hash == _research_share_token_hash(candidate)
            )
        ).scalar_one_or_none()
        if exists is None:
            return candidate
    raise HTTPException(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        detail="Không thể tạo share token.",
    )


def _serialize_research_share(
    share: WorkspaceConversationShare,
    *,
    job_id: str,
    issued_token: str | None = None,
) -> ResearchTier2ShareResponse:
    return ResearchTier2ShareResponse(
        job_id=job_id,
        share_id=share.id,
        share_token=issued_token,
        public_url=_research_share_public_url(issued_token) if issued_token else None,
        is_active=bool(share.is_active),
        expires_at=share.expires_at,
        created_at=share.created_at,
        updated_at=share.updated_at,
    )


@router.post("/tier2/jobs/{job_id}/share")
def share_research_tier2_job(
    job_id: str,
    payload: WorkspaceConversationShareCreateRequest | None = None,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> ResearchTier2ShareResponse:
    """Create (or rotate) a read-only share link for a research report (R16.3).

    Default-off behind ``RESEARCH_SHARE_ENABLED``: when disabled the surface does
    not exist (404), preserving legacy behavior. Owner-isolated. Reuses the
    ``WorkspaceConversationShare`` mechanism (``share_token`` + ``/share/{token}``
    public URL), keyed on the research job instead of a chat session.
    """

    if not _research_share_enabled():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Research share chưa được bật.",
        )

    request = payload or WorkspaceConversationShareCreateRequest()

    user = _get_user_by_token(db, token)
    job = db.execute(
        select(ResearchJob).where(
            ResearchJob.job_id == job_id,
            ResearchJob.user_id == user.id,
        )
    ).scalar_one_or_none()
    if job is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Research job không tồn tại.",
        )

    share = db.execute(
        select(WorkspaceConversationShare).where(
            WorkspaceConversationShare.user_id == user.id,
            WorkspaceConversationShare.research_job_id == job.id,
        )
    ).scalar_one_or_none()

    should_rotate = bool(request.rotate) or share is None
    issued_token: str | None = None
    if share is None:
        issued_token = _generate_research_share_token(db)
        share = WorkspaceConversationShare(
            user_id=user.id,
            session_id=None,
            research_job_id=job.id,
            token_hash=_research_share_token_hash(issued_token),
            is_active=True,
        )
    else:
        share.is_active = True
        if should_rotate:
            issued_token = _generate_research_share_token(db)
            share.token_hash = _research_share_token_hash(issued_token)

    share.expires_at = datetime.now(tz=UTC) + timedelta(hours=int(request.expires_in_hours))

    db.add(share)
    db.commit()
    db.refresh(share)
    return _serialize_research_share(share, job_id=job_id, issued_token=issued_token)


def _build_research_job_stream_headers() -> dict[str, str]:
    return {
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
        "X-Accel-Buffering": "no",
    }


def _sse_event(event_name: str, payload: dict[str, Any], event_id: str | None = None) -> str:
    data = json.dumps(payload, ensure_ascii=False, default=str)
    lines: list[str] = []
    if event_id:
        lines.append(f"id: {event_id}")
    lines.append(f"event: {event_name}")
    lines.append(f"data: {data}")
    return "\n".join(lines) + "\n\n"


@router.get("/tier2/jobs/{job_id}/stream")
async def stream_research_tier2_job(
    request: Request,
    job_id: str,
    heartbeat_seconds: int = 10,
    poll_interval_seconds: float = 0.8,
    token: TokenPayload = Depends(require_roles("normal", "researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
):
    user = _get_user_by_token(db, token)
    user_id = int(user.id)
    existing = db.execute(
        select(ResearchJob).where(
            ResearchJob.job_id == job_id,
            ResearchJob.user_id == user_id,
        )
    ).scalar_one_or_none()
    if existing is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Research job không tồn tại.",
        )

    safe_poll_interval = min(max(float(poll_interval_seconds), 0.25), 5.0)
    safe_heartbeat = 5 if heartbeat_seconds < 5 else heartbeat_seconds

    def _load_job_snapshot() -> dict[str, Any] | None:
        # Use a fresh session for each read to avoid stale identity-map cache
        # during long-lived SSE connections.
        with SessionLocal() as fresh_db:
            current = fresh_db.execute(
                select(ResearchJob).where(
                    ResearchJob.job_id == job_id,
                    ResearchJob.user_id == user_id,
                )
            ).scalar_one_or_none()
            if current is None:
                return None
            return _serialize_research_job(current, role=token.role).model_dump(mode="json")

    async def event_stream():
        last_heartbeat_at = time.monotonic()
        last_signature = ""
        sequence = 0
        yield ": connected\n\n"

        while True:
            if await request.is_disconnected():
                break

            snapshot = _load_job_snapshot()
            if snapshot is None:
                sequence += 1
                yield _sse_event(
                    "error",
                    {"message": "Research job không còn khả dụng."},
                    event_id=str(sequence),
                )
                break

            signature = json.dumps(
                {
                    "status": snapshot.get("status"),
                    "updated_at": snapshot.get("updated_at"),
                    "completed_at": snapshot.get("completed_at"),
                    "error": snapshot.get("error"),
                    "progress": snapshot.get("progress"),
                },
                ensure_ascii=False,
                sort_keys=True,
                default=str,
            )

            if signature != last_signature:
                last_signature = signature
                sequence += 1
                yield _sse_event("progress", snapshot, event_id=str(sequence))
                last_heartbeat_at = time.monotonic()

            status_text = str(snapshot.get("status") or "").lower()
            if status_text in {"completed", "failed"}:
                sequence += 1
                yield _sse_event("done", snapshot, event_id=str(sequence))
                break

            if time.monotonic() - last_heartbeat_at >= safe_heartbeat:
                yield ": keepalive\n\n"
                last_heartbeat_at = time.monotonic()

            await asyncio.sleep(safe_poll_interval)

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers=_build_research_job_stream_headers(),
    )


@router.get("/source-hub/catalog")
def source_hub_catalog(
    token: TokenPayload = Depends(require_roles("researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> dict[str, list[SourceHubCatalogEntry]]:
    _ = token
    return {"sources": _load_source_hub_catalog(db)}


@router.get("/source-hub/records")
def source_hub_records(
    source: str = "all",
    query: str = "",
    limit: int = 80,
    token: TokenPayload = Depends(require_roles("researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> SourceHubRecordsResponse:
    user = _get_user_by_token(db, token)
    items = _load_source_hub_records(db, user.id)

    normalized_source = source.strip().lower()
    if normalized_source and normalized_source != "all":
        items = [item for item in items if item.source == normalized_source]

    normalized_query = query.strip().lower()
    if normalized_query:
        items = [
            item
            for item in items
            if normalized_query in (item.title or "").lower()
            or normalized_query in (item.snippet or "").lower()
            or normalized_query in (item.query or "").lower()
        ]

    safe_limit = max(1, min(500, limit))
    return SourceHubRecordsResponse(records=items[:safe_limit])


@router.post("/source-hub/sync")
def source_hub_sync(
    payload: SourceHubSyncRequest,
    token: TokenPayload = Depends(require_roles("researcher", "doctor", "admin")),
    db: Session = Depends(get_db),
) -> SourceHubSyncResponse:
    user = _get_user_by_token(db, token)
    query = payload.query.strip()
    if not query:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="Query không được rỗng."
        )
    catalog_by_key = {entry.key: entry for entry in _load_source_hub_catalog(db)}
    selected_source = catalog_by_key.get(payload.source)
    if selected_source is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Source '{payload.source}' chưa được cấu hình trong catalog.",
        )
    if not selected_source.supports_live_sync:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Source '{payload.source}' đang tắt live sync.",
        )

    safe_limit = max(3, min(500, int(payload.limit)))
    try:
        records, warnings = _fetch_source_hub_records(payload.source, query, safe_limit)
    except httpx.TimeoutException:
        raise HTTPException(
            status_code=status.HTTP_504_GATEWAY_TIMEOUT,
            detail="Timeout khi đồng bộ dữ liệu từ nguồn ngoài.",
        ) from None
    except httpx.HTTPStatusError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Lỗi nguồn ngoài ({exc.response.status_code}) khi sync {payload.source}.",
        ) from exc
    except Exception as exc:  # pragma: no cover - defensive fallback
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Không thể sync nguồn {payload.source}: {exc}",
        ) from exc

    # Ensure query/sync metadata always present even if source omitted them.
    synced_at = datetime.now(tz=UTC).isoformat()
    normalized_records = [
        SourceHubRecord(
            **{
                **record.model_dump(),
                "query": record.query or query,
                "synced_at": record.synced_at or synced_at,
            }
        )
        for record in records
    ]

    existing = _load_source_hub_records(db, user.id)
    merged = _merge_source_hub_records(existing, normalized_records)
    _save_source_hub_records(db, user.id, merged)

    return SourceHubSyncResponse(
        source=payload.source,
        query=query,
        fetched=len(normalized_records),
        stored=len(merged),
        records=normalized_records,
        warnings=warnings,
    )
