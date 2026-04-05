from __future__ import annotations

import ast
import base64
import io
import re
from typing import Any

import httpx
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.token import Token

_MERMAID_RENDER_TIMEOUT_SECONDS = 8.0
_MERMAID_MAX_BYTES = 4_500_000
_MERMAID_INFO_PATTERN = re.compile(r"^mermaid(\s+|$)", flags=re.IGNORECASE)
_CHART_SPEC_INFO_PATTERN = re.compile(
    r"^(chart-spec|vega-lite|echarts-option)(\s+|$)",
    flags=re.IGNORECASE,
)
_UNICODE_BULLET_PATTERN = re.compile(r"^(\s*)[•●▪◦]\s+(.*)$")

_MERMAID_START_PREFIXES = (
    "flowchart",
    "graph ",
    "sequencediagram",
    "classdiagram",
    "statediagram",
    "erdiagram",
    "journey",
    "gantt",
    "pie",
    "mindmap",
    "timeline",
)


def build_docx_bytes_from_markdown(markdown_text: str) -> bytes:
    doc = Document()
    _configure_document_styles(doc)

    md = MarkdownIt("commonmark", {"html": False}).enable("table")
    tokens = md.parse(_normalize_markdown(markdown_text))
    _render_markdown_tokens(doc, tokens)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _configure_document_styles(doc: DocumentObject) -> None:
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal_style.font.size = Pt(11)


def _normalize_markdown(value: str) -> str:
    normalized = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    normalized = _normalize_unicode_bullets(normalized)
    return _auto_fence_special_blocks(normalized)


def _normalize_unicode_bullets(text: str) -> str:
    lines = text.split("\n")
    in_fence = False
    out: list[str] = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_fence = not in_fence
            out.append(line)
            continue
        if in_fence:
            out.append(line)
            continue

        match = _UNICODE_BULLET_PATTERN.match(line)
        if match:
            indent, content = match.groups()
            out.append(f"{indent}- {content}")
            continue
        out.append(line)
    return "\n".join(out)


def _auto_fence_special_blocks(text: str) -> str:
    lines = text.split("\n")
    out: list[str] = []
    in_fence = False
    prev_nonempty = ""
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_fence = not in_fence
            out.append(line)
            if stripped:
                prev_nonempty = stripped
            index += 1
            continue

        if in_fence:
            out.append(line)
            if stripped:
                prev_nonempty = stripped
            index += 1
            continue

        if _is_mermaid_start_line(stripped):
            block, next_index = _collect_nonempty_block(lines, index)
            out.append("```mermaid")
            out.extend(block)
            out.append("```")
            if next_index < len(lines) and lines[next_index].strip() == "":
                out.append(lines[next_index])
                next_index += 1
            prev_nonempty = "```mermaid"
            index = next_index
            continue

        if _is_chart_spec_start_line(stripped, prev_nonempty):
            block, next_index = _collect_nonempty_block(lines, index)
            out.append("```chart-spec")
            out.extend(block)
            out.append("```")
            if next_index < len(lines) and lines[next_index].strip() == "":
                out.append(lines[next_index])
                next_index += 1
            prev_nonempty = "```chart-spec"
            index = next_index
            continue

        out.append(line)
        if stripped:
            prev_nonempty = stripped
        index += 1

    return "\n".join(out)


def _collect_nonempty_block(lines: list[str], start: int) -> tuple[list[str], int]:
    block: list[str] = []
    cursor = start
    while cursor < len(lines):
        candidate = lines[cursor]
        stripped = candidate.strip()
        if not stripped:
            break
        if stripped.startswith("```"):
            break
        if cursor > start and stripped.startswith("#"):
            break
        block.append(candidate)
        cursor += 1
    return block, cursor


def _is_mermaid_start_line(value: str) -> bool:
    lowered = (value or "").strip().lower()
    if not lowered:
        return False
    return any(lowered.startswith(prefix) for prefix in _MERMAID_START_PREFIXES)


def _is_chart_spec_start_line(value: str, prev_nonempty: str) -> bool:
    lowered = (value or "").strip().lower()
    prev = (prev_nonempty or "").strip().lower()
    if not lowered.startswith("type:"):
        return False
    return "chart spec" in prev or "chart-spec" in prev


def _render_markdown_tokens(doc: DocumentObject, tokens: list[Token]) -> None:
    index = 0
    while index < len(tokens):
        token = tokens[index]
        token_type = token.type

        if token_type == "heading_open":
            level = _heading_level(token.tag)
            inline = tokens[index + 1] if index + 1 < len(tokens) else None
            paragraph = doc.add_heading(level=level)
            _append_inline_tokens(paragraph, inline.children if inline else [])
            index += 3
            continue

        if token_type == "paragraph_open":
            inline = tokens[index + 1] if index + 1 < len(tokens) else None
            paragraph = doc.add_paragraph()
            _append_inline_tokens(paragraph, inline.children if inline else [])
            index += 3
            continue

        if token_type in {"bullet_list_open", "ordered_list_open"}:
            ordered = token_type == "ordered_list_open"
            index = _render_list_block(doc, tokens, index, ordered=ordered)
            continue

        if token_type == "fence":
            info = (token.info or "").strip()
            code = token.content or ""
            if _MERMAID_INFO_PATTERN.match(info):
                _append_mermaid_block(doc, code)
            elif _CHART_SPEC_INFO_PATTERN.match(info):
                _append_chart_spec_block(doc, code)
            else:
                _append_code_block(doc, code)
            index += 1
            continue

        if token_type == "table_open":
            index = _render_table_block(doc, tokens, index)
            continue

        if token_type == "hr":
            doc.add_paragraph("────────────────────────")
            index += 1
            continue

        if token_type in {"blockquote_open", "blockquote_close", "paragraph_close"}:
            index += 1
            continue

        index += 1


def _heading_level(tag: str) -> int:
    try:
        parsed = int(str(tag or "h2").replace("h", ""))
    except ValueError:
        parsed = 2
    return max(1, min(parsed, 4))


def _render_list_block(
    doc: DocumentObject,
    tokens: list[Token],
    start_index: int,
    *,
    ordered: bool,
) -> int:
    closing_type = "ordered_list_close" if ordered else "bullet_list_close"
    style_name = "List Number" if ordered else "List Bullet"

    index = start_index + 1
    while index < len(tokens):
        token = tokens[index]
        if token.type == closing_type:
            return index + 1

        if token.type == "list_item_open":
            item_inline: Token | None = None
            cursor = index + 1
            while cursor < len(tokens) and tokens[cursor].type != "list_item_close":
                current = tokens[cursor]
                if current.type == "inline":
                    item_inline = current
                    break
                cursor += 1
            paragraph = doc.add_paragraph(style=style_name)
            _append_inline_tokens(paragraph, item_inline.children if item_inline else [])
            index = cursor

        index += 1

    return index


def _render_table_block(doc: DocumentObject, tokens: list[Token], start_index: int) -> int:
    rows: list[list[str]] = []
    current_row: list[str] = []
    index = start_index + 1

    while index < len(tokens):
        token = tokens[index]
        token_type = token.type

        if token_type == "table_close":
            break

        if token_type == "tr_open":
            current_row = []
            index += 1
            continue

        if token_type == "tr_close":
            if current_row:
                rows.append(current_row)
            current_row = []
            index += 1
            continue

        if token_type in {"th_open", "td_open"}:
            cell_value = ""
            cursor = index + 1
            while cursor < len(tokens) and tokens[cursor].type not in {"th_close", "td_close"}:
                if tokens[cursor].type == "inline":
                    cell_value = _flatten_inline_text(tokens[cursor].children or [])
                cursor += 1
            current_row.append(cell_value.strip())
            index = cursor + 1
            continue

        index += 1

    if not rows:
        return index + 1

    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            table.cell(row_index, column_index).text = value

    return index + 1


def _append_inline_tokens(paragraph: Any, children: list[Token] | None) -> None:
    if not children:
        return

    bold = False
    italic = False
    link_href: str | None = None

    for child in children:
        token_type = child.type

        if token_type == "strong_open":
            bold = True
            continue
        if token_type == "strong_close":
            bold = False
            continue
        if token_type == "em_open":
            italic = True
            continue
        if token_type == "em_close":
            italic = False
            continue
        if token_type == "link_open":
            link_href = _token_attr(child, "href")
            continue
        if token_type == "link_close":
            link_href = None
            continue
        if token_type in {"softbreak", "hardbreak"}:
            paragraph.add_run("\n")
            continue
        if token_type == "code_inline":
            run = paragraph.add_run(child.content or "")
            run.bold = bold
            run.italic = italic
            run.font.name = "Consolas"
            run.font.size = Pt(10.5)
            continue
        if token_type != "text":
            continue

        run = paragraph.add_run(child.content or "")
        run.bold = bold
        run.italic = italic
        if link_href:
            run.underline = True
            run.font.color.rgb = RGBColor(0x05, 0x61, 0xC3)


def _flatten_inline_text(children: list[Token]) -> str:
    chunks: list[str] = []
    for child in children:
        if child.type == "text":
            chunks.append(child.content or "")
        elif child.type == "code_inline":
            chunks.append(child.content or "")
        elif child.type in {"softbreak", "hardbreak"}:
            chunks.append(" ")
    return "".join(chunks)


def _append_code_block(doc: DocumentObject, code: str) -> None:
    lines = (code or "").splitlines() or [""]
    for line in lines:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10.5)


def _append_chart_spec_block(doc: DocumentObject, code: str) -> None:
    parsed = _parse_chart_spec(code)

    title = doc.add_paragraph(parsed.get("title") or "Chart Spec")
    title.runs[0].bold = True

    labels = parsed.get("x") or []
    values = parsed.get("y") or []
    row_count = min(len(labels), len(values))
    if row_count:
        table = doc.add_table(rows=row_count + 1, cols=3)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        table.cell(0, 2).text = "Bar"

        numeric_values = [item for item in (_coerce_float(v) for v in values) if item is not None]
        max_value = max((abs(item) for item in numeric_values), default=0.0)

        for idx in range(row_count):
            label = str(labels[idx])
            value = values[idx]
            numeric = _coerce_float(value)
            bar = ""
            if numeric is not None and max_value > 0:
                bar_units = max(1, int(round((abs(numeric) / max_value) * 20)))
                bar = "█" * bar_units
            table.cell(idx + 1, 0).text = label
            table.cell(idx + 1, 1).text = str(value)
            table.cell(idx + 1, 2).text = bar
        return

    _append_code_block(doc, code)


def _parse_chart_spec(code: str) -> dict[str, Any]:
    lines = [line.rstrip() for line in (code or "").splitlines() if line.strip()]
    parsed: dict[str, Any] = {"title": "Chart Spec", "x": [], "y": []}
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if ":" not in line:
            index += 1
            continue
        key, raw_value = line.split(":", 1)
        normalized_key = key.strip().lower()
        value = raw_value.strip()

        if normalized_key == "title" and value:
            parsed["title"] = value
            index += 1
            continue
        if normalized_key == "x":
            parsed["x"] = _parse_inline_list(value)
            index += 1
            continue
        if normalized_key == "y":
            if value.startswith("["):
                parsed["y"] = _parse_inline_list(value)
                index += 1
                continue

            y_values: list[Any] = []
            cursor = index + 1
            while cursor < len(lines):
                candidate = lines[cursor].strip()
                if not candidate.startswith("-"):
                    break
                y_values.append(_parse_scalar(candidate[1:].strip()))
                cursor += 1
            parsed["y"] = y_values
            index = cursor
            continue

        index += 1
    return parsed


def _parse_inline_list(raw: str) -> list[Any]:
    value = (raw or "").strip()
    if not value:
        return []
    try:
        parsed = ast.literal_eval(value)
    except Exception:
        parsed = None
    if isinstance(parsed, list):
        return [_parse_scalar(item) for item in parsed]

    if value.startswith("[") and value.endswith("]"):
        inner = value[1:-1]
    else:
        inner = value
    items = [part.strip() for part in inner.split(",") if part.strip()]
    return [_parse_scalar(item) for item in items]


def _parse_scalar(raw: Any) -> Any:
    if isinstance(raw, (int, float)):
        return raw
    text = str(raw).strip().strip("'\"")
    if not text:
        return ""
    number = _coerce_float(text)
    return number if number is not None else text


def _coerce_float(value: Any) -> float | None:
    try:
        return float(str(value).strip())
    except Exception:
        return None


def _append_mermaid_block(doc: DocumentObject, code: str) -> None:
    title = doc.add_paragraph("Mermaid Diagram")
    title.runs[0].bold = True

    png_bytes = _render_mermaid_png(code)
    if png_bytes:
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.2))
        return

    warning = doc.add_paragraph("[Mermaid render unavailable - included as code]")
    warning.runs[0].italic = True
    _append_code_block(doc, code)


def _render_mermaid_png(code: str) -> bytes | None:
    payload = (code or "").strip()
    if not payload:
        return None

    encoded = base64.urlsafe_b64encode(payload.encode("utf-8")).decode("ascii").rstrip("=")
    url = f"https://mermaid.ink/img/{encoded}"

    try:
        with httpx.Client(timeout=_MERMAID_RENDER_TIMEOUT_SECONDS, follow_redirects=True) as client:
            response = client.get(url)
            if response.status_code != 200:
                return None
            content_type = str(response.headers.get("content-type") or "").lower()
            if not content_type.startswith("image/"):
                return None
            data = response.content
            if not data or len(data) > _MERMAID_MAX_BYTES:
                return None
            return data
    except Exception:
        return None


def _token_attr(token: Token, key: str) -> str | None:
    attrs = token.attrs or {}
    if isinstance(attrs, dict):
        value = attrs.get(key)
        return str(value) if value is not None else None
    if isinstance(attrs, list):
        for entry in attrs:
            if isinstance(entry, (list, tuple)) and len(entry) == 2 and entry[0] == key:
                return str(entry[1])
    return None
