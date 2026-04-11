#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def build_reference_docx(output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]

    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(12)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    for style_name, size in [
        ("Title", 18),
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 13),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True

    header_para = section.header.paragraphs[0]
    header_para.text = "CLARA-Care | Báo cáo kỹ thuật và thuyết minh đề tài"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header_para.runs:
        header_para.runs[0].font.name = "Times New Roman"
        header_para.runs[0].font.size = Pt(10)

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Trang ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    add_field(footer_para, "PAGE")
    run = footer_para.add_run(" / ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    add_field(footer_para, "NUMPAGES")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Build pandoc reference DOCX (A4/header/footer).")
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    build_reference_docx(args.output)
    print(f"Created reference DOCX: {args.output}")


if __name__ == "__main__":
    main()
