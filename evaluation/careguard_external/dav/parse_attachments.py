"""Extract attachment text with explicit failures, never fabricated table rows."""
from __future__ import annotations

import argparse
import json
import subprocess
from pathlib import Path

from .core import Store


def parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("parser_dependency_unavailable:python-docx") from exc
    document = Document(path)
    lines = [paragraph.text for paragraph in document.paragraphs]
    lines.extend("\t".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows)
    return "\n".join(lines)


def parse_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("parser_dependency_unavailable:openpyxl") from exc
    workbook = load_workbook(path, read_only=True, data_only=False)
    try:
        return "\n".join(
            "\t".join("" if value is None else str(value) for value in row)
            for worksheet in workbook.worksheets
            for row in worksheet.iter_rows(values_only=True)
        )
    finally:
        workbook.close()


def parse(root: Path, retrieval_date: str) -> list[dict]:
    store = Store(root, retrieval_date); rows = [json.loads(line) for line in (store.manifests / "source_inventory.jsonl").read_text(encoding="utf-8").splitlines()]
    results = []
    for row in rows:
        path = Path(row["raw_path"])
        if "attachments" not in path.parts: continue
        result = {"raw_artifact_sha256": row["sha256"], "source_url": row["url"], "document_type": row.get("document_type"), "text": None, "parse_status": "UNPARSED", "parse_error": None}
        try:
            # pdftotext preserves Unicode and is available on common research hosts.
            if path.suffix.lower() == ".pdf":
                completed = subprocess.run(["pdftotext", "-layout", str(path), "-"], capture_output=True, check=True, timeout=120)
                result["text"] = completed.stdout.decode("utf-8", "replace"); result["parse_status"] = "PARSED"
            elif path.suffix.lower() == ".csv":
                result["text"] = path.read_text(encoding="utf-8-sig"); result["parse_status"] = "PARSED"
            elif path.suffix.lower() == ".docx":
                result["text"] = parse_docx(path); result["parse_status"] = "PARSED"
            elif path.suffix.lower() == ".xlsx":
                result["text"] = parse_xlsx(path); result["parse_status"] = "PARSED"
            else: result["parse_error"] = f"no_parser_for:{path.suffix.lower() or 'no_extension'}"
            if result["parse_status"] == "PARSED" and not result["text"].strip():
                result["parse_status"] = "PARSED_EMPTY"; result["parse_error"] = "no_extractable_text"
        except Exception as exc:  # noqa: BLE001 - persist parser failures without inventing rows
            result["parse_error"] = str(exc)
        results.append(result)
        if result["parse_status"] != "PARSED": store.append("failed_or_unavailable.jsonl", {"raw_artifact_sha256": row["sha256"], "url": row["url"], "reason": result["parse_error"]})
    output = store.manifests / "attachment_parse.jsonl"; output.write_text("".join(json.dumps(item, ensure_ascii=False, sort_keys=True) + "\n" for item in results), encoding="utf-8")
    return results

if __name__ == "__main__":
    parser = argparse.ArgumentParser(); parser.add_argument("--root", type=Path, default=Path("data/restricted/dav")); parser.add_argument("--date", required=True); args = parser.parse_args(); print(len(parse(args.root, args.date)))
